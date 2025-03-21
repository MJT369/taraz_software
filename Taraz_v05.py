# -*- coding: utf-8 -*-
"""
Name: Taraz Software
Version: 5.0.0
Author: MJT369
https://www.apple.com
Date Created: 2025
Description: text edit and translator and books search_internet etc .
Email : mj.taraz@yahoo.com
"""
print('بسم الله الرحمن الرحیم')
import datetime
import sys
import os
try:
 import pyperclip
except:
    pass
import re
from PyQt6.QtWidgets import (
    QAbstractItemView, QApplication, QDialogButtonBox, QHBoxLayout, QMainWindow, QLabel, QPushButton, QComboBox, QSpinBox, QStyle, QTableWidget, QTableWidgetItem,
    QTextEdit, QFileDialog, QMessageBox, QCheckBox, QInputDialog, QTabWidget,
    QDialog, QVBoxLayout, QColorDialog, QLineEdit, QWidget,
)
from PyQt6.QtGui import QFont, QTextCursor, QTextCharFormat, QBrush, QColor, QIcon, QTextDocument, QAction
from PyQt6.QtCore import Qt,QTimer
from functools import lru_cache
class lz:
    def __init__(self, module_name):
        self.module_name = module_name
        self.module = None
    def __getattr__(self, name):
        if self.module is None:
            self.module = __import__(self.module_name, fromlist=[name])
        return getattr(self.module, name)
Sequence = lz("difflib")
langdetec = lz("langdetect")
pdf2= lz("pdf2docx")
dox = lz("docx")
SpellChecke = lz("spellchecker")
Translator1 = lz("translatepy")
Translator2 = lz("argostranslate.translate")
package=lz('argostranslate')
Translator5 = lz("deep_translator")
gc = lz("gc")
shutil = lz("shutil")
farsi_tool = lz("farsi_tools")
Speller = lz("autocorrect")
requests = lz("requests")
urllib = lz("urllib.parse")
hashlib = lz("hashlib")
unicodedata = lz("unicodedata")
zip = lz("zipfile")
ftfy = lz("ftfy")
pd=lz("pandas")
camelot=lz("camelot")
npyxl=lz("openpyxl")
pypdf=lz("pypdf")

def importer():
    import pypdf
    import openpyxl
    import camelot
    import pandas
    import zipfile
    from difflib import SequenceMatcher
    import re
    from langdetect import detect
    import argostranslate
    import plistlib
    import pymupdf # type: ignore
    import argostranslate.translate
    import argostranslate.package
    from pdf2docx import Converter
    import ghostscript # type: ignore
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from spellchecker import SpellChecker
    import translatepy
    from translatepy.translators.bing import BingTranslate
    import deep_translator
    import threading
    import time
    import requests
    import hashlib
    import gc
    import shutil
    from farsi_tools import replace_ascii_digits_with_farsi,stop_words
    from autocorrect import Speller
class a(QMainWindow):
    def __init__(self):
        super().__init__()
        self.nam = "Taraz Software 313                                                                                                                    بِسْمِ اللهِ الرَّحْمنِ الرَّحِیم"
        self.clr4 = '#d3ecfa'
        self.font1 = 'Arial'
        self.size1 = 12
        self.text_font = QFont(self.font1, self.size1)
        self.num = 'default'
        self.lang_in1 = 'English'
        self.lang_out1 = 'فارسی'
        self.dir = os.getcwd()
        self.bookdir = os.path.join(self.dir, 'book_search')
        self.tab_widget = QTabWidget()
        self.thread = None
        self.worker = None
        self.dc=0
        self.translation_cache = {}
        self.thm()
        self.initUI()
    def initUI(self):
        self.setWindowTitle(self.nam)
        self.setGeometry(40, 40, 1280, 660)
        try:
            self.setStyleSheet(open(self.them1).read())
        except FileNotFoundError:
            pass
        self.clr = "background-color: #e2e8fa; color: blue;"
        self.qq = QMessageBox.question
        self.q = QMessageBox
        self.qb = QPushButton
        self.qc = QComboBox
        self.file_button = self.qb('انتخاب فایل', self)
        self.file_button.clicked.connect(self.select_file)
        self.file_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogOpenButton))
        self.translate_button = self.qb('ویرایش و ترجمه متن', self)
        self.translate_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.translate_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.translate_button.clicked.connect(self.translate)

        self.learn_button = self.qb('آموزش نیمه خودکار مترجم', self)
        self.learn_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.learn_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.learn_button.clicked.connect(self.semi_learn)

        self.translate_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay))
        self.trans_file_button = self.qb('ویرایش و ترجمه فایل', self)
        self.trans_file_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.trans_file_button.clicked.connect(self.trans_file)
        self.pdf_converting = self.qb('تبدیل PDF', self)
        self.pdf_converting.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.pdf_converting.clicked.connect(self.pdf_converter)
        self.export_button = self.qb('ذخیره ', self)
        self.export_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.export_button.clicked.connect(self.export_docx)
        self.export_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DialogSaveButton))
        self.clear_button = self.qb('پاک کردن متن', self)
        self.clear_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.clear_button.clicked.connect(self.clear)
        self.clear_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_TrashIcon))
        self.ch_tran_lang = self.qb('\u2190  \n  \u2192', self)
        self.ch_tran_lang.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.ch_tran_lang.clicked.connect(self.change_lang_translate)
        self.BS = self.qb('جستجو در کتاب', self)
        self.BS.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.BS.clicked.connect(self.run_searching)
        self.size_box = QSpinBox(self)
        self.size_box.setRange(0, 100)
        self.output_console = QTextEdit(self)
        self.input_console = QTextEdit(self)
        self.input_console.setPlaceholderText('متن خود  یا فایل  را اینجا وارد کنید ')
        self.input_console.setAcceptDrops(True)
        self.input_console.dragEnterEvent = self.dragEnterEvent
        self.input_console.dropEvent = self.dropEvent
        self.create_translation_tab()
        self.create_settings_tab()
        self.fileMode=False
        try:
            self.file_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
            self.input_console.setFont(QFont(self.font1, int(self.size1)))
            self.output_console.setFont(QFont(self.font1, int(self.size1)))
            self.size_box.setValue(int(self.size1))
            self.input_console.textChanged.connect(self.translate_starter)
        except:
            self.size_box.setValue(int(0))
            self.output_console.setFont(QFont(self.font1, int(0)))
            self.input_console.setFont(QFont(self.font1, int(0)))
            self.file_button.setFont(QFont('Arial', 14 + int(0), QFont.Weight.Bold))    
        self.size_box.valueChanged.connect(self.update_font)
        self.size_box.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.size_app = QSpinBox(self)
        self.size_app.setRange(0, 14)
        try:
          self.size_app.setValue(int(self.app_size)) 
        except: 
            self.app_size=2
            self.size_app.setValue(int(self.app_size))  
        self.size_app.valueChanged.connect(self.update_font2)
        self.size_app.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.book_box = QComboBox(self)
        self.book_box.addItems(["کتاب", "قرآن", 'نهج البلاغه', "اصول کافي", "نهج الفصاحه", "موعظه", "سه دقیقه در قیامت", "تمنای وصال", 'شعر تمنای وصال'])
        self.book_box.setFont(QFont('Arial', 11 , QFont.Weight.Bold))
        self.book_box.currentTextChanged.connect(self.gift)
        self.font_box = QComboBox(self)
        self.font_box.addItems(["Arial", 'Arial (Arabic)', 'Simplified Arabic Fixed', 'Courier New (Arabic)', 'Urdu Typesetting', 'Sakkal Majalla', 'Simplified Arabic', 'Traditional Arabic'])
        self.font_box.setFont(QFont('Arial', 11 , QFont.Weight.Bold))
        self.font_box.setCurrentText(self.font1)
        self.font_box.currentTextChanged.connect(self.update_font)
        self.stop_button = self.qb('توقف پردازش', self)
        self.stop_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.stop_button.clicked.connect(self.stop)
        self.help_button = self.qb('راهنما', self)
        self.help_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.help_button.clicked.connect(self.show_help_message)
        self.lang_button = self.qb('فارسی', self)
        self.lang_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.lang_button.clicked.connect(self.Lang)
        self.lang2_button = self.qb('English', self)
        self.lang2_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.lang2_button.clicked.connect(self.Lang2)
        self.output_console.setStyleSheet(self.clr)
        self.copy_button = self.qb('کپی', self)
        self.copy_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.copy_button.clicked.connect(self.copy_to_clipboard)
        self.copy_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogDetailedView))
        self.paste_button = self.qb(' جایگذاری', self)
        self.paste_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.paste_button.clicked.connect(self.paste_from_clipboard)
        self.paste_button.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_TitleBarNormalButton))
        self.reset_button = self.qb('بازنشانی', self)
        self.reset_button.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.reset_button.clicked.connect(self.restart_program1)
        self.lang_select()
        self.input_text_label = QLabel(self.M200, self)
        self.input_text_label.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.output_text_label = QLabel(self.M202, self)
        self.space_corect = QCheckBox(self.M11, self)
        self.space_corect.stateChanged.connect(self.aktive_space)
        self.learned_trance = QCheckBox(self.M44, self)
        self.learned_trance.stateChanged.connect(self.aktive_learn)
        self.semi_correct = QCheckBox(self.M13, self)
        self.semi_correct.stateChanged.connect(self.semi_corrections_state)
        self.virastar = QCheckBox(self.M94, self)
        self.auto_correction = QCheckBox(self.M15, self)
        self.auto_correction.setFont(QFont('Arial', 11 , QFont.Weight.Bold))
        self.auto_correction.stateChanged.connect(self.auto_corrections_state)
        self.virastar.stateChanged.connect(self.aktive_virast)
        self.source_language_label = QLabel(self.M002, self)
        self.source_language_label.setFont(QFont('Arial', 11 , QFont.Weight.Bold))
        self.target_language_label = QLabel(self.M003, self)
        self.target_language_label.setFont(QFont('Arial', 11 , QFont.Weight.Bold))
        self.translator_menu = self.qc(self)
        self.translator_menu.addItems([self.M118, self.M122, self.M119, self.M201, self.M123, self.M121, self.M259, self.M260, self.M261, self.M120])
        self.translator_menu.setFont(QFont('Arial', 13 , QFont.Weight.Bold))
        self.translator_menu.setCurrentText(self.translator1)
        self.translator_menu.currentTextChanged.connect(self.lang_code)
        self.them_num = self.qc(self)
        self.them_num.addItems([
            'Defult.CSS', 'VisualScript.qss', 'None', 'Irrorater.qss', 'Darkeum.qss',
            'Filmovio.qss', 'Gravira.qss', 'Fibers.qss', 'Diffnes.qss', 'Adaptic.qss',
            'Flower.css', 'Lovely.css'
        ])
        self.them_num.setFont(QFont('Arial', 13 ))
        self.them_num.setCurrentText(self.them1)
        self.them_num.currentTextChanged.connect(self.them)
        self.document_mode_box = self.qc(self)
        self.document_mode_box.addItems([self.M96, self.M98, self.M99, self.M97])
        self.document_mode_box.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.source_language_combo = self.qc(self)
        self.source_language_combo.addItems([
            self.M102, self.M100, self.M103, self.M110, self.M101, self.M104, self.M109, self.M113,
            self.M106, self.M107, self.M108, self.M111, self.M112, self.M113, self.M114, self.M115,
            self.M116, self.M117
        ])
        self.source_language_combo.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.source_language_combo.setCurrentText(self.lang_in1)
        self.source_language_combo.currentTextChanged.connect(self.lang_code)
        self.target_language_combo = self.qc(self)
        self.target_language_combo.addItems([
            self.M102, self.M100, self.M103, self.M113, self.M101, self.M109, self.M106, self.M107,
            self.M108, self.M110, self.M111, self.M112, self.M113, self.M114, self.M115, self.M116,
            self.M117
        ])
        self.target_language_combo.setFont(QFont('Arial', 14 , QFont.Weight.Bold))
        self.target_language_combo.setCurrentText(self.lang_out1)
        self.target_language_combo.currentTextChanged.connect(self.lang_code)
        self.notif_console = QTextEdit(self)
        self.notif_console.setFont(QFont('Arial', 12 , QFont.Weight.Bold))
        self.notif_console.setStyleSheet(self.clr)
        self.notif_console.setReadOnly(True)
        self.info_console = QTextEdit(self)
        self.info_console.setFont(QFont('Arial', 12 , QFont.Weight.Bold))
        self.info_console.setStyleSheet(self.clr)
        self.info_console.setReadOnly(True)
        self.info2_console = QTextEdit(self)
        self.info2_console.setFont(QFont('Arial', 12 , QFont.Weight.Bold))
        self.info2_console.setStyleSheet(self.clr)
        self.info2_console.setReadOnly(True)
        self.info3_console = QTextEdit(self)
        self.info3_console.setFont(QFont('Arial', 12 , QFont.Weight.Bold))
        self.info3_console.setStyleSheet(self.clr)
        self.info3_console.setReadOnly(True)
        self.MainWindow()
        self.start()
        self.BooleanVar_menue()
    def create_settings_tab(self):
        settings_tab = QWidget()
        self.tab_widget.addTab(settings_tab, "Settings")
        layout = QVBoxLayout(settings_tab)
        # Font and Theme

    def create_translation_tab(self):
        translation_tab = QWidget()
        self.tab_widget.addTab(translation_tab, "Translation")
        layout = QVBoxLayout(translation_tab)
        font_theme_layout = QHBoxLayout()
        font_theme_layout.addWidget(QLabel("Font:"))

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if urls and urls[0].isLocalFile():
            self.file_path = urls[0].toLocalFile()
            self.trans_file()
    def BooleanVar_menue(self):
        self.timer = QTimer()
        self.CWL2="Translated list  \n\n"            
        self.perian_num = False
        self.Farsi_text_edit = False
        self.virast = False
        self.space_word_ = False
        self.using_orginal_text_enable = False  
        self.defult_format = False
        self.text_edit = False
        self.nazar = False
        self.learn=0
        self.jomlesazi = False
        self.deghat = False
        self.sjmle = False
        self.gf = False
        self.deepgf = False
        self.rtl_format_true=False
        self.dc2=False
        self.coorrect_state=False
        self.auto_correct=False
        self.corrections=False
        menubar = self.menuBar()
        self.file_menu = menubar.addMenu(self.M1)
        self.edit_menu = menubar.addMenu(self.M8)
        self.rakb_menu = menubar.addMenu(self.M150)
        self.frg = menubar.addMenu(self.M212)
        self.file_menu.addAction(self.M26, self.select_file)
        self.file_menu.addAction(self.M2, self.new_file)
        self.file_menu.addAction(self.M3, self.open_file)
        self.file_menu.addAction(self.M4, self.save_file)
        self.file_menu.addAction(self.M411, self.save_file_in)
        self.file_menu.addAction(self.M6, self.close)
        self.file_menu.addSeparator()
        self.rakb_menu.addAction(self.M10, self.choose_color)
        self.edit_menu.addAction(self.M19, self.cut_text)
        self.edit_menu.addAction(self.M20, self.copy_text)
        self.edit_menu.addAction(self.M21, self.paste_text)
        self.add_checkable_action(self.frg, self.M178, self.deghat, self.sent_degh)
        self.add_checkable_action(self.frg, self.M180, self.sjmle, self.sent_degh_2)
        self.add_checkable_action(self.frg, self.M255, self.gf, self.show2)
        self.add_checkable_action(self.frg, self.M269, False,self.sent_degh_3)  
        dic_labels= [self.M177,self.M125, self.M154,self.M159, self.M160, self.M166,self.M207,
                            self.M167, self.M168, self.M169, self.M173,self.M188,self.M213]      
        self.dict = self.qc(self)
        self.dict.addItems(dic_labels)
        self.dict.setCurrentText(self.M177)
        self.dict.setFont(QFont('Arial', 11+int(self.app_size)))
        self.dict.currentTextChanged.connect(self.dictunary)
        self.stopBook = False
        self.minimax=False
        self.update_font()
        self.add_checkable_action(self.edit_menu, self.M16, self.perian_num,self.persianNum)
        self.add_checkable_action(self.edit_menu, self.M17, self.defult_format,self.add_origin_text)
        self.add_checkable_action(self.edit_menu, self.M18, self.using_orginal_text_enable,self.using_orginal_text)
        self.add_checkable_action(self.edit_menu, self.M12, self.rtl_format_true, self.fom_at)
        self.add_checkable_action(self.edit_menu, self.M270, self.dc2, self.perian2)
        self.add_checkable_action(self.edit_menu, '  آموزش نیمه خودکار لغت (در حالت اصلاح نیمه خودکار لغت لیست اصلاح نیمه خودکار نمایش داده می شود )', self.dc2, self.learning2)
    def positions_and_size(self):
        """Update the positions and sizes of UI elements based on the app size."""
        screen = QApplication.primaryScreen()
        screen_geometry = screen.geometry()
        screen_width, screen_height = screen_geometry.width(), screen_geometry.height()
        base_width, base_height = 1280, 660
        scale_factor = min(screen_width / base_width, screen_height / base_height) * (self.app_size*0.70)
        scaled_width = int(base_width * scale_factor)
        scaled_height = int(base_height * scale_factor)
        self.setGeometry(40, 40, scaled_width, scaled_height)
        def move_resize(widget, x, y, width, height):
            widget.move(int(x * scale_factor), int(y * scale_factor))
            widget.resize(int(width * scale_factor), int(height * scale_factor))
        widgets = [
            (self.size_app, 11, 40, 110, 40),
            (self.them_num, 140, 40, 110, 40),
            (self.book_box, 260, 40, 96, 40),
            (self.dict, 360, 40, 190, 40),
            (self.translator_menu, 600, 40, 230, 40),
            (self.document_mode_box, 850, 40, 230, 40),
            (self.source_language_label, 10, 90, 114, 40),
            (self.source_language_combo, 140, 90, 180, 40),
            (self.target_language_label, 630, 90, 114, 40),
            (self.target_language_combo, 750, 90, 180, 40),
            (self.ch_tran_lang, 470, 90, 135, 40),
            (self.BS, 1100, 250, 160, 35),
            (self.info2_console, 1085, 30, 190, 200),
            (self.input_text_label, 350, 110, 90, 25),
            (self.output_text_label, 960, 110, 90, 25),
            (self.input_console, 10, 135, 520, 380),
            (self.output_console, 550, 135, 520, 380),
            (self.trans_file_button, 1100, 370, 160, 38),
            (self.pdf_converting, 1100, 420, 160, 35),
            (self.lang2_button, 1100, 460, 160, 35),
            (self.paste_button, 10, 480, 90, 35),
            (self.copy_button, 980, 480, 90, 35),
            (self.lang_button, 1100, 500, 160, 35),
            (self.help_button, 1100, 540, 160, 35),
            (self.reset_button, 1100, 580, 160, 35),
            (self.file_button, 10, 560, 96, 50),
            (self.virastar, 720, 620, 160, 40),
            (self.space_corect, 230, 620, 200, 40),
            (self.learned_trance,  5, 620, 230, 40),
            (self.semi_correct, 410, 620, 200, 40),
            (self.auto_correction, 580, 620, 150, 40),
            (self.size_box, 870, 620, 90, 35),
            (self.font_box, 985, 620, 166, 35),
            (self.clear_button, 790, 560, 166, 50),
            (self.stop_button, 670, 560, 110, 50),
            (self.export_button, 985, 560, 85, 50),
            (self.translate_button, 120, 560, 330, 50),
            (self.learn_button, 500, 560, 160, 50),
            (self.info3_console, 1100, 290, 160, 80),
            (self.info_console, 550, 517, 520, 35),
            (self.notif_console, 10, 517, 520, 35),
        ]
        for widget, x, y, width, height in widgets:
            move_resize(widget, x, y, width, height)
        self.thm2()
    def change_lang_translate(self):
        lang1=self.target
        lang2=self.source
        self.source_language_combo.setCurrentText(lang1)
        self.target_language_combo.setCurrentText(lang2)
        self.from_code = self.language_codes.get(lang1)
        self.to_code = self.language_codes.get(lang2)
        self.lang_code()
    def add_checkable_action(self, menu, label, variable, callback=None):
        action = QAction(label, self, checkable=True)
        action.setChecked(int(variable))
        if callback:
            action.triggered.connect(callback)
        menu.addAction(action)
    def sent_degh(self):
        self.sjmle = False
        self.gf = False
    def sent_degh_2(self):
        self.deghat = False
        self.gf = False
    def sent_degh_3(self):
        if self.deepgf==False:
            self.deepgf=True
        else:
            self.deepgf=False
    def show2(self):
        self.input_console.setExtraSelections([])  # Clear previous highlights
        search_query, ok = QInputDialog.getText(self, "Find", "Enter your text")
        if ok and search_query:
            format = QTextCharFormat()
            format.setForeground(QBrush(QColor("blue")))
            extra_selections = []
            self.input_console.moveCursor(QTextCursor.Start)
            while self.input_console.find(search_query, QTextDocument.FindCaseSensitively): # type: ignore
                selection = QTextEdit.ExtraSelection()
                selection.cursor = self.input_console.textCursor()
                selection.format = format
                extra_selections.append(selection)
            self.input_console.setExtraSelections(extra_selections)
    def read_them_lines(self):
        try:
            with open('them.jsonl', 'r', encoding="utf-8") as f:
                return f.read().splitlines()
        except:
            try:
                with open('them.jsonl', 'w', encoding="utf-8") as f:
                        f.write(self.default)
                with open('them.jsonl', 'r', encoding="utf-8") as f:
                    return f.read().splitlines()
            except:
                return self.default.splitlines()
    def thm(self):
        self.default = f"""
Defult.CSS
فارسی
فارسی
0
#141414
14
Arial
مترجم آفلاین
spaceCorrect1
0
0
0
0
0
0
0
''
{os.path.expanduser("~")}
0
0
0
0
0
0
"""
        self.default2 = f"""
Defult.CSS
English
English
0
#141414
14
Arial
Google
spaceCorrect1
0
0
0
0
0
0
0
''
{os.path.expanduser("~")}
0
0
0
0
0
0
"""
        try:
            lines = self.read_them_lines()
            self.them1, self.lang_in1, self.lang_out1 = lines[1], lines[2], lines[3]
            self.dc, self.color1, self.size1 = lines[4], lines[5], lines[6]
            self.font1, self.translator1, self.spaceCorrect1 = lines[7], lines[8], lines[9]
            self.lastDir, self.app_size = lines[18], lines[10]
        except Exception:
            try:
                lines = self.read_them_lines()
                self.them1, self.lang_in1, self.lang_out1 = lines[1], lines[2], lines[3]
                self.dc, self.color1, self.size1 = lines[4], lines[5], lines[6]
                self.font1, self.translator1, self.spaceCorrect1 = lines[7], lines[8], lines[9]
                self.lastDir, self.app_size = lines[18], lines[10]
               
            except Exception:
                self.update_notification(self.M254)
        self.num = self.them1
    def thm2(self):
        try:
            lines = self.read_them_lines()
            self.them1, self.lang_in1, self.lang_out1 = lines[1], lines[2], lines[3]
            self.dc, self.color1, self.size1 = lines[4], lines[5], lines[6]
            self.font1, self.translator1, self.spaceCorrect1 = lines[7], lines[8], lines[9]
            self.lastDir, self.app_size = lines[18], lines[10]
            self.virastar.setCheckState(Qt.CheckState(int(lines[12])))
            self.space_corect.setCheckState(Qt.CheckState(int(lines[13])))
            self.learned_trance.setCheckState(Qt.CheckState(int(lines[19])))
            self.semi_correct.setCheckState(Qt.CheckState(int(lines[15])))
            self.auto_correction.setCheckState(Qt.CheckState(int(lines[14])))
            self.rtl_format_true = int(lines[16])
            self.file_path = lines[17]
            file_name = os.path.basename(self.file_path)
            self.dc2= int(lines[4])
            self.perian_num=int(lines[19])
            self.defult_format=int(lines[20])
            self.using_orginal_text_enable=int(lines[21])
            self.info3(file_name)
        except Exception:
            try:
                os.remove('them.jsonl')
                lines = self.read_them_lines()
                self.them1, self.lang_in1, self.lang_out1 = lines[1], lines[2], lines[3]
                self.dc, self.color1, self.size1 = lines[4], lines[5], lines[6]
                self.font1, self.translator1, self.spaceCorrect1 = lines[7], lines[8], lines[9]
                self.lastDir, self.app_size = lines[18], lines[10]
                self.virastar.setCheckState(Qt.CheckState(int(lines[12])))
                self.space_corect.setCheckState(Qt.CheckState(int(lines[13])))
                self.learned_trance.setCheckState(Qt.CheckState(int(lines[19])))
                self.semi_correct.setCheckState(Qt.CheckState(int(lines[15])))
                self.auto_correction.setCheckState(Qt.CheckState(int(lines[14])))
                self.rtl_format_true = int(lines[16])
                self.file_path = lines[17]
                file_name = os.path.basename(self.file_path)              
                self.dc2 = int(lines[4])
                self.perian_num=int(lines[19])
                self.defult_format=int(lines[20])
                self.using_orginal_text_enable=int(lines[21])
                self.info3(file_name)
            except Exception:
                self.update_notification(self.M254)
        self.num = self.them1
    def lang_select(self):
        try:
            with open('index_lang.json', 'r') as f:
                lang = f.read()
                if lang == 'en':
                    self.lang_en()
                else:
                    self.lang_fa()
        except:
            try:
                with open('index_lang.json', 'w', encoding="utf-8") as f:
                    f.write('fa')
            except:
                self.update_notification(self.M254)
            self.lang_fa()
    def jamal(self):
        self.info2("https://www.iranlawclinic.com")
        try:
            os.startfile('کلینیک حقوقی ایران.html')
        except:
             self.info0("\n https://www.iranlawclinic.com")
    def them(self):
        t=self.them_num.currentText()
        self.save_state(1,t)
        result =self.qq(self,self.M258,self.M87)
        if result ==16384:
           self.restart_program()
        else:
            pass
        return
    def off_ketab(self):
        self.end = True
        self.search_Active = False
        self.separate_search = False
        self.internet_aktive = False
        self.book_aktive = False
        self.dict_aktive = False
        self.dict.setCurrentText(self.M177)
        self.info2('\n\n')
        return
    def dictunary(self):
        if self.dict.currentText() != self.M177:
            self.spm = False
            self.separate_search = False
            self.search_Active = True
            self.internet_aktive = False
            self.book_aktive = False
            self.dict_aktive = True
            if self.dict.currentIndex() != self.M125:
                self.info2(self.dict.currentText())
                try:
                    if not self.thread_active:
                        self.separate_search = True
                        self.ketab_ = ""
                        src_book_dic = self.searching_book_options()
                        if self.spm or src_book_dic:
                            self.run_book_search()
                        elif self.dict.currentText() in [self.M207]:
                            self.search_process = self.dict.currentText()
                            self.run_book_search()
                except:
                    pass
            else:
                self.info2(self.M125)
        else:
            self.off_ketab()
    def persianNum(self):
        if  self.perian_num==False:
            self.perian_num=True
            t=1
        else:
            t=0
            self.perian_num=False
        self.save_state(19,t)
    def perian2(self):
        if  self.dc2==False:
            self.dc2=True
            t=1
        else:
            self.dc2=False
            t=0
        self.save_state(4,t)
    def add_origin_text(self):
        if  self.defult_format==False:
            self.defult_format=True
            t=1
        else:
            t=0
            self.defult_format=False
        self.save_state(20,t)
    def fom_at(self):
        if  self.rtl_format_true==False:
            self.rtl_format_true=True
            t=1
        else:
            t=0
            self.rtl_format_true=False
        self.save_state(16,t)
    def using_orginal_text(self):
            if  self.using_orginal_text_enable==False:
                self.using_orginal_text_enable=True
                t=1
            else:
              t=0
              self.using_orginal_text_enable=False
            self.save_state(21,t)
    def aktive_virast(self):
        if self.virast !=True:
            self.virast=True
            self.info2(self.M95)
            t=2
        else:
            self.virast=False
            self.info2(" ")
            t=0
        self.save_state(12,t)
    def aktive_space(self):
        if self.space_word_ != 0:
            self.space_word_ = 0
            self.space = 0
            t=0
        else:
            t=2
            self.space_word_ = 1
        self.space = 1
        self.save_state(13,t)
    def aktive_learn(self):
        if self.learn != 0:
            self.learn = 0
            t=0
        else:
            t=2
            self.learn = 1
        self.save_state(19,t)
        
    def auto_corrections_state(self):
        if self.auto_correct != True:
            self.semi_correct.setCheckState(Qt.CheckState.Unchecked)
            self.semi_corrections = False
            self.corrections = True
            self.auto_correct = True
            self.info2_console.setReadOnly(False)
            self.info2_console.clear()
            self.info2(self.M129)
            t = 2
        else:
            t = 0
            self.auto_correction.setCheckState(Qt.CheckState.Unchecked)
            self.corrections = False
            self.semi_corrections = False
            self.auto_correct = False
            self.info2_console.setReadOnly(False)
            self.info2_console.clear()
        self.save_state(14, t)

    def semi_corrections_state(self):
        try:
            if self.semi_corrections != True:
                self.auto_correction.setCheckState(Qt.CheckState.Unchecked)
                self.corrections=True
                self.auto_correct=False
                self.semi_corrections = True
                self.info2(self.M128)
                t=2
            else:
                t=0
                self.semi_correct.setCheckState(Qt.CheckState.Unchecked)
                self.info2(" \n")
                self.corrections=False
                self.semi_corrections = False
                self.auto_correct=False
            self.save_state(15,t)
        except:
            pass
         
    def save_state(self,l,t):
        try:
            with open('them.jsonl', 'r', encoding="utf-8") as f:
                        lines = f.read().splitlines()
                        lines[int(l)] =str(t)
            with open('them.jsonl', 'w', encoding="utf-8") as f:
                        for line in lines:
                            f.write(line)
                            f.write('\n')
        except:
            self.update_notification(self.M254)
    def run_book_search(self):
        try:
            self.stopBook = False
        except:
            pass
        try:
            self.cunter = 0
            self.info2('')
        except:
            pass
        self.update()
        self.book_search("", "", "")
    def start(self):
        self.word_regexes = {
    'ar':r"[آ-ی]+", 
    "en": r"[A-Za-z]+",
    "de": r"[A-Za-z]+",
    "nl": r"[A-Za-z]+",
    "pl": r"[A-Za-zęĘóÓąĄśŚłŁżŻźŹćĆńŃ]+",
    "ru": r"[АаБбВвГгДдЕеЁёЖжЗзИиЙйКкЛлМмНнОоПпРрСсТтУуФфХхЦцЧчШшЩщЪъЫыЬьЭэЮюЯя]+",
    "tr": r"[a-zA-ZçÇğĞüÜöÖşŞıİ]+",
    "es": r"[A-Za-zÁáÉéÍíÓóÚúÜüÑñ]+",
    "pt": r"[a-zA-ZãáàâçéêíõóôúüÃÁÀÂÇÉÊÍÕÓÔÚÜ]+",
    "it": r"[a-zA-ZãáàâçéêíõóôúüÃÁÀÂÇÉÊÍÕÓÔÚÜ]+",
    "fr": r"[a-zA-ZãáàâçéêíõóôúüÃÁÀÂÇÉÊÍÕÓÔÚÜ]+",
}
        self.filter = ['Download', 'Herunterladen', 'تحميل', 'Télécharger','اکستروژن','فلز','game','بازی','Punch-tera','موسیقی','Crush','جادوگر',
                        '&quot;','Descargar', 'Загружать', 'Scaricare', 'İndirmek', 'Baixar','دانلود','کامپیوتری','Britney' ,'Bastard','Punch']
        self.dir_path = os.path.join(os.path.expanduser('~'),'.local','cache','argos-translate','downloads')
        self.base_path = os.path.join(os.path.expanduser('~'), '.local','share','argos-translate','packages')
        self.destination_path = os.path.join(os.path.expanduser('~'), '.local','cache','argos-translate')
        self.source_path = os.path.join(self.dir, 'index.json')
        self.file_name = os.path.basename(self.source_path)
        self.special_chars2 = "[$#%&()*+-./:;«»,؟!؛<=>?@[\\]^_`{|}،\u200c ]"
        self.destination_file_path = os.path.join(self.destination_path, self.file_name)
        self.special_chars = ['$', '#', '%', '&', '(', ')', '*', '+', ',', '-', '.', '/', ':', ';','«','»' ,'؟','!','؛','"','.', '. ',' .',  
        '•','<', '=', '>', '?', '@', '[', '\\', ']', '^', '_', '`', '{', '|', '}','، ',' ،','،',', ',' ,',',',' ؟','؟ ']
        self.url_downloads='url_downloads.json'
        self.sin = ['~','!','@','#','$','%','^','&','*','(',')','_','-','=','.','/','','+','<',
                    '>','{','}','?','؟','|','"',"'",
                    ':',';',',','حح"', '"حح', 'حححح',]
        self.invalid_languages = {'', 'Exception', 'No features in text', 'id', 'ur', 'ch', 'af', 'sl', 'se', 'sr', 'sk', 'su',
                                'hy', 'as', 'av', 'ay', 'bn', 'bg', 'ch', 'cv', 'cr', 'cs', 'dv', 'so', 't1', 'ca','lv','lt','pl',
                                'et', 'ee', 'no', 'ro', 'fy', 'gu', 'ha', 'kn', 'kk', 'km', 'fi','sv','da','hr','t1','tl',
                                'kj', 'ko', 'ms', 'nn', 'uk', 'nb', 'pa', 'rn', 'tk', 'xh','uz','bo','cy','vi','ro','sw'}
        self.time= 0.23
        self.patterns = [
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            r'\b\d{5}(?:-\d{4})?\b', 
            r'\b\d{1,16}\b',  
            r'\b\d+\b', 
            r'\b00\d{2}\d{7,12}\b',
            r'\+\d{2}\d{7,12}\b', 
            r'\b0\d{2}\d{7,12}\b', 
            r'^[a-zA-Z]',
            r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b'
        ]
        self.start_with=['الله','ب','می', 'نمی', 'بی', 'نا','ال','پر','داد','آ','جا','محمد','میر','نمی‌','می‌',
                            'با','به','نی','ل','خو','نیک','پاک','بر','دار','هر',
                            'در','سر','پیش','پس','ان','علی','خان','خواجه','نی','نیا','بار','خوش','بد','الی','فر','آقا',
                            'یک','دوی','سی','چهار','پان','پنج','چهل','پنجاه','شش','شصت','هفت','هفتاد','زشت','والا','و','از',
                            'هشت','هشتاد','نه','نود','هزار','میلیون','ملیارد','ترلیون','بهر','بهره','زیبا','من',
                            ]
        self.start_with.sort(key=len, reverse=False)

        self.suffixes = [ 'الله', 'هایمان', 'هایشان', 'ستان', 'طور', 'طوری', 'گاری', 'گذار','شان',
                            'ستانی','خو','سر','وند',  'انی', 'یمان', 'هایم', 'هایش', 'های', 'آلات','ریز', 'که','می‌',
                            'علی','خان','خواجه','گی','بیک','بیگ','اژ','خور','ار','زاد','راد','سیرت','بار','فر',
                            'یک','دوی','سی','چهار','پان','پنج','چهل','پنجاه','شش','شصت','هفت','سرشت','یار','آن',
                            'هفتاد','هشت','هشتاد','نه','نود','هزار','میلیون','ملیارد','ترلیون','شاه','رت',
                            'دار', 'زا', 'هایت', 'گار', 'مین','بر','فت','افت','فند','تان', 'سرا', 'های','داد',
                            'یند', 'نامه', 'آموزی', 'آموز', 'وار', 'کار', 'مند', 'گرا','ست','یافت','پیش',
                            'پس','داری', 'گیری', 'آور', 'ستان', 'گری', 'گاه', 'بین', 'زاده', 'واری', 'منش',
                            'یان','ییان','دا','بود','خوار', 'آوازه', 'بند', 'بندی', 'نواز', 'انه', 'پذیر','اید',
                            'ترین', 'پسین', 'یه', 'چه','ک','وه','مان','در','ات', 'یم', 'گر', 'یت', 'یش','اش',
                            'بندی', 'بند', 'ان', 'ای','نامه', 'بان', 'بانی','ند','اند','شان','سیرت','ین',
                            'ید', 'نی', 'می','تر' ,'ها', 'ی','ش','را','انداز',
                        ]
        self.persian_conjunctions = [
        # حروف ربط ساده
        "و", "یا", "پس", "اگر", "نه", "را", "چون", "چه", "تا", "اما", "باری", "خواه", "زیرا", "که", "لیکن", "نیز", "ولی", "هم",
        # حروف ربط مرکب
        "بالعکس", "ولو", "به جز", "سپس", "از این گذشته", "همچنین", "چون‌ که", "چندان ‌که", "زیرا که", "همان‌ که", "بلکه", "چنانچه", "تا این که", "تا آن که", "آنگاه که", "از آن‌جا که", "از این ‌رو", "از بس", "از بهر آن که", "اکنون که", "اگرچه", "اگر", "مگر این که", "با این حال", "با این‌ که", "با وجود این", "بس که", "به شرط آن که",
        # حروف ربط گسسته ربطی
        "خواه", "چه", "یا", "نه", "هم"
    ]
        self.suffixes.sort(key=len, reverse=True)
        self.char_groups = [ ['ز', 'ظ'], ['ز', 'ض'], ['ذ', 'ز'], ['ض', 'ظ'], ['ظ', 'ض'], ['ذ', 'ظ'],
                                ['ذ', 'ض'], ['ز', 'ذ'], ['ظ', 'ز'], ['ض', 'ز'], ['ط', 'ت'], ['ت', 'ط'],
                                ['ر', 'ز'], ['ز', 'ر'], ['د', 'ذ'], ['ذ', 'د'], ['ح', 'ه'], ['ه', 'ح'],
                                    ['ج', 'چ'], ['چ', 'ج'], ['ح', 'خ'], ['خ', 'ح'], ['ح', 'خ'], ['ح', 'ج'],
                                    ['ج', 'ح'], ['ص', 'س'], ['س', 'ص'], ['س', 'ث'], ['ص', 'ث'], ['س', 'ص'],
                                    ['ش', 'س'], ['ص', 'ث'], ['ص', 'س'], ['ث', 'ص'], ['ث', 'س'], ['ص', 'س'],
                                    ['س', 'ش'], ['ع', 'ا'], ['ا', 'ع'], ['ا', 'ع'],['ا', 'ع'],
                                    ['ق', 'غ'], ['غ', 'ق'], ['ق', 'ف'], ['ف', 'ق'],
                                    ['ک', 'گ'], ['یی', 'ت'], ['یی', 'ی'], ['یی', ''], ['ک', 'گ'], ['گ', 'ک'], ['خوا', 'خا'], ['خا', 'خوا'] ]
        self.chars = ['','آ','ا','ا','ب', 'پ', 'ت', 'ث', 'ج', 'چ', 'ح', 'خ', 'د', 'ذ', 'ر',
                        'ز', 'ژ', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ک', 'گ',
                        'ل', 'م', 'ن', 'و', 'ه', 'ی']
        self.on_off()
    def on_off(self):
        self.text=""
        self.language_codes_2 = {
            "en":self.M100,
            "auto":self.M104,
            "fa":self.M102,
            "de":self.M101,
            "ar":self.M109,
            "fr":self.M108,
            "zh": self.M103,
            "es": self.M110,
            "ru": self.M111,
            "it": self.M106,
            "tr":self.M112,
            "pt":self.M113,
            "id": self.M114,
            "nl": self.M107,
            "hi": self.M115,
            "ja": self.M116,
            "ur":self.M117}
        self.language_codes = {
            self.M100: "en",
            self.M104:"auto",
            self.M102: "fa",
            self.M101: "de",
            self.M109: "ar",
            self.M108: "fr",
            self.M103: "zh-CN",
            self.M110: "es",
            self.M111: "ru",
            self.M106: "it",
            self.M112: "tr",
            self.M113: "pt",
            self.M114: "id",
            self.M107: "nl",
            self.M115: "hi",
            self.M116: "ja",
            self.M117: "ur",}
        self.semi_corrections = False
        self.space=1
        self.space_word_=1
        self.internet_aktive=False
        self.book_aktive=False
        self.dict_aktive=False
        self.search_Active=False
        self.separate_search=True
        self.offline_installed=False
        self.ketab_=""
        self.file_path=''
        self.tranc_err=False
        self.word_office_not_installed=False
        self.search_process=''
        self.save_word=False
        self.clicked=False
        self.clicked_trueWords=False
        self.pdf=False
        self.excel=False
        self.words2=''
        self.translate_pay=False
        self.arg_para=False
        self.extract=False
        self.pdf_convert=False
        self.empty_1=False
        self.i=0
        self.color_code = (0, 0, 0)
        self.search_opt='select'
        self.file_content = {}
        self.reerror_packages = False
        self.reinstalled = False
        self.thread_active=False
        self.pack_install = False
        self.argose_err = False
        self.pack_downloaded=False
        self.rtl_mode = False
        self.packerror=False
        self.er=False
        self.last_detected_languages ="en"
        self.tra=False
        self.reerror_packages=False
        self.reinstalled=False
        self.dfiscancel=False
        self.rev=False
        self.export_docx_=False
        self.sp1=False
        self.sp2=False
        self.sp3=False
        self.sp4=False
        self.S5=False
        self.S6=False
        self.S7=False
        self.S8=False
        self.S9=False
        self.S10=False
        self.S11=False
        self.S12=False
        self.S13=False
        self.S14=False
        self.S15=False
        self.S16=False
        self.S17=False
        self.S18=False
        self.S19=False
        self.S20=False
        self.S21=False
        self.S22=False
        self.S23=False
        self.S24=False
        self.S25=False
        self.S26=False
        self.S27=False
        self.S28=False
        self.S29=False
        self.S30=False
        self.S31=False
        self.spm=False
        self.S32=False
        self.S33=False
        self.S34=False
        self.S35=False
        self.S36=False
        self.S37=False
        self.S38=False
        self.S40=False
        self.S41=False
        self.S42=False
        self.S43=False
        self.S44=False
        self.S45=False
        self.S46=False
        self.S47=False
        self.S48=False
        self.S49=False
        self.S50=False
        self.S51=False
        self.S52=False
        self.S53=False
        self.S54=False
        self.S55=False
        self.S56=False
        self.S57=False
        self.S58=False
        self.S59=False
        self.S60=False
        self.S61=False
        self.S62=False
        self.farhang=[]
        self.sp63=False
        self.S63=False
        self.S64=False
        self.S65=False
        self.S66=False
        self.S67=False
        self.S68=False
        self.S69=False
        self.S70=False
        self.S71=False
        self.stops=False
        self.thread_active=False
        self.skip_rtl=False
        self.console=False
        self.skip_save=False
        self.docx=False
        self.docx2=False
        self.CWL='لیست کلمات جایگزین شده \n'
        self.total_para=5
        self.lang_code()
        self.update_output('mj.taraz@yahoo.com')
    def open_3(self):
        self.update_notification(self.M162)
        with open('mjt.json', 'r', encoding='utf-8') as f:
            self.synonyms_str = f.read().splitlines()
            self.S5=True
            return
    def open_4(self):
        try:
            self.not_found_file = 'not_found_fa.json'
            txt_per_comp1 = 'Book_fa.json'
            self.replace = 'replace.json'
            self.replace2 = 'replace2.json'

            # Read the replace file only once and reuse its contents.
            with open(self.replace, 'r', encoding='utf-8') as f:
                replaced_content = f.read()
            self.replaced_lines = replaced_content.splitlines()
            self.replaced_words = set(re.split(r'[`\n]+', replaced_content))

            # Similarly for replace2.
            with open(self.replace2, 'r', encoding='utf-8') as f:
                replaced2_content = f.read()
            self.replaced_lines2 = replaced2_content.splitlines()
            self.replaced_words2 = set(re.split(r'[`\n]+', replaced2_content))

            # Load additional files only when S6 is False.
            if not self.S6:
                with open(txt_per_comp1, 'r', encoding='utf-8') as f:
                    self.fa_words = f.read().splitlines()
                with open(self.not_found_file, 'r', encoding='utf-8') as f:
                    self.fa_w_not_found = f.read().splitlines()

            self.S6 = True

        except Exception as e:
            # Create missing files if they do not exist.
            for filename in ['not_found_fa.json', 'replace.json', 'found_fa.json', 'replace2.json']:
                if not os.path.exists(filename):
                    with open(filename, 'w', encoding='utf-8') as f:
                        f.write('')
            # Optional: Log the exception for further debugging.
            self.info0(f"Error in open_4: {e}")


    def open_5(self):
        try:
            self.lo = f"{self.from_code}_{self.to_code}_learn.json"
            with open(self.lo, 'r', encoding='utf-8') as f:
                content = f.read()
            self.learn_lines = content.splitlines()
            self.learn_texts = re.split(r'[`\n]+', content)
            
        except Exception as e:
            try:
                # Create the learning file if it doesn't exist.
                if not os.path.exists(self.lo):
                    with open(self.lo, 'w', encoding='utf-8') as f:
                        f.write('')
                    # Read back the (now empty) file.
                    with open(self.lo, 'r', encoding='utf-8') as f:
                        content = f.read()
                    self.learn_lines = content.splitlines()
                    self.learn_texts = re.split(r'[`\n]+', content)
                else:
                    # If the file exists but another error occurred, trigger a notification.
                    self.info0('برای آموزش دیدن نرم افزار آن را در حالت ادمین اجرا کنید')
            except Exception as ex:
                self.info0(f"Error in open_5 file creation: {ex}")
    def update_notification(self, message):
        self.notif_console.setReadOnly(False)
        self.notif_console.clear()
        self.notif_console.setPlainText(message)
        self.notif_console.setReadOnly(True)
        self.wait()

    def update_input(self, text):
        self.input_console.insertPlainText(text)
        self.input_console.moveCursor(QTextCursor.MoveOperation.End)
        self.wait()

    def update_output(self, text):
        self.output_console.setReadOnly(False)
        self.output_console.insertPlainText(text)
        self.output_console.moveCursor(QTextCursor.MoveOperation.End)
        self.wait()

    def info0(self, text):
        self.info_console.setReadOnly(False)
        self.info_console.setPlainText(text)
        self.info_console.setReadOnly(True)
        self.wait()

    def info2(self, text):
        self.info2_console.setReadOnly(False)
        self.info2_console.insertPlainText('\n')
        self.info2_console.insertPlainText(text)
        self.info2_console.moveCursor(QTextCursor.MoveOperation.End)
        self.info2_console.setReadOnly(True)
        self.wait()

    def info3(self, text):
        self.info3_console.setReadOnly(False)
        self.info3_console.insertPlainText('\n')
        self.info3_console.setPlainText(text)
        self.info3_console.setReadOnly(True)
        self.wait()

    def show1(self):
        self.input_console.setExtraSelections([])  # Clear previous highlights
        search_query, ok = QInputDialog.getText(self, "Find", "Enter your text")
        if ok and search_query:
            format = QTextCharFormat()
            format.setForeground(QBrush(QColor("blue")))
            extra_selections = []
            self.input_console.moveCursor(QTextCursor.MoveOperation.Start)
            while self.input_console.find(search_query, QTextDocument.FindFlag.FindCaseSensitively):
                selection = QTextEdit.ExtraSelection()
                selection.cursor = self.input_console.textCursor()
                selection.format = format
                extra_selections.append(selection)
            self.input_console.setExtraSelections(extra_selections)

    def Lang(self):
        try:
            with open('index_lang.json', 'r', encoding='utf-8') as f:
                lang = f.read()
                if lang == 'en':
                    confirmation = self.qq(self,'بازنشانی برنامه برای ایجاد تغییر زبان فارسی','برای ایجاد تغییرات زبان برنامه باید بازنشانی شود آیا تایید می فرمایید؟')
                    if confirmation==16384:
                        with open('index_lang.json', 'w', encoding='utf-8') as f:
                            f.write('')
                        self.restart_program()
                else:
                    self.update_notification('زبان برنامه فارسی است' )
        except:
            self.update_notification(self.M254)
    def Lang2(self):
        try:
            with open('index_lang.json', 'r') as f:
                    lang = f.read()
                    if lang != 'en':
                        result = self.qq(self, "Program need restart to take effect for English",
                                        "Do you want to restart program for take effect?",
                                        self.q.StandardButton.Yes | self.q.StandardButton.No, self.q.StandardButton.No)
                        if result == self.q.StandardButton.Yes:
                            self.defultthem()
                            with open('index_lang.json', 'w', encoding="utf-8") as f:
                                f.write('en')
                            # self.start1()
                            self.restart_program()
                    else:
                        self.update_notification("Program language already English")
        except:
            self.update_notification(self.M254)
    def lang_fa(self):
        self.M001='نمایه ها'
        self.M002='زبان ورودی'
        self.M003='زبان خروجی'
        self.M004="جستجو در کتاب"
        self.M005='تغییر سایز برنامه'
        self.file_button.setText("انتخاب فایل")
        self.translate_button.setText("ویرایش و ترجمه متن")
        self.trans_file_button.setText("ویرایش و ترجمه فایل")
        self.pdf_converting.setText('تبدیل PDF')
        self.export_button.setText('ذخیره ')
        self.clear_button.setText('پاک کردن متن')
        self.paste_button.setText(' جایگذاری')
        self.help_button.setText('راهنما')
        self.reset_button.setText('بازنشانی')
        self.copy_button.setText('کپی')
        self.stop_button.setText("توقف پردازش")
        self.M1='فایل'
        self.M2='پاک کردن متن ورودی'
        self.M3='وارد نمودن فایل با فرمت  تی ایکس تی'
        self.M4='ذخیره متن خروجی با فرمت  تی ایکس تی '
        self.M411='ذخیره متن ورودی با فرمت  تی ایکس تی '
        self.M6='خروج از برنامه'
        self.M8='ویرایش'
        self.M9='نوع تبدبل فایل'
        self.M10='انتخاب رنگ متن'
        self.M11= 'اصلاح فاصله'
        self.M12=' تغییر چیدمان متن PDF'
        self.M13='اصلاح نیمه خودکار لغت'
        self.M14='ویرایش متن فارسی '
        self.M15='اصلاح خودکار لغت  '
        self.M16=' تبدیل اعداد از انگلیسی به  فارسی '
        self.M17='فرمت و فونت پیش فرض برای فایل های پی دی اف و ورد'
        self.M18='اضافه نمودن متن اصلی به متن ترجمه شده در فایل های پی دی اف و ورد '
        self.M19='بریدن'
        self.M20='رونوشت'
        self.M21=' جایگذاری'
        self.M22='خطای دانلود'
        self.M23=' خطا '
        self.M24='لطفا صبر کنید تا پردازش قبلی کامل شود'
        self.M25='هیچ فایلی  انتخاب نشده است'
        self.M26='انتخاب فایل'
        self.M27='خطا در انتخاب فایل'
        self.M28='متن از فایل پی دی اف وارد شد'
        self.M29='خطا در پردازش پی دی اف'
        self.M30='استخراج جدول از پی دی اف شروع خواهد شد.'
        self.M31=' خطا در استخراج متن از پی دی اف'
        self.M32='هیچ فایلی  برای ذخیره ایجاد یا انتخاب نشد.'
        self.M33='پایان پردازش'
        self.M34=' حذف فایل'
        self.M35='خطا در تبدیل پردازش پی دی اف'
        self.M36=' پردازش متن  پی دی اف'
        self.M37=' پردازش   پی دی اف'
        self.M38='  در حال پردازش متن'
        self.M39='اتمام پردازش '
        self.M40='  در حال پردازش ورد'
        self.M41='ترجمه و ذخیره فایل'
        self.M42='  پایان پردازش  '
        self.M43=' خطا در  پردازش ورد'
        self.M44="آموزش مترجم آفلاین و استفاده"
        self.M45='پردازش پی دی اف و استخراج جدول'
        self.M46='لغو پردازش پی دی اف و استخراج جدول'
        self.M47='لغونصب برنامه کمکی استخراج جدول'
        self.M48='در حال پردازش پی دی اف و استخراج جدول'
        self.M49='جدول اسخراج شده '
        self.M50=' اسخراج جدول تکمیل شد'
        self.M51=' فایل  برای ذخیره ایجاد یا انتخاب نشد'
        self.M52='ذخیره خروجی در فایل '
        self.M53='جدول اسخراج شده '
        self.M54=' استخراج جدول از فایل انتخاب شده امکان پذیر نیست'
        self.M55='کتاب مورد نظر یافت نشد '
        self.M56='خطا در ترجمه'
        self.M57='در حال ترجمه'
        self.M58='لطفا صبر فرمایید'
        self.M59='خطا در ترجمه یا پردازش متن '
        self.M60=' زبان مورد نظر یافت نشد '
        self.M61='خطا در ترجمه آنلاین لطفا اینترنت را چک یا مترجم آفلاین را امتحان کنید '
        self.M62='خطا در مترجم دوم گوگل'
        self.M63='مترجم ارگوس'
        self.M64='در حال بررسی بسته زبان مورد نظر'
        self.M65='خطا در نرجمه آفلاين'
        self.M66='بسته زبان مورد نظر یافت نشد'
        self.M67='خطا در نصب بسته زبان دوباره دانلود می شود'
        self.M68='خطا در نصب بسته زبان '
        self.M69='یک خطا رخ داد  بسته زبان دوباره نصب و برنامه بازنشانی گردد؟  '
        self.M70='دانلود مجدد'
        self.M71=' زبان درخواست شده برای مترجم آفلاین در دسترس نیست'
        self.M72='یک مترجم دیگر را از لیست انتخاب فرمایید '
        self.M73='دانلود  '
        self.M74='لغو گردید لطفا یک مترجم دیگر را از لیست انتخاب فرمایید '
        self.M75= 'دانلود'
        self.M76='بسته زبان'
        self.M77='دانلود  پایان یافت  '
        self.M78='خطا در دانلود'
        self.M79='لفا دوباره سعی فرمایید'
        self.M80='2 مترجم گوگل'
        self.M81='خطای مترجم مای مموری'
        self.M82='خطای مترجم دیپ'
        self.M83='خطای مترجم یاندکس'
        self.M84=' انتخاب حالت اصلاح لغت'
        self.M85=' فایل  اجرا شود ؟'
        self.M86=' اطلاع رسانی'
        self.M87='نیاز به بازنشانی سیستم است تایید می فرمایید ؟'
        self.M88='برای استخراج جدول از فایل پی دی اف نرم افزار گوییست اسکریپ باید نصب و سسیستم بازنشانی گردد تایید می فرمایید ؟'
        self.M89='نصب بسته زبان تکمیل شد برنامه بازنشانی  می شود'
        self.M90='بسته زبان مورد نظر یافت نشد بسته  دانلود شود ؟'
        self.M91='در حال آماده سازی فایل لطفا صبر فرمایید:'
        self.M92='نوع فایل انتخاب شده برای تبدیل با قالب مشابه مناسب نیست آیا تمایل دارید فقط متن از فایل استخراج شود ؟'
        self.M93= "اصلاح لغت:    "
        self.M94='ویرایش  بدون ترجمه'
        self.M95='ویرایش متن'
        self.M96='تبدیل فایل و ترجمه'
        self.M97='استخراج جداول از فایل پی دی اف'
        self.M98='استخراج  متن از فایل و ترجمه'
        self.M99='استخراج  متن از فایل'
        self.M100='انگلیسی'
        self.M101='آلمانی'
        self.M102='فارسی'
        self.M103='چینی'
        self.M104='شناسایی خودکار زبان '
        self.M105='لطفا از تشابه زبان فایل انتخابی با زبان مبدا اطمینان حاصل فرمایید '
        self.M106='ایتالیایی'
        self.M107='هلندی'
        self.M108='فرانسوی'
        self.M109='عربی'
        self.M110='اسپانیولی'
        self.M111='روسی'
        self.M112='ترکی'
        self.M113='پرتقالی'
        self.M114='اندونزی'
        self.M115='هندی'
        self.M116='ژاپنی'
        self.M117='اردو'
        self.M118='مترجم آفلاین'
        self.M119='مترجم گوگل '
        self.M120='مترجم بینگ'
        self.M121='مترجم مای مموری'
        self.M122='مترجم دیپ'
        self.M123='مترجم یاندکس'
        self.M124= 'انتخاب مترجم'
        self.M125=' کلمات متضاد و مترادف'
        self.M126='mj.taraz@yahoo.com'
        self.M127=' وارد نمودن متن و اصلاح جملات'
        self.M128='اصلاح نيمه خودکار کلمات '
        self.M129='اصلاح خودکار لغت'
        self.M130='لغت درست را انتخاب کنيد'
        self.M131='لغت صحيح را انتخاب و يا در پنجره زير بنويسيد '
        self.M132='تاييد '
        self.M133="کلمه یا متن مورد نظر را پیدا و اصلاح یا جایگزین کنید تغیرات شما بعد از فشردن تایید  در برنامه اعمال می شود لیست متناسب با انتخاب زبان است"
        self.M134="ذخيره لغت جديد در مرجع لغت !!!"
        self.M137=" به مرجع لغت اضافه شود؟"
        self.M135="لیست کلمات مشابه"
        self.M136='عدم نمایش این پنجره '
        self.M138="تایید"
        self.M139= "بنویسید"
        self.M140= "زبان مبدا و مقصد يکي است  عدم ترجمه  (حالت ويرايش متن )"
        self.M141='خطا در تبدیل فایل به پی دی اف اگر آفیس نصب نیست آن را نصب فرمایید'
        self.M142='عدم ذخیره خودکار کلمه جایگزین'
        self.M143=' جستجو گر کتاب راه اندازی نشد'
        self.M144="لغو ذخیره فایل ترجمه شده؟"
        self.M145='در حال ذخیره فایل'
        self.M146="2 تغییر چیدمان حروف"
        self.M150="رنگ متن"
        self.M151="اصلاح لغت"
        self.M152="3  تغییر چیدمان حروف"
        self.M153=" عدم تغییر چیدمان متن و حروف"
        self.M154=" فرهنگ لغت فارسی"
        self.M155="براي اصلاح لغت غير فارسي بايد يک زبان غير فارسي را انتخاب نماييد"
        self.M156=" زبان مبدا  با زبان دیکشنری یا کتاب  یکی نیست "
        self.M157="متاسفانه کلمات مترادف فقط برای فارسی موجود است"
        self.M158= 'لغت یا متن مورد نظر  را بنویسید و  یا از متن زیر کپی  کنید'
        self.M159= 'فرهنگ لغت ابجد به فارسی'
        self.M160= 'فرهنگ لغت عربی به فارسی'
        self.M162= "بارگذاری فایل از شکیبایی شما سپاسگزاریم "
        self.M163= "گوگل"
        self.M164= "ویکی پدیا"
        self.M165= 'دانلود / اینترنت'
        self.M166= 'فرهنگ لغت انگلیسی به فارسی'
        self.M167= 'فرهنگ لغت فارسی به انگلیسی'
        self.M168= 'فرهنگ لغت فارسی به عربی'
        self.M169= 'فرهنگ لغت عربی'
        self.M170= " ترجمه قرآن"
        self.M171= ' حافظ'
        self.M172= ' قرآن'
        self.M173= "فرهنگ لغت انگلیسی"
        self.M174= " نهج البلاغه"
        self.M175= " نهج الفصاحه"
        self.M176= "کتب مذهبی"
        self.M177= "فرهنگ لغت"
        self.M178= "جستجوی دقیق"
        self.M179= "صحیفه سجادیه"
        self.M180="جستجوی جمله"
        self.M181="فرهنگ لغت دهخدا"
        self.M182="4  تغییر چیدمان حروف"
        self.M183="بحارالانوار"
        self.M184="قانون در طب"
        self.M185="طب سنتی"
        self.M186="علم رجال"
        self.M187="آیین دادرسی"
        self.M188="فرهنگ لغت آلمانی به فارسی"
        self.M189="شاهنامه"
        self.M190="اصول کافی"
        self.M191="دیوان سعدی"
        self.M192="قطره"
        self.M193="قطره فارسی"
        self.M194="قرآن با اعراب"
        self.M195="صفحه جستجو"
        self.M196="تغییر چیدمان کلمات متن در فایل خروجی"
        self.M197=" هیچ متنی پیدا نشد یا متن ورودی خالی است"
        self.M198="جستجو"
        self.M199="قانون طب ابوعلی سینا"
        self.M200="متن ورودی"
        self.M201="ترانسلیت کام"
        self.M202="متن خروجی"
        self.M203="کتب مهندسی"
        self.M204="درمان بیماری داخلی"
        self.M205="درمان بیماری  اندام"
        self.M207="محاسبه ابجد"
        self.M208="المیزان"
        self.M209="مفردات راغب"
        self.M210="مفردات راغب با ترجمه"
        self.M211="ترجمه المیزان"
        self.M212="گزینه جستجو"
        self.M213="غیر فعال"
        self.M214="نمایش صفحه جستجو جداگانه برای "
        self.M215="فایل با فرمت پی دی اف ذخیره نشد اگر آفیس نصب نیست آن را نصب و دوباره سعی کنید در حال حاضر فایل با فرمت نات پد ذخیره می شود اگر متن درست نیست دوباره سعی فرمایید"
        self.M216="SMD  کدهای قطعات الکترونیکی"
        self.M217='استاندارد آلمان DIN بخش یک'
        self.M218='استاندارد آلمان DIN بخش دو'
        self.M219='استاندارد آلمان DIN بخش سه'
        self.M220='استاندارد آلمان DIN بخش چهار'
        self.M221='استاندارد آلمان DIN بخش پنج'
        self.M222='مهندسی مکانیک'
        self.M223='مهندسی برق'
        self.M224='مهندسی الکترونیک'
        self.M225='مهندسی عمران'
        self.M226='مهندسی کامپیوتر'
        self.M227='کتب شعر'
        self.M228='شهریار'
        self.M229='نیما یوشیج'
        self.M230='مولوی'
        self.M231='سهراب سپهری'
        self.M232='خیام'
        self.M233='بابا طاهر'
        self.M234='عنصری بلخی'
        self.M235=' رودکی'
        self.M236='مترجم آفلاین نصب نیست اگر فایل نصب آفلاین را ندارید لطفا برای دریافت مترجم آفلاین ایمیل بزنید mj.taraz@yahoo.com'
        self.M237='این کتاب  نصب نیست اگر فایل نصب  را ندارید لطفا برای دریافت   ایمیل بزنید mj.taraz@yahoo.com'
        self.M238='کلنیک حقوقی ایران'
        self.M239="دانلود شروع شود ؟"
        self.M240= ' لینک دانلود را وارد یا جایکذاری و دکمه دانلود را بزنید \n (هشدار از اعتبار لینک جهت دانلود اطمینان حاصل فرمایید که حاوی ویروس یا تروجان نباشد)'
        self.M241="کتب طب"
        self.M242="طب الرضا"
        self.M243="رساله"
        self.M244="عیون الرضا"
        self.M245="اصلاح همه کلمات مشابه"
        self.M246="یافتن کلمه صحیح"
        self.M247="ویرایش لیست کلمات اضافه شده"
        self.M248=" ! ذخیره" 
        self.M249="ویرایش لیست کلمات جایگزین شده"
        self.M250="کلمه جایگزین پیشنهادی :"
        self.M251="پردازش  پی دی اف برای زبانهای عربی فارسی و اردو درست انجام نمی شود و ممکن است متن درست استخراج نشود آیا ادامه می دهید ؟"
        self.M252="بارگذاری کامل شد"
        self.M253="بازنشانی نرم افزار"
        self.M254=' برنامه را در حالت Administrator   اجرا کنید'
        self.M255="جستجو در متن ورودی"
        self.M256="در حال استخراج جدول از پی دی اف"
        self.M257 ="جایگزین شد با "
        self.M258 ="اجرای تغییرات"
        self.M259="ریورسو"
        self.M260="MJT Gv2"
        self.M261="MJT Gv1"
        self.M262="الغدیر"
        self.M263="توضیح المسایل سیستانی"
        self.M264="گوهر خود را هویدا کن"
        self.M265="موعظه آیت الله ضیاء آبادی"
        self.M266="حج برنامه تکامل"
        self.M267="کتاب آشپزی"
        self.M268='جدول استخراج شده ترجمه شود ؟'
        self.M269='جستجوی عمیق'
        self.M270='اصلاح عمیق'
    def lang_en(self):
        messages = {
        "M001": "Style:", "M002": "From Language:", "M003": "To Language:", "M004": "Search in Book", "M005": "Change app size", "M1": "File",
        "M2": "New", "M3": "Open Text File", "M4": "Save as TXT File", "M411": "Save Input Text as TXT File", "M6": "Exit", "M7": "Reverse Word (2)",
        "M8": "Edit", "M9": "convert:", "M10": "Change Text Color", "M11": "Space Correction", "M12": "Reverse Word PDF", 
        "M13": "Farsi correction (semi)", "M14": "Farsi Text Edition", "M15": "Word correction (Auto)", "M16": "Farsi Number convert", 
        "M17": "Farsi Default Font and format for PDF and word File", "M18": "Add Source Text in Translated Text", "M19": "Cut", "M20": "Copy", 
        "M21": "Paste", "M22": "GUI Error", "M23": "An Error Occurred", "M24": "Please Wait Until the Other Part of the Translator Finishes Work", 
        "M25": "No File Selected", "M26": "Insert file", "M27": "Select File Error", "M28": "Text from PDF Inserted to the Entry", 
        "M29": "PDF RE Process Error", "M30": "Table Extract Enabled from Edit Menu Extracting and Translation Table from PDF Will Start", 
        "M31": "PDF Text Error", "M32": "No File Selected for Save Translated", "M33": "Process is Finished.", "M34": "Removed", 
        "M35": "PDF Process Error", "M36": "PDF re Processing in Text Mode Only", "M37": "PDF Process", "M38": "Document Processing", 
        "M39": "Document Process Finished", "M40": "Error process_docx", "M41": "Translation and Save File", "M42": "Successful!", 
        "M43": "Docx rapidfuzz.process Error", "M44": "Use learning translate", "M45": "PDF Process: Extracting Tables...", "M46": "PDF Process: Extracting Tables Canceled", 
        "M47": "PDF Process: Installation Was Canceled", "M48": "Extracting Tables from PDF Please Wait...", "M49": "Table Extracted:", 
        "M50": "Tables Extracted Successfully!", "M51": "No Save Directory Selected for Translation", "M52": "Output File Saved To", 
        "M53": "PDF Table Extract:", "M54": "No table can be extracted from selected file", "M55": "Book Not Found:", "M56": "Translate Error:", 
        "M57": "Translation", "M58": "Please Wait...", "M59": "translate_doc Error:", "M60": "Translation Error! Please Make Sure Requested Packages Downloaded or Using an Online Translator", 
        "M61": "Translation Error! Please Check the Internet Connection or Using Offline Translator", "M62": "Bing: Internet Connection", 
        "M63": "Argos Translator", "M64": "Checking Language Package if Necessary", "M65": "Unknown Error with Argos Offline Translator!", 
        "M66": "The Language Package Not Found", "M67": "Installation Error While Downloading:", "M68": "Installation Language Error", 
        "M69": "Reinstall Language Packages?", "M70": "Retry Downloading...", "M71": "Not Available Language Support for Argos:", 
        "M72": "Change Translator and Retry", "M73": "Download", "M74": "Canceled Try Translate with Selecting Other Translators", 
        "M75": "Downloading", "M76": "Language Package...", "M77": "Download Finished", "M78": "Download Error", "M79": "Please Retry", 
        "M80": "Google", "M81": "MyMemory Error", "M82": "Deep Error", "M83": "Yandex Error", "M84": "Corrections", "M85": "Would You Like to Open", 
        "M86": "Confirmation", "M87": "Gs Installed System Must Restart to Take Effect, Would You Like to Restart?", "M88": "For Extract Table Ghostscript Must Install in (Default Directory) and Restart System Would You Like to Install Ghostscript?", 
        "M89": "Completed Program will be Restarted to Take Effect", "M90": "Requested Package Not Installed! Would You Like to Download Language Package", 
        "M91": "Working on File Please Wait:", "M92": "Selected PDF File Not Supported or Some Unknown Error for docx2.convert with the Same Structure, Would You Like to Extract Text Only?", 
        "M93": "Corrections WORD:", "M94": "Edit text without translate", "M95": "Correcting Text", "M96": "Convert and Translate", 
        "M97": "Table Extract and Translate (For PDF Files Only)", "M98": "Text Extract from File and Translate", "M99": "Text Extract from File", 
        "M100": "English", "M101": "German", "M102": "Farsi", "M103": "Chinese", "M104": "Auto_Detect_Language", 
        "M105": "Please Make Sure the Source Language is the Same as Your PDF File", "M106": "Italian", "M107": "Dutch", "M108": "French", 
        "M109": "Arabic", "M110": "Spanish", "M111": "Russian", "M112": "Turkish", "M113": "Portuguese", "M114": "Indonesian", 
        "M115": "Hindi", "M116": "Japanese", "M117": "Urdu", "M118": "Argos", "M119": "Google", "M120": "Bing", "M121": "MyMemory", 
        "M122": "Deep Translator", "M123": "Yandex", "M124": "Translator:", "M125": "Synonymous Word (Farsi)", 
        "M126": "mj.taraz@yahoo.com \n call +989914604366", "M127": "Entry Text and Correction", "M129": "Word Correction", 
        "M128": "Word Correction (Semi Auto)Farsi only", "M130": "Select Correct Word", "M131": "Enter or Select Correct Word", "M132": "OK", 
        "M133": "Word Correction", "M134": "Save New Word in Dictionary!", "M135": "Continue Word Correction for Next Text", 
        "M136": "Do not show this window", "M137": "Add in Dictionary?", "M138": "Confirm", "M139": "Write Text", 
        "M140": "The Source and Target Language are the Same No Translations (Text Correction Mode)", "M141": "Microsoft Office Not Installed for docx2.convert Word Office to PDF File", 
        "M142": "Skip Auto Save Replaced Words", "M143": "Search book Not Started", "M144": "Really Cancel Save docx2.convert and Translated File?", 
        "M145": "Saving File", "M146": "Reverse Word (2)", "M150": "Text Color", "M151": "Word Correction", "M152": "Reverse Word (3)", 
        "M153": "No Change Text and Char", "M154": "Farsi Dictionary", "M155": "For Latin Word Correction Must Select Other Languages Not Farsi", 
        "M156": "Please Set Source Language with Dictionary or Book", "M157": "Synonymous Words Only Work for Farsi. Select Farsi from Language", 
        "M158": "Write Word or Copy from Text", "M159": "Abjad to Farsi DICTIONARY", "M160": "Arabic to Farsi DICTIONARY", 
        "M162": "Loading File...", "M163": "Google", "M164": "Wikipedia", "M165": "SEARCH and Download INTERNET.", 
        "M166": "English to Farsi Dictionary.", "M167": "Farsi to English Dictionary.", "M168": "Farsi to Arabic Dictionary.", 
        "M169": "Arabic Dictionary.", "M170": "Quran Farsi", "M171": "HAFEZ", "M172": "Quran Arabic", "M173": "English Dictionary.", 
        "M174": "Nahj_albalaqah", "M175": "Nahj_alfasaha", "M176": "SEARCH IN BOOK", "M177": "DICTIONARY", "M178": "SEARCHING THE SAME WORD", 
        "M179": "Sahife sajjadieh", "M180": "SEARCH BY SENTENCE", "M181": "Dehkhoda Dictionary.", "M182": "Reverse Word (4)", "M183": "beharolanvar", 
        "M184": "Ghanoon_teb","M185": "Teb_sonati", "M186": "Elm-rijal", "M187": "Ayin-dadrasi", "M188": "German_to_Farsi_Dictionary",
        "M189": "Shahname_Ferdosi","M190": "Osul_KAfi","M191": "saadi","M192": "AL_Qatrah","M193": "AL_Qatrah_farsi","M194": "Quran with erab",
        "M195": "SEARCH WINDOW","M196": "Reverse_Text","M197": "Not Found or Text is Empty","M198": "SEARCH","M199": "Qanoun fe teb aboAliSina",
        "M200": "Input text","M201": "translatorCom","M202": "Output text","M203": "Engineering Books","M204": "Diseases and Treatment 1","M205": "Diseases and Treatment 2",
        "M207": "ABJAD Calc","M208": "AL_MIZAN", "M209": "AL_Mofradat", "M210": "AL_Mofradat_Farsi", "M211": "AL_MIZAN_Farsi", "M212": "SEARCH OPTIONS",
        "M213": "INACTIVE","M214": "SHOW SEARCH WINDOW IN SEPARATE WINDOW?","M215": "File Not Saved as PDF. If Microsoft Office Not Installed Please Install and Try Again. In a Moment File Will Save as Notepad. If Format is Wrong Try Again",
        "M216": "SMD Electronics Code","M217": "DIN STANDARD PART 1.", "M218": "DIN STANDARD PART 2.","M219": "DIN STANDARD PART 3.",
        "M220": "DIN STANDARD PART 4.","M221": "DIN STANDARD PART 5.","M222": "MECHANICAL ENG.","M223": "ELECTRICAL ENG.","M224": "ELECTRONIC ENG.",
        "M225": "CONSTRUCTION ENG.","M226": "COMPUTER ENG.","M227": "Poetry Books.","M228": "Shahriar.","M229": "Nima youshij.","M230": "Molana.",
        "M231": "Sohrab sepehri.","M232": "Khayyam.","M233": "Baba Taher.","M234": "Onsori balkhi.","M235": "Roudaki.",
        "M236": "Offline Translator is Not Installed. If You Have No Install File Please Send Email to mj.taraz@yahoo.com to Receive Offline Translator. Thank You.",
        "M237": "This Book is Not Installed. If You Have No Install File Please Send Email to mj.taraz@yahoo.com to Receive file. Thank You.",
        "M238": "IranClinicLaw.","M239": "Would You Like to download?",
        "M240": "Please Insert Download Link or Paste from Clipboard and Select Download Button. (Note: Please Ensure Link is Safe and There is No Virus.)",
        "M241": "Search in Teb","M242": "Teb Reza","M243": "Resaleh","M244": "Oyoun Areza","M245": "Replace All","M246": "Results Corrected Word May:",
        "M247": "New Words List","M248": "Save!","M249": "Saved Replaced Words List","M250": "Order Correct Word is:",
        "M251": "PDF Process for Source Language Not Supported. It Can Be Incorrect Extracting Text. Would You Like Continue?","M252": "Loading complete",
        "M253": "Restarting","M254": "Please run the program in administrator mode to support this feature","M255": "search in input text",
        "M256": "Table extract enabled from edit menu. Extracting and translation table from PDF will start","M257": "Replace with",
        "M258": "Make changes","M259": "Reverso","M260": "googlev2","M261": "MJT Gv1","M262": "algadir","M263": "tozih_almasael","M264": "Gohare khod ra yad kon",
        "M265": "Moezeh","M266": "Haj barnamrh e takamol","M267": "Cooking book","M268": "Would You Like to translate extracted tables?",
        "M269": "Deep search","M270": "Deep corrections"
            }
        for key, value in messages.items():
                setattr(self, key, value)
        button_texts = {
                self.file_button: "Select",
                self.translate_button: "Translate text",
                self.pdf_converting: "PDFconvert",
                self.export_button: "Save",
                self.clear_button: "Text Clear",
                self.paste_button: "Paste",
                self.help_button: "Help",
                self.reset_button: "Restart",
                self.copy_button: "Copy",
                self.stop_button: "Stop",
                self.trans_file_button: "Translate File"
            }
        for button, text in button_texts.items():
                button.setText(text)

    def choose_color(self):
        self.color = QColorDialog.getColor()
        if self.color.isValid():
            self.color_code = self.color.name()
            t = self.color_code
            self.save_state(5,t)
            self.input_console.setStyleSheet(f"color: {self.color_code};")
            self.output_console.setStyleSheet(f"color: {self.color_code};")
    def update_font2(self):
            t = self.size_app.value()
            self.app_size = float(1 + float(t * 0.1)) 
            self.save_state(10,t) 
            self.positions_and_size()       
    def update_font(self):
        try:
            try:
                self.input_console.setStyleSheet(f"color: {self.color1};")
                self.output_console.setStyleSheet(f"color: {self.color1};")
            except:
                self.input_console.setStyleSheet('color:#0000ff;')
                self.output_console.setStyleSheet('color:#0000ff;')
            size = self.size_box.value()
            size_app = self.size_app.value()
            font = self.font_box.currentText()
            new_font = QFont(font, int(size))
            self.output_console.setFont(new_font)
            self.input_console.setFont(new_font)
            t = str(size)
            t2 = self.font_box.currentText()
            self.app_size = float(1 + float(size_app * 0.1))
            self.positions_and_size() 
            self.save_state(6,t)
            self.save_state(7,t2)
        except FileNotFoundError:
            try:
                self.defultthem()
            except Exception as e:
                self.update_notification(self.M254)
                print(f"Failed to write default theme: {e}")
        except Exception as e:
            print(f"An error occurred: {e}")
            self.update_notification(self.M254)


    def new_file(self):
        self.input_console.clear()
        self.current_file = None

    def open_file(self):
        filetypes = "Text File (*.txt);;All Files (*);"
        self.file, _ = QFileDialog.getOpenFileName(self, "Select File", self.lastDir, filetypes)
        if self.file:
            encodings = ['utf-8', 'latin-1', 'ascii'] 
            for encoding in encodings:
                try:
                    with open(self.file, 'r', encoding=encoding) as file:
                        content = file.read()
                        self.update_input(content)
                        self.current_file = self.file
                    self.lastDir = os.path.dirname(self.file)
                    t = self.lastDir  # replaced the sixth line with a variable 't'
                    self.save_state(18, t)
                    self.process_finish()
                except Exception as e:
                    self.update_input(str(e))
        self.process_finish()

    def open_file2(self):
        encodings = ['utf-8', 'latin-1', 'ascii'] 
        for encoding in encodings:
            try:
                with open(self.file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    self.update_input(content)
                    self.current_file = self.file_path
                self.lastDir = os.path.dirname(self.file_path)
                t = self.lastDir  # replaced the sixth line with a variable 't'
                self.save_state(18, t)
                self.process_finish()
            except Exception as e:
                self.update_input(str(e))
                self.process_finish()

    def save_file(self):
        text_to_save = self.output_console.toPlainText()
        
        self.file_p, _ = QFileDialog.getSaveFileName(self, "Save File", "",
                                                "Text files (*.txt);;All Files (*)")
        if self.file_p:
            try:
                with open(self.file_p, "w", encoding="utf-8") as file:
                    file.write(text_to_save)
                self.current_file = self.file_p
                self.q.information(self,self.M42, "File saved successfully.")
                os.startfile(self.file_p)
            except Exception as e:
                self.info0(self.M254)
                self.q.information(self,"Error", f"Saving the file error: {str(e)}")
    def SaveTranslations(self):
        text_to_save = self.CWL2
        if len(text_to_save)<50:
            return
        self.file_p, _ = QFileDialog.getSaveFileName(self, "Save corrected words File", "", "Text files (*.txt);;All Files (*)")
        if self.file_p:
            try:
                with open(self.file_p, "w", encoding="utf-8") as file:
                    file.write(text_to_save)
                self.current_file = self.file_p
                os.startfile(self.file_p)
            except Exception as e:
                self.info0(self.M254)
                self.q.information(self,"Error", f"Saving the file error: {str(e)}")
    def SaveCorrections(self):
        text_to_save = self.CWL
        if len(text_to_save)<50:
            return
        self.file_p, _ = QFileDialog.getSaveFileName(self, "Save corrected words File", "", "Text files (*.txt);;All Files (*)")
        if self.file_p:
            try:
                with open(self.file_p, "w", encoding="utf-8") as file:
                    file.write(text_to_save)
                self.current_file = self.file_p
                os.startfile(self.file_p)
            except Exception as e:
                self.info0(self.M254)
                self.q.information(self,"Error", f"Saving the file error: {str(e)}")
    def save_file_in(self):
        text_to_save = self.input_console.toPlainText()
        self.file_p, _ = QFileDialog.getSaveFileName(self, "Save corrected words File", "", "Text files (*.txt);;All Files (*)")
        if self.file_p:
            try:
                with open(self.file_p, "w", encoding="utf-8") as file:
                    file.write(text_to_save)
                self.current_file = self.file_p
                self.q.information(self,self.M42, "File saved successfully.")
                os.startfile(self.file_p)
            except Exception as e:
                self.info0(self.M254)
                self.q.information(self,"Error", f"Saving the file error: {str(e)}")
    def cut_text(self):
        try:
            selected_text = self.input_console.toPlainText()
            pyperclip.copy(selected_text)
            self.input_console.clear()
        except Exception as e:
            self.update_notification(f"Error: {str(e)}")
    def copy_text(self):
        try:
            selected_text = self.input_console.toPlainText()
            pyperclip.copy(selected_text)
        except Exception as e:
            self.update_notification(f"Error: {str(e)}")
    def paste_text(self):
        if not self.thread_active:
            try:
                clipboard_text = pyperclip.paste()
                self.update_input(' '+clipboard_text)
            except Exception as e:
                self.update_notification(f"Error: {str(e)}")
        else:
            self.update_notification(self.M24)
    def MainWindow(self):
        try:
            icon_path = os.path.join(self.dir, 'icon.png')
            self.setWindowIcon(QIcon(icon_path))
        except Exception as e:
            self.update_notification(f"Icon path: {str(e)}")  
    def show_help_message(self):
        if  self.thread_active==False:
            try:
                text_path = os.path.join(self.dir, 'Help.json')
                text_path2 = os.path.join(self.dir, 'Help2.json')
                with open(text_path, 'r', encoding='utf-8') as file:
                    text_content = file.read()
                with open(text_path2, 'r', encoding='utf-8') as file:
                    text_content2 = file.read()
                self.clear()
                self.update_input( text_content)
                self.q.information(self,"Help", text_content2)
                os.startfile('help.pdf')
            except :
                pass
    def copy_to_clipboard(self):
        try:
            clipboard_text = self.output_console.toPlainText()
            QApplication.clipboard().setText(clipboard_text)
            self.update_notification(" متن کپی شد")
        except:
            self.update_notification("کپی نشد")
    def paste_from_clipboard(self):
        try:
            if not self.thread_active:
                self.input_console.clear()
                clipboard_text = QApplication.clipboard().text()
                self.input_console.insertPlainText(clipboard_text)
            else:
                self.update_notification(self.M24)
        except:
            self.update_notification("هیچ متنی کپی نشده")
    def select_file(self):
        self.output_console.setReadOnly(True)
        self.CWL = 'لیست کلمات جایگزین شده \n'
        self.extract = False
        try:
            filetypes = "Supported Files (*.pdf *.docx *.xlsx);;PDF files (*.pdf);;Word files (*.docx);;Excel files (*.xlsx);;All Files (*);"
            self.file_p, _ = QFileDialog.getOpenFileName(self, "Select File", self.lastDir, filetypes)
            if not self.file_p:
                self.process_finish()
                self.thread_active = False
                return
            self.lastDir = os.path.dirname(self.file_p)
            self.save_state(18, self.lastDir)
            self.file_path = self.file_p
            if not os.path.exists(self.file_path):
                self.process_finish()
                self.q.information(self, "Error", f"The selected file does not exist: {self.file_path}")
                self.thread_active = False
                return
            file_name = os.path.basename(self.file_path)
            self.info3(file_name)
            self.save_state(17, self.file_path)
            if not self.thread_active:
                if self.file_path:
                    self.thread_active = True
                    self.tranc_err = False
                    self.run_document()
                else:
                    self.update_notification("No file path selected.")
            else:
                self.update_notification("Thread is already active.")
        except Exception as e:
            self.process_finish()
            self.update_notification(f"Error: {str(e)}")
    def update_progress(self, message):
        print(message)
    def trans_file(self):
        if self.thread_active==True:
           self.update_notification(self.M24)
           return
        self.CWL='لیست کلمات جایگزین شده \n'
        self.extract=False
        try:
            if len(self.file_path)>3:
                self.thread_active=True
                self.tranc_err=False
                self.run_document()
            else:
                self.update_notification(self.M25)
                self.process_finish()
        except Exception as e:
            self.process_finish()
            self.update_notification(f"  {self.M27}   : {str(e)}")
    def run_document(self):
        self.input_console.textChanged.disconnect()

        if self.source == self.M104 :
            self.detect_language()
        self.open_5()
        self.fileMode=True
        self.argose_err = False
        self.stops=False
        try:
            self.clear()
            self.thread_active=True
            file_name = os.path.splitext(os.path.basename(self.file_path))[0]
            timestamp = datetime.datetime.now().strftime("%Y %m %d _ %H-%M")
            filename2 = f"EDITED {timestamp} {file_name}"
            directory = os.path.dirname(self.file_path)
            if self.file_path.endswith(".pdf") or self.file_path.endswith(".PDF") or self.pdf_convert==True:
                if self.from_code in ['fa','ar','ur']:
                    confirmation = self.qq(self,self.M86,f"{self.M251} {self.source}")
                    if not confirmation==16384:
                        self.stop()
                        return
                self.docx2=True
                self.file_path2 = os.path.join(directory, filename2+'.docx')
                self.file_path3 = os.path.join(directory, filename2+'.xlsx')
                self.file_path4 = os.path.join(directory, filename2+'.xlsx')
                self.process_pdf_file()
            elif self.file_path.endswith(".docx") or self.file_path.endswith(".DOCX") :
                self.docx=True
                self.docx2=True
                self.skip_rtl = True
                self.file_path2 = os.path.join(directory, filename2+'.docx')
                self.process_docx_file()
            elif self.file_path.endswith(".xlsx") or self.file_path.endswith(".XLSX") :
                self.skip_rtl = True
                self.docx=True
                self.file_path2 = os.path.join(directory, filename2+'.xlsx')
                self.file_path3 = os.path.join(directory, filename2+'ext.xlsx')
                self.process_xlsx_file()
            else:
                self.open_file2
            if self.corrections==True:
                self. SaveCorrections()
            self.input_console.textChanged.connect(self.iventtext)
        except:
            self.process_finish()
            self.update_notification(self.M25)
            self.input_console.textChanged.connect(self.iventtext)

    def process_pdf_file(self):
        self.stops = False
        try:
            if self.document_mode_box.currentText() in [self.M96] or self.pdf_convert:
                self.process_pdf2()
            elif self.document_mode_box.currentText() == self.M97:
                self.file_path2 = self.file_path4
                self.extract_tables()
            elif self.document_mode_box.currentText() == self.M98:
                self.skip_rtl = True
                self.PDF_Text_RTL()
            else:
                self.skip_rtl = True
                self.process_pdf()
            self.update_notification(self.M42)
        except Exception as e:
            self.skip_rtl = False
            self.update_notification(f"  {self.M23}   : {str(e)}")
    def process_docx_file(self):
        self.thread_active = True
        if self.document_mode_box.currentText() in [self.M96, self.M97]:
            self.process_docx2()
            self.rev = True
        else:
            self.process_docx()
        self.update_notification(self.M42)
    def process_xlsx_file(self):
        self.thread_active = True
        if self.document_mode_box.currentText() in [self.M96, self.M97]:
            self.process_xlsx2()
        else:
            self.process_xlsx()
        self.update_notification(self.M42)
    def pdf_converter(self):
        try:
            self.pdf_convert = True
            if self.from_code in ['fa', 'ar', 'ur']:
                confirmation = self.qq(self, self.M86, f"{self.M251} {self.source}", self.q.StandardButton.Yes | self.q.StandardButton.No)
                if confirmation == self.q.StandardButton.No:
                    self.stop()
                    return
            self.select_file()
        except Exception as e:
            self.update_notification(f"  {self.M23}  : {e}")
    def process_pdf(self):
        try:
            self.update_notification(self.M37)
            if self.document_mode_box.currentText()==self.M97 :
               self.update_notification(self.M35)
               self.file_path2 = self.file_path4
               self.extract_tables()
               return
            else:
                reader = pypdf.PdfReader(self.file_path)
                full_text = ""
                for page in reader.pages:
                    full_text += page.extract_text() + "\n"
                    if self.stops==True:
                        break
                    self.update_input(full_text+'\n')
                if self.pdf_convert==True:
                    self.full_text=full_text
                    self.pdf_convert.set(1)
                    self.thread_active = False
                    self.export_docx()
                    return
                self.process_finish()
        except Exception as e:
                self.update_notification(f"  {self.M31}: {str(e)}")
                self.process_finish()
    def process_pdf2(self):
            try:
                file_info = (f"{self.M91}   {  os.path.basename(self.file_path)}")
                self.update_notification(file_info)
                directory = os.path.dirname(self.file_path)
                file_name = os.path.splitext(os.path.basename(self.file_path))[0]
                docx_file = os.path.join(directory, file_name+'_temp.docx')
                cv = pdf2.Converter(self.file_path)
                cv.convert(docx_file, start=0, end=None)
                cv.close()
                doc = dox.Document(docx_file)
                translated_texts = {}
                processed_runs = []
                for para in doc.paragraphs:
                    self.lang_code()
                    if para.text.strip():
                        para.text1=self.space_correction(f"{para.text}")
                        if self.from_code in ["fa", "ar", "ur"]:
                            if  self.rtl_format_true==True:
                                para.text1=self.auto_reverse_mix_text(para.text1)
                        if self.from_code != self.to_code:
                            if self.from_code in ["fa", "ar", "ur"]:
                                if self.to_code not in ["fa", "ar", "ur"]:
                                    para_properties = para._element.get_or_add_pPr()
                                    para.alignment = 0
                                    para_properties.rtl = False
                            elif self.from_code not in ["fa", "ar", "ur"] and self.to_code in ["fa", "ar", "ur"]:
                                para_properties = para._element.get_or_add_pPr()
                                para.alignment = 2
                                para_properties.rtl = True
                        self.update_input( para.text1 + '\n')
                        if para.text1 not in translated_texts:
                            translated_texts[para.text1]=self.translate_text(para.text1)
                        translated_text = translated_texts[para.text1]
                        translation_applied = False
                        for run in para.runs:
                            if self.stops==True:
                                break
                            if run in processed_runs:
                                continue
                            if not 'graphicData'  in run._r.xml:
                                if run.text.strip():
                                    if run.font.size:
                                        font_size = run.font.size.pt
                                        run.font.size = dox.shared.Pt(font_size)
                                    else:
                                        font_size = 12
                                    run.font.size = dox.shared.Pt(font_size)
                                    if not self.using_orginal_text_enable:
                                       run.clear()
                                    if not translation_applied:
                                        if  self.rtl_format_true==True:
                                            if  (not self.from_code in ["fa", "ar","ur"] and  self.to_code in ["fa", "ar","ur"]):
                                                run.font.rtl = True
                                        run.add_text(translated_text)
                                        translation_applied = True
                        processed_runs.append(run)
                        if self.using_orginal_text_enable:
                           run.clear()
                docx_file2 = self.file_path2
                if not docx_file2:
                    self.process_finish()
                    os.remove(docx_file)
                    return
                self.update_notification(self.M145)
                translated_docx_file = os.path.splitext(docx_file2)[0] + '.docx'
                try:
                  doc.save(translated_docx_file)
                except:
                    self.update_notification(f"فایل ادیت شده قبلی باز است ")
                    translated_docx_file = os.path.splitext(docx_file2)[0] + '_2.docx'
                    doc.save(translated_docx_file)
                file_name_2 = os.path.basename(translated_docx_file)
                confirmation = self.qq(self,self.M86,f"{self.M42} {self.M85}  {file_name_2}")
                os.remove(docx_file)
                if confirmation==16384:
                    os.startfile(translated_docx_file)
                self.process_finish()
                self.update_notification(f"  {self.M41} {self.M42}")
            except Exception as e:
                self.process_finish()
                self.update_notification( f"{self.M35} : {str(e)}")
    def PDF_Text_RTL(self):
        self.tranc_err=False
        try:
            doc = dox.Document()
            pdf = pypdf.PdfReader(self.file_path)
            for page in pdf.pages:
                self.text1 = page.extract_text('')
                self.text2=self.space_correction(self.text1)
                self.update_input(self.text2 +'\n')
                self.thread_active = True
                if self.stops==True:
                    break
                translated_text = self.translate_text(self.text2 )
                para = doc.add_paragraph()
                if self.target in [self.M102,self.M109]:
                    para.alignment = dox.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
                para.add_run(translated_text)  # Add translated text to the paragraph
            docx_file2 = self.file_path2
            if not docx_file2:
                self.process_finish()
                self.update_notification(self.M32)
                return
            translated_docx_file = os.path.splitext(docx_file2)[0] + '.docx'
            try:
                doc.save(translated_docx_file)
            except:
                self.update_notification(f"فایل ادیت شده قبلی باز است ")
                translated_docx_file = os.path.splitext(docx_file2)[0] + '_2.docx'
                doc.save(translated_docx_file)

            self.process_finish()
            confirmation = self.qq(self,self.M86,f"{self.M85} {translated_docx_file}")
            if confirmation==16384:
                os.startfile(translated_docx_file)
        except Exception as e:
            self.process_finish()
            self.update_notification(f"  {self.M37} :{str(e)}")
    def process_docx(self ):
        self.update_notification(self.M38)
        self.skip_rtl=True
        try:            
            doc = dox.Document(self.file_path)            
            for para in doc.paragraphs:
                if self.stops==True:                    
                    break 
                para.text2=self.space_correction(para.text)      
                self.update_input( para.text2 + '\n')
                if self.document_mode_box.currentText() ==self.M98:
                    self.translate_text(para.text2)
            self.thread_active =False
            if self.document_mode_box.currentText() ==self.M98: 
                self.export_docx()  
            self.process_finish()                    
            self.update_notification(self.M39)
            return
        except Exception as e:            
            self.update_notification(f"  {self.M40}  {str(e)}")
            self.process_finish() 
    def process_docx2(self):
        self.skip_rtl = True
        try:
            file_info = f"{self.M91} {os.path.basename(self.file_path)}"
            self.update_notification(file_info)
            doc = dox.Document(self.file_path)
            translated_texts = {}
            processed_runs = []
            for para in doc.paragraphs:
                if self.stops:
                    break

                if para.text.strip():
                    para.text1 = self.space_correction(para.text)
                    self.update_input(para.text1 + '\n')

                    if para.text1 not in translated_texts:
                        translated_texts[para.text1] = self.translate_text(para.text1)
                    translated_text = translated_texts[para.text1]

                    translation_applied = False
                    if self.from_code != self.to_code:
                        para_properties = para._element.get_or_add_pPr()
                        if self.from_code in ["fa", "ar", "ur"] and self.to_code not in ["fa", "ar", "ur"]:
                            para.alignment = 0
                            para_properties.rtl = False
                        elif self.from_code not in ["fa", "ar", "ur"] and self.to_code in ["fa", "ar", "ur"]:
                            para.alignment = 2
                            para_properties.rtl = True

                    for run in para.runs:
                        if self.stops:
                            break
                        if run in processed_runs:
                            continue

                        if 'graphicData' not in run._r.xml and run.text.strip():
                            font_size = run.font.size.pt if run.font.size else 12
                            run.font.size = dox.shared.Pt(font_size)

                            if not self.using_orginal_text_enable:
                                run.clear()
                            if not translation_applied:
                                if self.rtl_format_true and (self.from_code not in ["fa", "ar", "ur"] and self.to_code in ["fa", "ar", "ur"]):
                                    run.font.rtl = True
                                run.add_text(translated_text)
                                translation_applied = True
                            if run.font.superscript:
                                run.font.superscript = True

                            processed_runs.append(run)
                            
                        if self.using_orginal_text_enable:
                            run.clear()
                        self.wait()
            docx_file2 = self.file_path2
            if not docx_file2:
                self.process_finish()
                return

            self.update_notification(self.M145)
            translated_docx_file = os.path.splitext(docx_file2)[0] + '.docx'
            try:
                doc.save(translated_docx_file)
            except:
                self.update_notification(f"فایل ادیت شده قبلی باز است ")
                translated_docx_file = os.path.splitext(docx_file2)[0] + '_2.docx'
                doc.save(translated_docx_file)


            confirmation = self.qq(self, self.tr(self.M42), f"{self.tr(self.M85)} \n\n  {self.tr(translated_docx_file)}", self.q.StandardButton.Yes | self.q.StandardButton.No, self.q.StandardButton.No)
            if confirmation == self.q.StandardButton.Yes:
                os.startfile(translated_docx_file)

            self.update_notification(f"{self.M41} {self.M42}")
            self.process_finish()

        except Exception as e:
            self.process_finish()
            self.q.information(self, self.M43, f"{self.M23} : {str(e)}")
            self.update_notification(f"{self.M43} : {str(e)}")

    def process_xlsx2(self):
        self.excel = True
        try:
            # Extract file info and update notification
            file_info = f"{self.M91}  {os.path.basename(self.file_path)}"
            self.update_notification(file_info)

            # Load workbook and prepare translated file path
            workbook = npyxl.load_workbook(self.file_path)
            translated_file = self.file_path2
            if not translated_file:
                self.process_finish()
                self.update_notification(self.M32)
                return

            # Iterate through sheets in the workbook
            for sheet_name in workbook.sheetnames:
                try:
                    sheet = workbook[sheet_name]
                    data = sheet.values
                    columns = next(data)[0:]
                    df = pd.DataFrame(data, columns=columns)
                except Exception:
                    continue

                # Process and translate data in the sheet
                translated_data = []
                if self.stops:
                    break
                for row in df.values:
                    translated_row = []
                    if self.stops:
                        break
                    for cell in row:
                        self.excel = True
                        try:
                            if self.stops:
                                break
                            if isinstance(cell, (float, int)):
                                translated_text = cell
                                cell = str(cell)
                                self.update_input(cell + '\n')
                            elif cell is None or len(str(cell)) < 1:
                                translated_text = ""
                            else:
                                if self.extract:
                                    self.thread_active = True
                                    self.stops = False
                                self.update_input(cell + '\n')
                                translated_text = self.translate_text(cell)
                        except Exception as e:
                            self.update_notification(f"{self.M23} : {str(e)}")
                            self.update_input(cell + '\n')
                            translated_text = cell
                        translated_row.append(translated_text)
                    translated_data.append(translated_row)

                # Write translated data back to the sheet
                translated_df = pd.DataFrame(translated_data)
                for r_idx, row in enumerate(translated_df.values, 2):
                    for c_idx, value in enumerate(row, 1):
                        try:
                            sheet.cell(row=r_idx, column=c_idx, value=value)
                            self.wait()
                        except Exception:
                            pass

            # Save translated workbook and open it
            workbook.save(translated_file)
            self.process_finish()
            os.startfile(translated_file)

        except Exception as e:
            self.process_finish()
            self.q.information(self, self.M23, f"{self.M23} : {str(e)}")
            self.update_notification(f"{self.M23} : {str(e)}")

    def extract_tables(self):
        try:
            self.update_notification(self.M45)
            if not self.find_folders_2():
                if self.qq(self, self.M86, f"{self.M88}"):
                    source_path = os.path.join(self.dir, 'gs10012w64.exe')
                    os.startfile(source_path)
                    if self.find_folders_2():
                        if self.qq(self, self.M86, f"{self.M87}"):
                            os.system("shutdown /r /t 3")
                    else:
                        self.update_notification(self.M45)
                        self.process_finish()
                        return
        except Exception as e:
            self.update_notification(f"{self.M47} {str(e)}")
            self.process_finish()
            return

        try:
            self.update_notification(self.M48)
            # Read tables from the PDF file using Camelot
            tables = camelot.read_pdf(self.file_path, flavor='lattice', pages='all')
            destination = self.file_path3
            if not destination:
                self.update_notification(self.M47)
                self.process_finish()
                return

            all_data = pd.DataFrame()  # Initialize an empty DataFrame to hold all data

            # Iterate through tables and concatenate data
            for i, table in enumerate(tables):
                if self.stops:
                    break
                df = table.df
                all_data = pd.concat([all_data, df], ignore_index=True)  # Concatenate data
                self.wait()

            # Write the concatenated data to a single sheet
            with pd.ExcelWriter(destination, engine='openpyxl') as writer:
                all_data.to_excel(writer, sheet_name='MergedTables', index=False)

            workbook = npyxl.load_workbook(destination)
            worksheet = workbook['MergedTables']

            # Adjust column widths
            for col in worksheet.columns:
                max_length = 0
                column = col[0].column_letter  # Get the column name
                for cell in col:
                    try:
                        if cell.value and not isinstance(cell.value, (float, int)) and len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except Exception:
                        pass
                self.wait()
                adjusted_width = max_length + 2
                worksheet.column_dimensions[column].width = adjusted_width

            workbook.save(destination)
            self.process_finish()
            self.update_notification(f"{len(tables)} {self.M50}")

            # Open the file if conversion is enabled
            if self.pdf_convert:
                self.q.information(self, self.M42, f"{len(tables)} {self.M50}!")
                self.update_notification(self.M33)
                os.startfile(destination)
                self.process_finish()
                return

            # Handle RTL languages
            if self.from_code in ['fa', 'ar', 'ur']:
                self.extract = True
                self.skip_rtl = False

            if not destination:
                self.update_notification(self.M51)
                self.process_finish()
                return

            self.update_notification(f"{self.M52} {destination}")
            self.file_path = destination

            # Confirmation before processing XLSX
            confirmation = self.qq(self, self.M86, f"{self.M268} {len(tables)} {self.M50}!", self.q.StandardButton.Yes | self.q.StandardButton.No, self.q.StandardButton.No)
            if confirmation == self.q.StandardButton.Yes:
                self.process_xlsx2()
                self.process_finish()
                return

            os.startfile(destination)
            self.update_notification(self.M42)
            return

        except Exception as e:
            self.process_finish()
            self.q.information(self, "PDF Table extract", f"{self.M54}: {e}")
            self.update_notification(f"{self.M54}: {e}")

    def find_folders_2(self):
        base_path = r"C:\Program Files\gs\gs10.01.2"
        base_path_2 = r"C:\Program Files (x86)\gs\gs10.01.2"
        if os.path.exists(base_path) or os.path.exists(base_path_2):
            return True
        else:
            return False
    def process_xlsx(self):
        
        self.output_console.clear()
        try:
            workbook = npyxl.load_workbook(self.file_path, data_only=True)
            for sheet in workbook.sheetnames:
                #self.update_input( f"Sheet: {sheet}\n")
                worksheet = workbook[sheet]
                #row_text)
                row_text=self.space_correction(sheet)
                if self.stops==True:
                    break
                for row in worksheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                   # row_text_formatted=self.format_text(row_text)
                    self.update_input( row_text + '\n')
            self.process_finish()
        except Exception as e:
            self.process_finish()
            self.update_input( f"{self.M23}: {e}\n")
            return
    def export_docx(self):
        if not self.thread_active:
            self.export_docx_ = True
            try:
                doc = dox.Document()
                size = int(self.size_box.value())
                new_font = self.font_box.currentText()
                if self.color_code and isinstance(self.color_code, str) and len(self.color_code) == 7:  # Ensure valid hex color code
                    rgb_color = tuple(int(self.color_code[i:i+2], 16) for i in (1, 3, 5))  # Get the RGB color from hex
                else:  # If no color is selected or invalid color code, use a default color
                    rgb_color = (0, 0, 0)

                if self.pdf_convert:
                    text = self.full_text.split('\n')
                else:
                    text = self.output_console.toPlainText().split('\n')

                if self.pdf_convert or len(text) > 1:
                    
                    file_pat, _ = QFileDialog.getSaveFileName(self, "Export Translated Text", "",
                                                            "Word files (*.docx);;All Files (*)")
                    if not file_pat:
                        return
                    self.file_pat = file_pat

                    for para in text:
                        para1 = self.space_correction(para)
                        paragraph = doc.add_paragraph(para1)
                        for run in paragraph.runs:
                            run.font.size = dox.shared.Pt(size)
                            run.font.name = new_font
                            run.font.color.rgb = dox.shared.RGBColor(*rgb_color)
                        self.wait()

                    doc.save(self.file_pat)
                    os.startfile(self.file_pat)
                    self.process_finish()
                    self.rev = False
                else:
                    self.update_notification(self.M197)
            except Exception as e:
                self.process_finish()
                self.rev = False
                self.q.information(self, self.M23, f"{str(e)}")
        else:
            self.update_notification(self.M24)
        return
    def run_searching(self):
        try:
            file_path=os.path.join(self.dir, 'TarazBook')
            os.startfile(file_path)
        except:
            self.update_notification(self.M143)
    def gift(self):
        try:
            selected_book = self.book_box.currentText()
            if selected_book =='قرآن':
               file_path=os.path.join(self.bookdir, 'FarsiQuran_Vista')
               os.startfile(file_path)
               return
            if selected_book =='اصول کافي' :
               file_path_2=os.path.join(self.bookdir, 'kafi_j1.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='نهج الفصاحه' :
               file_path_2=os.path.join(self.bookdir, 'nf.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='نهج البلاغه' :
               file_path_2=os.path.join(self.bookdir, 'nb.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='موعظه' :
               file_path_2=os.path.join(self.bookdir, 'moezeh.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='سه دقیقه در قیامت' :
               file_path_2=os.path.join(self.bookdir, '3d.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='تمنای وصال':
               file_path_2=os.path.join(self.bookdir, 'ts.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='شعر تمنای وصال':
               file_path_2=os.path.join(self.bookdir, 'ts2.pdf')
               os.startfile(file_path_2)
               return
            else:
               self.update_notification(self.M55)
               return
        except Exception as e:
            self.update_notification(f"   {self.M55}:{e}")
    def searching_book_options(self):
        try:
            book_options = {
                self.M168: ("fa_ar_dic.json", "fa_ar_dic", "S10"),
                self.M167: ("en_fa.json", "en_fa", "S9"),
                self.M159: ("ab_dic.json", "ab_dic", "sp3"),
                self.M160: ("ar_fa_dic.json", "ar_fa_dic", "sp4"),
                self.M154: ("fa_dic.json", "fa_dic", "sp2"),
                self.M169: ("ar_dic.json", "ar_dic", "S11"),
                self.M166: ("en_fa_dic.json", "en_fa_dic", "S11"),
                self.M173: ("en_dic.json", "en_dic", "S15"),
                self.M188: ("gr_to_Fa_dic.json", "gr_to_Fa_dic", "S26"),
                       }
            if self.dict.currentText() != self.M177:
                variant = self.dict.currentText()
            else:
                return False
            for book, (book_name, book_attr, Sattr) in book_options.items():
                self.wait()
                if self.stops:
                    break
                if book == variant:
                    self.ktb = variant
                    self.path_book =os.path.join(self.bookdir, book_name)
                    self.search_process = variant
                    if getattr(self, Sattr):
                        self.spm = True
                        self.ketab_ = set(getattr(self, book_attr))
                        return True
                    split0001 = [self.M184, self.M185, self.M186, self.M187, self.M190, self.M192, self.M262, self.M263, self.M265, self.M266, self.M264, self.M199, self.M242, self.M243, self.M244, self.M267, self.M174, self.M175, self.M193, self.M205, self.M200, self.M202, self.M204]
                    with open(self.path_book, 'r', encoding="utf-8") as f:
                        if variant in split0001:
                            text = f.read()
                            parts = text.split('۰۰۰۱')
                            setattr(self, book_attr, parts)
                        else:
                            setattr(self, book_attr, f.read().splitlines())
                        setattr(self, Sattr, True)
                    self.ketab_ = set(getattr(self, book_attr))
                    return True
            return False
        except Exception as e:
            self.update_notification(f"   {self.M23}:{str(e)}")
            self.info0(self, self.M23, self.M237)
            return False
    def replace_ascii_digits_with_farsi(self,text):
        words = re.split(r'(\W+)', text)
        if self.to_code in ['fa','ar','ur']:
            for i in range(len(words)):
                words[i] = words[i].replace('0', '۰')
                words[i] = words[i].replace('1', '۱')
                words[i] = words[i].replace('2', '۲')
                words[i] = words[i].replace('3', '۳')
                words[i] = words[i].replace('4', '۴')
                words[i] = words[i].replace('5', '۵')
                words[i] = words[i].replace('6', '۶')
                words[i] = words[i].replace('7', '۷')
                words[i] = words[i].replace('8', '۸')
                words[i] = words[i].replace('9', '۹')
        else:
            for i in range(len(words)):
                words[i] = words[i].replace( '۰','0')
                words[i] = words[i].replace( '۱','1')
                words[i] = words[i].replace( '۲','2')
                words[i] = words[i].replace( '۳','3')
                words[i] = words[i].replace( '۴','4')
                words[i] = words[i].replace( '۵','5')
                words[i] = words[i].replace( '۶','6')
                words[i] = words[i].replace( '۷','7')
                words[i] = words[i].replace( '۸','8')
                words[i] = words[i].replace( '۹','9')
        return ''.join(words)
    def text_process(self):
        try:
            if  self.perian_num==True:
                if self.from_code in ["fa","ur"] :
                    self.text=self.replace_ascii_digits_with_farsi(self.text)
            if self.corrections==True:
                if self.from_code in ["fa","ur"] :
                  #  self.text = self.penglish_to_farsi(self.text)
                    self.text=self.Farsi_correction_words(self.text)
                else:
                    self.text=self.correction_text(self.text)
            if self.dict_aktive==True:
                if self.dict.currentText()==self.M125:
                    if self.from_code in ['fa']:
                       self.text=self.motaradef(self.text)
                    else:
                        self.q.information(self,self.M23,self.M157)
                elif self.dict.currentText()==self.M207:
                     self.text=self.search(self.text)
                else:
                    src_book_dic=self.searching_book_options()
                    if self.spm==True or src_book_dic==True :
                        self.text=self.search(self.text)
            return self.text
        except:
            self.info0('text process error')
            return self.text
    def translate_auto(self):
        if  self.thread_active==False:
            self.thread_active = True
            self.stops=False
            self.tranc_err=False
            self.translate_starter()
        else:
            self.update_notification(self.M24)

    def translate(self):
        if  self.thread_active==False:
            self.thread_active = True
            self.stops=False
            self.tranc_err=False
            self.argose_err = False
            self.translate_starter()
        else:
            self.update_notification(self.M24)
    def semi_learn(self):
        if  self.thread_active==False:
            self.thread_active = True
            self.stops=False
            self.tranc_err=False
            self.argose_err = False
            self.semi_learn_start()
        else:
            self.update_notification(self.M24)
    def semi_learn_start(self):
        try:
            self.open_5()
            t = self.learning(self.learn_lines)
            text_to_save = "\n".join(t)  # Join each line with a newline.
            with open(self.lo, 'w', encoding="utf-8") as f:  # Use 'w' to write (overwrite) instead of 'a'
                f.write(text_to_save)
            with open(self.lo, 'r', encoding='utf-8') as f:
                self.learn_lines=f.read().splitlines()
            self.process_finish()
        except:
            self.process_finish()
    def iventtext(self):
        self.timer.start(1500)
        self.translate_starter()
    def translate_starter(self):
        if self.source == self.M104 :
            self.detect_language()
        self.open_5()
        self.rev=False
        self.CWL='لیست کلمات جایگزین شده \n'
        try:
            self.output_console.setReadOnly(False)
            self.output_console.clear()
            input_text=self.input_console.toPlainText()
            text_parts = input_text.split('\n')
            text_groups = [[part] for part in text_parts]
            for group in text_groups:
                self.console=True
                group_text = '\n'.join(group)
                if self.stops==True:
                    break
                self.thread_active = True
                group_text=self.space_correction(group_text)
                self.translate_text(group_text)
              #  self.info_console.setPlainText('__')
            if self.corrections==True:
                self. SaveCorrections()
            
            self.process_finish()
            self.update_notification(self.M42)
        except Exception as e:
            self.process_finish()
            self.update_notification(f"   {self.M56}:{str(e)}")
    def translate_text(self, text: str) -> str:
        self.info_console.setPlainText('.')
        self.text = text
        try:
            # Preprocess the text.
            self.text = self.text_process()
            # Skip translation if error flag or no translation needed.
            if self.virast or self.from_code == self.to_code:
                self.update_notification(self.M140)
                final_text = self.text + '\n'
                self.update_output(final_text)
                return final_text
            # Validate text content.
            if len(self.text) < 2 or isinstance(self.text, float) or self.text.strip() == "":
                final_text = self.text + '\n'
                self.update_output(final_text)
                return final_text

            self.thread_active = True
            # Map menu items to translator functions.
            translators_map = {
                self.M119: self.google,
                self.M120: self.bing,
                self.M121: self.mymemory,
                self.M122: self.deep,
                self.M123: self.yandex,
                self.M201: self.translatorCom,
                self.M259: self.reverso,
                self.M260: self.googlev2,
                self.M261: self.googlev1,
                self.M118: self.argos,
            }
            # Set the language to display.
            selected_lang = (
                self.language_codes_2.get(self.from_code)
                if self.source == self.M104
                else self.source
            )
            self.update_notification(f"  {self.M57}  {selected_lang}  ...   {self.target}    ")
            # Retrieve the translator function.
            translator_name = self.translator_menu.currentText()
            translate_func = translators_map.get(translator_name)
            if not translate_func:
                self.update_notification("Translator function not found.")
                final_text = self.text + '\n'
                return final_text
            # Process and translate text paragraph by paragraph.
            paragraphs = self.text.split('\n')
            translated_paragraphs = []
            for paragraph in paragraphs:
                trans = paragraph.strip()
                # Set the current paragraph for translation.
                self.group_text = paragraph
                # If source language must be detected, do so.
                if self.source == self.M104:
                    self.detect_language()
                # For empty lines, append an empty string.
                if not trans:
                    translated_paragraphs.append("")
                    continue
                if  translator_name == self.M118 and self.learn == 1 and len(paragraph)<77:
                    # Check for an existing offline learning translation.
                    offline_translation = self._lookup_offline_translation(
                        trans, text
                    )
                    if offline_translation is not None:
                        translated_paragraphs.append(offline_translation)
                        continue
                # Get translation from the selected provider.
                translation = translate_func()
                # Save valid translation in offline learning cache (except for specific translator).
                if self.learn == 1 and translation and len(translation)<77 and translator_name != self.M118 and translation.strip()!=paragraph.strip():
                    self.learn_offline(paragraph, translation)
                translated_paragraphs.append(translation)
            translated_text = "\n".join(translated_paragraphs) or text
            final_text = translated_text + '\n'
            self.update_output(final_text)
            return final_text
        except Exception as e:
            error_message = str(e)
            self.info3(error_message)
            self.update_notification(f"  {self.M59} : {error_message}")
            self.process_finish()
            if self.translator_menu.currentText() == self.M118:
                self.update_notification(self.M60)
            else:
                self.packerror = True
                self.update_notification(self.M61)
            return self.text + '\n'

    @lru_cache(maxsize=100)
    def _lookup_offline_translation(self, trans: str, full_text: str) -> str:
        try:
            norm_para = trans.replace(" ", "")
            if norm_para in self.translation_cache:
                return self.translation_cache[norm_para]
            for line in self.learn_lines:
                parts = line.split("`")
                if len(parts) >= 2:
                    norm_trans = parts[1].replace(" ", "")
                    if norm_para == norm_trans:
                        result = parts[0].strip()
                        self.translation_cache[norm_para] = result
                        self.info0(f"{full_text} {result}")
                        return result
            return None  
        except Exception as e:
            self.info0(f"An error occurred: {e}")
            return None
    def learn_offline(self, word, replacedWord):
        normalized_word = word.strip()
        if normalized_word not in self.learn_texts:
            try:
                with open(self.lo, 'a+', encoding='utf-8') as f:
                    f.write('\n' + replacedWord + "`" + word)
                    f.flush()  # Ensure data is physically written out.
                    f.seek(0)
                    content = f.read()
                    self.learn_lines = content.splitlines()
                    self.learn_texts = re.split(r'[`\n]+', content)
            except Exception as e:
                self.update_notification(self.M254)
                self.info0(f"Error occurred in learn_offline: {e}")

    def google(self):
        try:
            self.translated_group = Translator1.Translator().translate(
                self.group_text,
                source_language=self.from_code,
                destination_language=self.to_code
            ).result
            return self.translated_group
        except Exception as e:
            self._handle_translation_error(e, self.M62)
            return self.group_text
    def bing(self):
        try:
            self.translated_group = Translator1.translators.BingTranslate().translate(
                self.group_text,
                source_language=self.from_code,
                destination_language=self.to_code
            ).result
            return self.translated_group
        except Exception as e:
            # Fallback to default translator if Bing fails.
            try:
                self.translated_group = Translator1.Translator().translate(
                    self.group_text,
                    source_language=self.from_code,
                    destination_language=self.to_code
                ).result
                return self.translated_group
            except Exception as fallback_e:
                self._handle_translation_error(fallback_e, "bing")
                return self.group_text
    def mymemory(self):
        language_codes = {
            "en": "en-US", "fa": "fa-IR", "de": "de-DE", "ar": "ar-SA",
            "fr": "fr-FR", "zh-CN": "zh-CN", "es": "es-ES", "ru": "ru-RU",
            "it": "it-IT", "tr": "tr-TR", "pt": "pt-PT", "id": "id-ID",
            "nl": "nl-NL", "hi": "hi-IN", "ja": "ja-JP", "ur": "ur-PK"
        }
        if self.from_code not in language_codes or self.to_code not in language_codes:
            self.info0(f"{self.M81} Unsupported language code")
            return self.group_text
        try:
            self.translated_group = Translator5.MyMemoryTranslator(
                source=language_codes[self.from_code],
                target=language_codes[self.to_code]
            ).translate(self.group_text)
            return self.translated_group
        except Exception as e:
            self._handle_translation_error(e, self.M81)
            return self.group_text
    def transe_learn(self,text):
        language_codes = {
            "en": "en-US", "fa": "fa-IR", "de": "de-DE", "ar": "ar-SA",
            "fr": "fr-FR", "zh-CN": "zh-CN", "es": "es-ES", "ru": "ru-RU",
            "it": "it-IT", "tr": "tr-TR", "pt": "pt-PT", "id": "id-ID",
            "nl": "nl-NL", "hi": "hi-IN", "ja": "ja-JP", "ur": "ur-PK"
        }
        try:
            translated= Translator5.MyMemoryTranslator(
                source=language_codes[self.from_code],
                target=language_codes[self.to_code]
            ).translate(text)
            return translated
        except:
            try:
                translated = Translator5.GoogleTranslator(
                    source=self.from_code,
                    target=self.to_code
                ).translate(text)
                return translated
            except :
                try:
                    translated= Translator1.translators.YandexTranslate().translate(
                        text,
                        source_language=self.from_code,
                        destination_language=self.to_code
                    ).result
                    return translated
                except Exception as e:
                    self._handle_translation_error(e, self.M83)
                    return text
    def deep(self):
        try:
            translated = Translator5.GoogleTranslator(
                source=self.from_code,
                target=self.to_code
            ).translate(self.group_text)
            return translated
        except Exception as e:
            self._handle_translation_error(e, self.M82)
            return self.group_text
    def yandex(self):
        try:
            self.translated_group = Translator1.translators.YandexTranslate().translate(
                self.group_text,
                source_language=self.from_code,
                destination_language=self.to_code
            ).result
            return self.translated_group
        except Exception as e:
            self._handle_translation_error(e, self.M83)
            return self.group_text
    def translatorCom(self):
        try:
            self.translated_group = Translator1.translators.translatecom.TranslateComTranslate().translate(
                self.group_text,
                source_language=self.from_code,
                destination_language=self.to_code
            ).result
            return self.translated_group
        except Exception as e:
            self._handle_translation_error(e, "translatorCom")
            return self.group_text
    def reverso(self):
        try:
            self.translated_group = Translator1.translators.reverso.ReversoTranslate().translate(
                self.group_text,
                source_language=self.from_code,
                destination_language=self.to_code
            ).result
            return self.translated_group
        except Exception as e:
            self._handle_translation_error(e, "ReversoTranslate")
            return self.group_text
    def googlev2(self):
        try:
            self.translated_group = Translator1.translators.google.GoogleTranslateV2().translate(
                self.group_text,
                source_language=self.from_code,
                destination_language=self.to_code
            ).result
            return self.translated_group
        except Exception as e:
            self._handle_translation_error(e, "googlev2")
            return self.group_text
    def googlev1(self):
        try:
            self.translated_group = Translator1.translators.google.GoogleTranslateV1().translate(
                self.group_text,
                source_language=self.from_code,
                destination_language=self.to_code
            ).result
            return self.translated_group
        except Exception as e:
            self._handle_translation_error(e, "googlev1")
            return self.group_text
    def argos(self):
        if self.to_code == 'zh-CN':
            self.to_code = 'zh'
        if self.from_code == 'zh-CN':
            self.from_code = 'zh'
        try:
            translated_text = Translator2.translate(self.group_text, self.from_code, self.to_code)
            if any(word in self.filter for word in re.split(r'(\W+)', translated_text)):
                try:
                    translated_text = self.transe_learn(self.group_text)
                    self.learn_offline(self.group_text, translated_text)
                except Exception:
                    translated_text = self.group_text
            fixed_text = self.fix_repeated_phrase(str(translated_text))
            return fixed_text
        except Exception as e:
            error_message = str(e)
            if "NoneType" in error_message or "opening" in error_message:
                if not self.argose_err:
                    self.process_finish()
                    self.check_and_install_argos_pak()
                else:
                    self.info0(f"{self.M60} : {error_message}")
            else:
                if self.argose_err:
                    self.argose_err = False
                    self.info0(f"{self.M59}: {error_message}")
                else:
                    self.info0(f"{self.M63} : {error_message}")
            self.process_finish()
            return self.group_text
    def _handle_translation_error(self, exception, service_name):
        error_message = f"{service_name} {self.M61}: {str(exception)}"
        if not self.tranc_err:
            self.tranc_err = True
            self.info0(error_message)
        else:
            self.info2(error_message)
        self.process_finish()
    def check_and_install_argos_pak(self):
        self.update_notification(self.M64)
        pattern = f"{self.from_code}_en" if self.from_code != "en" else f"{self.from_code}_{self.to_code}"
        pattern_2 = f"en_{self.to_code}" if self.to_code != "en" else f"{self.from_code}_{self.to_code}"
        os.makedirs(self.dir_path, exist_ok=True)
        os.makedirs(self.base_path, exist_ok=True)
        if not os.path.exists(self.destination_file_path):
            shutil.copy(self.source_path, self.destination_path)
        if self.find_folders(pattern) is not None and self.find_folders(pattern_2) is not None:
            self.update_notification(self.M65)
            if not self.reerror_packages or self.reinstalled:
                self.reerror_packages = True
                return
            else:
                self.reinstalled=True
                confirmation = self.qq(self,self.M65, self.M69,self.q.StandardButton.Yes | self.q.StandardButton.No, self.q.StandardButton.No)
                if confirmation == self.q.StandardButton.Yes:
                    if self.find_folders(pattern) is not None:
                        paths = self.find_folders(pattern)
                        for path in paths:
                                shutil.rmtree(path)
                    elif self.find_folders(pattern_2) is not None:
                         paths=self.find_folders(pattern_2)
                         for path_2 in paths:
                            shutil.rmtree(path_2)
                    return
        if self.from_code == "en" or self.to_code == "en":
            self.ls = self.from_code
            self.t = self.to_code
        else:
           if not self.find_folders(pattern)!=None:
                  self.ls=self.from_code
                  self.t="en"
           else:
                  self.ls="en"
                  self.t=self.to_code
        available_package = self.get_available_package()
        self.thread_active = True
        self.stops=False
        download_path=self.down_path()
        if available_package:
           self.download_url = available_package.links[0]
           try:
                if download_path:
                    if self.pack_install == True and self.pack_downloaded == True and self.argose_err == True :
                        self.update_notification(f"  {self.M67} {available_package} :{self.M68}")
                        self.pack_install = False
                        try:
                            os.remove(download_path)
                        except:
                            pass
                        self.download_lang_pack()
                        return
                    try:
                        if self.pack_install==False:
                            self.pack_install= True
                            self.extract_zip(download_path, self.base_path)
                            self.pack_install= False
                            self.pack_downloaded = False
                            self.q.information(self,self.M42,f"Installation  {available_package} {self.M89}")
                            self.restart_program()
                            return
                    except:
                        self.pack_install= False
                        self.download_lang_pack()
                        return
                else:
                    self.download_lang_pack()
                self.process_finish()
           except Exception as e:
                self.info0(f"{self.M68} {str(e)}")
                self.update_notification(f"  {self.M68}: {str(e)} {self.M70}...")
        else:
            self.update_notification(f"   { self.M71} {self.source}  {self.target} { self.M72} ")
    def extract_zip(self,download_path, base_path):
       
        with zip.ZipFile(download_path, 'r') as zip_ref:
             zip_ref.extractall(base_path)
    def get_available_package(self):
        available_packages =package.package.get_available_packages()
        return next((pkg for pkg in available_packages if pkg.from_code == self.ls and pkg.to_code == self.t), None)
    def download_lang_pack(self):
        if self.argose_err==True:
            self.info0(f"    {self.M73}   {fromdlang}   ...    {tomdlang}   {self.M74}  ")
            return
        parsed_url = urllib.urlparse(self.download_url)
        #self.pack_install=False
        if not all([parsed_url.scheme, parsed_url.netloc]):
            self.update_notification("Download Error: Please check internet connection and Retry")
        save_path= self.down_path()
        fromdlang=self.language_codes_2.get(self.ls)
        tomdlang=self.language_codes_2.get(self.t)
        confirmation = self.qq(self, self.tr(f"{self.M86}"),  self.tr(f"{self.M90}     {fromdlang}  ...  {tomdlang} "),
                                            self.q.StandardButton.Yes | self.q.StandardButton.No, self.q.StandardButton.No)
        if confirmation == self.q.StandardButton.No:
            self.argose_err = True
            self.process_finish()
            self.update_notification(f"    {self.M73}   {fromdlang}   ...    {tomdlang}   {self.M74}  ")
            if not self.dfiscancel==True:
               self.dfiscancel=True
            return
        self.thread_active = True
        self.stops=False
        self.pack_downloaded = True
        headers = {}
        downloaded = 0
        if os.path.exists(save_path):
            downloaded = os.path.getsize(save_path)
            headers['Range'] = f'bytes={downloaded}-'
        try:
            response = requests.get(self.download_url, headers=headers, stream=True)
            total_size = int(response.headers.get('Content-Length', 0)) + downloaded
            if total_size > 200000000:
                raise ValueError("File too large")
            hash_object = hashlib.sha256()  # Change to desired hash function if needed
            with open(save_path, 'ab') as file:
                for data in response.iter_content(chunk_size=1048576):
                    if self.stops:
                        break
                    file.write(data)
                    hash_object.update(data)
                    downloaded += len(data)
                    total_size_MB = int(total_size / 1048576)
                    percent = (downloaded) / 1048576
                    self.info3(f"{self.M75}{fromdlang} {tomdlang} \n {percent:.0f}/{total_size_MB} Mb")
            if downloaded >= total_size:
                self.check_and_install_argos_pak()
            self.process_finish()
        except requests.exceptions.RequestException as e:
            self.info0( f"Internet connection {str(e)}")
            self.pack_downloaded = True
            self.process_finish()
            self.update_notification(f"  Download Error: {str(e)}. Please Retry")
    def down_path(self):
        try:
            download_path = os.path.join(
                os.path.expanduser("~"),
                ".local",
                "cache",
                "argos-translate",
                "downloads",
                f"translate-{self.ls}_{self.t}.argosmodel"
            )
            return download_path
        except Exception as e:
            self.update_output(f"Down_path: {str(e)}")
    def find_folders(self, pattern):
        base_path = os.path.join(
            os.path.expanduser("~"),
            ".local",
            "share",
            "argos-translate",
            "packages"
        )
        folders = []
        for file_name in os.listdir(base_path):
            full_path = os.path.join(base_path, file_name)
            if os.path.isdir(full_path):
                if pattern in file_name:
                    folders.append(full_path)
        if not folders:
            return None
        return folders
    def find_matching_files(self, directory, pattern):
        try:
            matching_folders = []
            pattern_regex = re.compile(pattern)
            for folder_name in os.listdir(directory):
                if os.path.isdir(os.path.join(directory, folder_name)) and pattern_regex.search(folder_name):
                    matching_folders.append(folder_name)
            return matching_folders
        except Exception as e:
            self.update_notification(f"find_matching_folders Error: {str(e)} and directory: {directory}")
    def lang_code(self):
        try:
            self.source = self.source_language_combo.currentText()
            self.target = self.target_language_combo.currentText()
            self.from_code = self.language_codes.get(self.source)
            self.to_code = self.language_codes.get(self.target)
            t = self.source
            t2 = self.target
            t3 = self.translator_menu.currentText()
            if self.source == self.M104:
                self.detect_language()
                try:
                    self.save_state(2,t)
                    self.save_state(3,t2)
                    self.save_state(8,t3)
                except:
                    self.update_notification(self.M254)
                    try:
                        self.defultthem()
                    except:
                        self.update_notification(self.M254)
                return
            try:
                self.save_state(2,t)
                self.save_state(3,t2)
                self.save_state(8,t3)
            except:
                try:
                    self.defultthem()
                except:
                    self.update_notification(self.M254)
            return
        except Exception as e:
            self.info3(str(e))
            self.process_finish()

    def detect_language(self):
        self.source = self.source_language_combo.currentText()
        self.target = self.target_language_combo.currentText()
        self.from_code = self.last_detected_languages
        if not len(self.text)<2:
            try:
                if self.translator_menu.currentText() in [self.M119, self.M120]:
                    self.from_code = 'auto'
                elif re.search(r'[\u0600-\u06FF]', self.group_text):
                    if self.last_detected_languages != 'ar':
                        self.last_detected_languages = 'fa'
                    else:
                        self.last_detected_languages = 'ar'
                    self.from_code = self.last_detected_languages
                else:
                    detected_language = langdetec.detect(self.group_text)  # use langdetect for language detection
                    if len(detected_language)<3 and not detected_language in self.invalid_languages and   detected_language!=None :
                        self.last_detected_languages = detected_language
                        self.from_code = self.last_detected_languages
                self.to_code = self.language_codes.get(self.target)
                self.info0(f"{self.source}   ...   {self.target}    {self.translator_menu.currentText()} ")
            except :
                self.last_detected_languages = 'en'
        else:self.last_detected_languages = 'en'
    def clear(self):
        self.input_console.clear()
        self.output_console.clear()
    def correction_text(self, text):
        try:
            if  self.from_code in ['en','fr','pt','de','it','ar','eu','nl','ar']:
                text=self.spellingcheck_de(text)
                return text
            if self.from_code in ['ur','zh-CN','id','ru',]:
                selected_lang=self.language_codes_2.get(self.from_code)
                self.info0(f"corrections not support for {selected_lang}")
                return text
            self.update_notification(f"{self.M84}:{text}")
            try:
                spell = Speller.Speller(self.from_code)
                txt=spell(text)
            except:
                txt=text
            return txt
        except Exception as e:
            self.info3(str(e))
            return text         
    def spellingcheck_de(self, text):
        words=re.split(r'(\W+)', text)
        corrected_words = []
        if self.from_code == "en":
            spell =SpellChecke.SpellChecker()
        else:
            spell =SpellChecke.SpellChecker(language=self.from_code)
        for word in words:
            corr_word=""
            if  len(word)<3 or re.search(r'\W', word) and not word.isspace() or isinstance(word, int) or word in self.special_chars:
                corrected_words.append(word)
                continue
            if  re.search(r'[\u0600-\u06FF]', word) and not self.from_code=='ar':
                corrected_words.append(word)
                continue
            try:
                corr_word = spell.correction(word)
                if str(corr_word) !=' ' or str(corr_word) != None:
                    corrected_word=corr_word
                else:corrected_word=word
            except Exception as e:
                self.update_notification(f"  {self.M23}  : {e}")
                corrected_word=word
            if corrected_word is None:
                corrected_word=word
            corrected_words.append(corrected_word)
        correct_text=''.join(corrected_words)
        return correct_text
    def replace_characters2(self, words):
        try:
            for i in range(len(words)): 
                if self.stopBook==True or self.stops==True:    
                    return words
                words[i] = words[i].replace('هـ', 'ه')
                if words[i] == "اهلل":
                    words[i]='الله'
                words[i] = words[i].replace("الل ه", "الله")
                words[i] = words[i].replace("اهلل", "الله")
        except Exception :
            words[i] = ""
        return words 
    def Farsi_correction_words(self, text):
        if not self.S6:
            self.open_4()
        self.not_save_active = False
        try:
            if not self.corrections:
                return text
            words1 = re.split(r'(\W+)', text)
            words1 = self.replace_characters2(words1)
            words = self.Farsi_check_singel_char(words1)
            corrected_words = []
            cache = {}
            fa_words_set = self.fa_words if isinstance(self.fa_words, set) else set(self.fa_words)
            fa_not_found_set = self.fa_w_not_found if isinstance(self.fa_w_not_found, set) else set(self.fa_w_not_found)
            for word in words:
                if word in cache:
                    corrected_words.append(cache[word])
                    continue
                self.gonext = False
                self.update_notification("  " + word)
                if not self.corrections:
                    cache[word] = word
                    corrected_words.append(word)
                    continue
                if (len(word) <= 3 or 
                    (re.search(r'\W', word) and not word.isspace()) or 
                    isinstance(word, int) or 
                    not re.search(r'[\u0600-\u06FF]', word) or 
                    any(re.search(pattern, word) for pattern in self.patterns) or 
                    word in self.special_chars):
                    cache[word] = word
                    corrected_words.append(word)
                    continue
                if word in fa_words_set:
                    cache[word] = word
                    corrected_words.append(word)
                    continue
                if self.semi_corrections == False and word in self.replaced_words:
                    rep_line_ok = False
                    for line in self.replaced_lines:
                        if word in line:
                            for w in line.split("`"):
                                if w != word and not rep_line_ok:
                                    rep_line_ok = True
                                    cache[word] = w
                                    corrected_words.append(w)
                                    self.info2(f"{word} \u2190 {w}")
                                    if word not in self.CWL:
                                        self.CWL += f"{word} \u2190 {w}\n"
                                    break
                    if rep_line_ok:
                        continue
                if word in self.replaced_words2:
                    rep_line_ok = False
                    for line in self.replaced_lines2:
                        if word in line:
                            for w in line.split("`"):
                                if w != word and not rep_line_ok:
                                    rep_line_ok = True
                                    cache[word] = w
                                    corrected_words.append(w)
                                    self.info2(f"{word} \u2190 {w}")
                                    if word not in self.CWL:
                                        self.CWL += f"{word} \u2190 {w}\n"
                                    break
                    if rep_line_ok:
                        continue
                if len(word) >= 3:
                    for suffix in self.suffixes:
                        if word.endswith(suffix):
                            new_word = word[:-len(suffix)]
                            if new_word in fa_words_set:
                                if len(word) > 4 and len(suffix) > 2:
                                    self.gonext = True
                                    similar_word = new_word + '\u200c' + suffix + ' '
                                    cache[word] = similar_word
                                    corrected_words.append(similar_word)
                                    self.saveReplacedWords(word, similar_word)
                                    break
                                elif self.semi_corrections == False and new_word in fa_not_found_set:
                                    self.gonext = True
                                    cache[word] = word
                                    corrected_words.append(word)
                                    if word not in fa_not_found_set:
                                        self.saveNewWord(word)
                                    break
                            else:
                                continue
                    if self.gonext:
                        continue
                    for prefix in self.start_with:
                        if word.startswith(prefix):
                            new_word = word[len(prefix):]
                            if new_word in fa_words_set:
                                if prefix == 'و':
                                    self.gonext = True
                                    word1 = prefix + ' ' + new_word
                                    self.saveReplacedWords(word, word1)
                                    cache[word] = word1
                                    corrected_words.append(word1)
                                elif self.semi_corrections == False and new_word in fa_not_found_set:
                                    self.gonext = True
                                    cache[word] = word
                                    corrected_words.append(word)
                                    if word not in fa_not_found_set:
                                        self.saveNewWord(word)
                                break
                        else:
                            continue
                    if self.gonext:
                        continue
                if self.semi_corrections == False and word in fa_not_found_set:
                    cache[word] = word
                    corrected_words.append(word)
                    continue
                if self.semi_corrections:
                    if word in self.replaced_words2:
                        cache[word] = word
                        corrected_words.append(word)
                        continue
                    else:
                        word1 = self.farsi_semi_auto_correction(word, text)
                        cache[word] = word1
                        corrected_words.append(word1)
                        self.saveReplacedWords(word, word1)
                        continue
                else:
                    self.update_notification(f"  {self.M93} : {word}")
                    similar_word = self.farsi_auto_correction(word)
                    if similar_word is not None:
                        cache[word] = similar_word
                        corrected_words.append(similar_word)
                        if word != similar_word:
                            self.info2(f"{word} \u2190 {similar_word}")
                            if word not in self.CWL:
                                self.CWL += f"{word} \u2190 {similar_word}\n"
                            self.saveReplacedWords(word, similar_word)
                        elif not re.search(r'\s', similar_word) and similar_word not in fa_not_found_set:
                            self.saveNewWord(similar_word)
                        continue
                    else:
                        self.info2(f"first  {word}")
                        word2 = self.replace_characters1(word, False)
                        self.info2(f"second  {word2}")
                        if word in fa_not_found_set or isinstance(word2, int) or word2 in fa_words_set:
                            cache[word] = word2
                            corrected_words.append(word2)
                            continue
                        similar_word = self.farsi_auto_correction(word2)
                        if similar_word is not None:
                            cache[word] = similar_word
                            corrected_words.append(similar_word)
                            if word != similar_word:
                                self.info2(f"{word} \u2190 ( {similar_word} )")
                                if word not in self.CWL:
                                    self.CWL += f"{word} \u2190 ( {similar_word}\n )"
                                self.saveReplacedWords(word, similar_word)
                            elif not re.search(r'\s', similar_word) and similar_word not in fa_not_found_set:
                                self.saveNewWord(similar_word)
                            continue
                        else:
                            cache[word] = word
                            corrected_words.append(word)
                            if word not in fa_not_found_set:
                                self.saveNewWord(word)
                            continue

            return ''.join(corrected_words)

        except Exception as e:
            return text

    def farsi_auto_correction(self, word):
        if not self.S6:
            self.open_4()
        if len(self.words2) < 2 :
            self.words2 = farsi_tool.stop_words()
        def replace_and_check(i, char):
            new_word = word[:i] + char + word[i+1:]
            if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words:
                return new_word
            return None
        try:
            similar_word = self.find_base_word(word)
            if similar_word is not None:
                return similar_word
            similar_word = self.find_similar_word(word)
            if similar_word is not None:
                return similar_word            
            for i in range(len(word)):
                for char_group in self.char_groups:
                    if word[i] in char_group:
                        for char in char_group:
                            new_word = word[:i] + char + word[i+1:]
                            if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words:
                                return new_word
            for i in range(len(word)):
                for char in self.chars:
                    new_word = replace_and_check(i, char)
                    if new_word is not None:
                        return new_word
            return None
        except:
            return None
    def find_base_word(self, word):
        self.info0('ریشه یابی')
        try:
            # Define constants for readability
            MIN_LENGTH_SHALLOW = 9
            MEDIUM_LENGTH_MAX = 12
            TRUNCATE_SHORT_MAX = 7
            TRUNCATE_LONG_MIN = 5
            TRUNCATE_LONG_MAX = 8
            
            # Helper function to check word validity
            def is_valid(w):
                return w in self.words2 or w in self.fa_w_not_found or w in self.fa_words
            
            # Helper function to check ZWNJ split words
            def check_split_words(w):
                split_words = w.split('\u200c')
                return any(is_valid(sw) for sw in split_words)
            
            # Helper function to check prefixes
            def check_prefixes(w):
                for prefix in self.start_with:
                    if w.startswith(prefix):
                        truncated = w[len(prefix):]
                        if is_valid(truncated):
                            self.save_word = True
                            return True
                return False
            
            # Character replacement helper
            def replace_chars(w):
                for i in range(len(w)):
                    for char_group in self.char_groups:
                        if w[i] in char_group:
                            for char in char_group:
                                new_word = w[:i] + char + w[i+1:]
                                if is_valid(new_word):
                                    self.save_word = True
                                    return new_word
                return None
            
            # Character insertion helper
            def insert_chars(w):
                for i in range(len(w) + 1):
                    for char in self.chars:
                        new_word = w[:i] + char + w[i:]
                        if is_valid(new_word):
                            self.save_word = True
                            return new_word
                return None
            
            # Truncation helper
            def truncate_word(w, min_truncate, max_truncate):
                for length in range(min_truncate, max_truncate + 1):
                    if length >= len(w):
                        continue
                    truncated = w[:-length]
                    if is_valid(truncated):
                        if len(truncated) > 3 and length > 1:
                            return f"{truncated}{w[-length:]}"
                        self.save_word = True
                        return w
                return None
            
            # Main logic flow
            if check_split_words(word):
                return word
            
            if len(word) < MIN_LENGTH_SHALLOW or self.dc2:
                if check_prefixes(word):
                    return word
                
                replaced = replace_chars(word)
                if replaced:
                    return replaced
                
                inserted = insert_chars(word)
                if inserted:
                    return inserted
            
            if MEDIUM_LENGTH_MAX > len(word) > 4:
                truncated = truncate_word(word, 1, TRUNCATE_SHORT_MAX)
                if truncated:
                    return truncated
            
            if len(word) > 11 or self.dc2:
                truncated = truncate_word(word, TRUNCATE_LONG_MIN, TRUNCATE_LONG_MAX)
                if truncated:
                    return truncated
            
            self.info0('اصلاح')
            return None
        
        except Exception as e:
            return None
    def find_similar_word(self, word):
        if len(word) < 12 or  self.dc2==True:
            try:
                cleaned_word = self.clean_text_arabic(word)
                if cleaned_word in self.fa_words:
                    return cleaned_word
                wordq=self.find_base_word(cleaned_word)
                if wordq is not None:
                    return wordq
                for x in self.words2:
                    ratio = Sequence.SequenceMatcher(None, cleaned_word, x).ratio()
                    if ratio > 0.75:
                       return x
                for x in self.fa_words:
                    ratio = Sequence.SequenceMatcher(None, cleaned_word, x).ratio()
                    if ratio > 0.75:
                       return x
            except Exception as e:
                return None
        return None
    def fix_repeated_phrase(self, text):
        words = re.split(r'(\W+)', text)  # Split text into words and non-word characters
        n = len(words)
        
        # Start with larger sequences (phrases) so that longer repeated blocks are removed first
        for seq_len in range(n // 2, 0, -1):
            i = 0
            # Ensure there is room for at least two sequences of length seq_len
            while i <= len(words) - 2 * seq_len:
                # The sequence we're going to compare, excluding spaces
                seq = [word for word in words[i:i + seq_len] if word.strip()]
                count = 1
                # Count how many times this exact sequence is repeated consecutively
                while (i + (count + 1) * seq_len) <= len(words) and \
                    [word for word in words[i + count * seq_len:i + (count + 1) * seq_len] if word.strip()] == seq:
                    count += 1

                if count > 1:
                    # We have one or more extra occurrences: keep only one occurrence
                    words = words[:i + seq_len] + words[i + count * seq_len:]
                    # Do not increment i so that any new adjacent sequence (after merging) is checked
                else:
                    i += 1
        
        # Join the tokens back into a string and return
        return "".join(words)

    def saveNewWord(self, newword):
        try:
            if not self.S6:
                self.open_4()
            with open(self.not_found_file, 'a', encoding="utf-8") as f:
                f.write('\n')
                f.write(newword)
            with open(self.not_found_file, 'r', encoding="utf-8") as f:
                self.fa_w_not_found = f.read().splitlines()
        except:
            self.update_notification(self.M254)

    def saveReplacedWords(self, word, replacedWord):
        if not word in self.replaced_words:     
            try:
                if not self.S6:
                    self.open_4()
                self.replace = 'replace.json'
                with open(self.replace, 'a', encoding="utf-8") as f:
                    f.write('\n')
                    f.write(replacedWord + "`" + word)
                with open(self.replace, 'r', encoding='utf-8') as f:
                    self.replaced_lines = f.read().splitlines()
                with open(self.replace, 'r', encoding='utf-8') as f:
                    self.replaced_words = set(re.split(r'[`\n]+', f.read()))
            except:
                self.update_notification(self.M254)
        elif not word in self.replaced_words2:
            try:
                if not self.S6:
                    self.open_4()
                self.replace2 = 'replace2.json'
                with open(self.replace2, 'a', encoding="utf-8") as f:
                    f.write('\n')
                    f.write(replacedWord + "`" + word)
                with open(self.replace2, 'r', encoding='utf-8') as f:
                    self.replaced_lines2 = f.read().splitlines()
                with open(self.replace2, 'r', encoding='utf-8') as f:
                    self.replaced_words2 = set(re.split(r'[`\n]+', f.read()))
            except:
                self.update_notification(self.M254)
    def check_single_char(self, words, i):
        try:
            if i > 0 and (words[i-1] + words[i]) in self.fa_words:
                return 'before'
            elif i < len(words) - 1 and (words[i] + words[i+1]) in self.fa_words:
                return 'after'
            else:
                return 'none'
        except:
             return 'none'
    def Farsi_check_singel_char(self, words):
        try:
            corrected_words = []
            i = 0
            while i < len(words):
                word = words[i]
                corrected_word = word  # Initialize corrected_word with the original word
                if len(word) == 1 and word not in['و','ه','ء'] :
                    check_result = self.check_single_char(words, i)
                    if check_result == 'before':
                        corrected_word = corrected_words[-1] + word  # Join with the word before
                        corrected_words[-1] = corrected_word  # Update the last word in corrected_words
                        words[i]=''
                    elif check_result == 'after':
                        corrected_word = word + words[i+1]  # Join with the word after
                        corrected_words.append(corrected_word)  # Add the corrected word to corrected_words
                        words[i+1]=''
                    else:
                        corrected_words.append(word)  # Do not join
                    i += 1
                else:
                    corrected_words.append(word)
                    i += 1
            return corrected_words
        except:
                return words

    def fix_space_word(self, text):
        def replacer(match):
            char = match.group(0)
            if all(c in self.special_chars for c in char):
                return char  # Return sequence of special characters as is
            elif any(c in self.special_chars for c in char):
                # Split the sequence of characters to handle individual special characters
                return ''.join(f' {c} ' if c in self.special_chars else c for c in char)
            return char
        try:
            # Use a regular expression to match sequences of special characters and words
            pattern = f"({'|'.join(re.escape(char) for char in self.special_chars)}|\\S+)"
            fixed_text = re.sub(pattern, replacer, text)
        except Exception as e:
            print(f"An error occurred: {e}")

        return fixed_text

    def space_correction(self,text):
        if  self.space_word_==0:
            return text
        try:
            text=self.fix_space_word(text)
            if not self.to_code in ['fa','ur','ar']:
                 text = ftfy.fix_text(text)
                 text = re.sub(r'([^\w\s])', r' \1' , text)
           # text = re.sub(r'(\d+)', r' \1 ', text)
           # text = re.sub(r'(\d+\))', r' \1 ', text)
            text = re.sub(r'(\b[A-Za-z]+@[A-Za-z]+\.[A-Z|a-z]{2,}\b)', r' \1 ', text)
            text = re.sub(r'("  ")', r' \1 ', text)
            return text
        except:
            return text
    def learning2(self):
        if not self.S6:
            self.open_4()
        if self.semi_corrections ==True:
           replace2 = 'replace2.json' 
        else:
           replace2 = 'replace.json'
        try:
            # Read the file and load lines into a set.
            with open(replace2, 'r', encoding='utf-8') as f:
                text_list =f.read().splitlines()
            size = int(self.size_box.value())
            font_name = self.font_box.currentText()
            font = QFont(font_name)
            font.setPointSize(size)
            # Create the main dialog.
            dialog = QDialog(self.tab_widget)
            dialog.setWindowTitle(self.M133)
            dialog.setGeometry(66, 66, 1000, 600)
            layout = QVBoxLayout(dialog)

            find_label = QLabel("Find")
            find_edit = QLineEdit()
            replace_label = QLabel("Replace:")
            replace_edit = QLineEdit()
            layout.addWidget(find_label)
            layout.addWidget(find_edit)
            layout.addWidget(replace_label)
            layout.addWidget(replace_edit)
            table = QTableWidget(dialog)
            table.setRowCount(len(text_list))
            table.setColumnCount(4)
            table.setFont(font)
            table.setHorizontalHeaderLabels([
                "کلمات نامشخص",
                "کلمات جایگزین ",
                "Translate",
                "Delete"
            ])
            def add_buttons_for_row(row):
                # --- Translate Button in Column 2 ---
                translate_button = QPushButton("Aouto Correct", dialog)
                def on_translate_clicked():
                    btn = dialog.sender()  # Identify the clicked button.
                    target_row = -1
                    # Look up the row by comparing the sender with the cell widget.
                    for r in range(table.rowCount()):
                        if table.cellWidget(r, 2) is btn:
                            target_row = r
                            break
                    if target_row == -1:
                        return
                    # Get the original text from Column 0.
                    original_item = table.item(target_row, 0)
                    original_text = original_item.text() if original_item else ""
                    translated_full = self.find_similar_word(original_text)
                    if table.item(target_row, 1) and translated_full is not None:
                        table.item(target_row, 1).setText(translated_full)
                translate_button.clicked.connect(on_translate_clicked)
                table.setCellWidget(row, 2, translate_button)
                # --- Delete Button in Column 3 ---
                delete_button = QPushButton("Delete", dialog)
                def on_delete_clicked():
                    # Identify the button that triggered the event
                    btn = dialog.sender()
                    target_row = -1

                    # Locate the row with the button in column 3
                    for row in range(table.rowCount()):
                        if table.cellWidget(row, 3) is btn:
                            target_row = row
                            break

                    # Remove the row if a match is found
                    if target_row != -1:
                        table.removeRow(target_row)

                # Connect the 'delete_button' clicked signal to the function
                delete_button.clicked.connect(on_delete_clicked)
                table.setCellWidget(row, 3, delete_button)
            for row, line in enumerate(text_list):
                parts = line.split("`", 1)
                original_text = parts[1] if len(parts) > 0 else ""
                replacement_text = parts[0] if len(parts) > 1 else ""
                table.setItem(row, 0, QTableWidgetItem(original_text))
                table.setItem(row, 1, QTableWidgetItem(replacement_text))
                add_buttons_for_row(row)
            table.resizeColumnsToContents()
            table.resizeRowsToContents()
            table.horizontalHeader().setStretchLastSection(True)

            # Scroll to the bottom after populating the table
            table.scrollToBottom()
            layout.addWidget(table)

            table.resizeColumnsToContents()
            table.resizeRowsToContents()
            table.horizontalHeader().setStretchLastSection(True)
            layout.addWidget(table)
            # Add a single "Insert New Row" button below the table.
            insert_button = QPushButton("Insert New Row", dialog)
            layout.addWidget(insert_button)
            # Setup the Find button and search field connections.
            # Initialize search state variables (they will persist between calls).
            last_search_text = ""
            match_positions = []  # List of tuples: (row, col)
            current_match_index = 0

            def find_text_in_table():
                nonlocal last_search_text, match_positions, current_match_index
                # Get and trim the search string.
                search_text = find_edit.text().strip()
                
                # If the search text is empty, exit.
                if not search_text:
                    return

                # If this is a new search string, recalc all match positions.
                if search_text != last_search_text:
                    last_search_text = search_text
                    match_positions = []
                    # Look in both columns (columns 1 and 2).
                    for row in range(table.rowCount()):
                        for col in range(1, 3):  # columns 1 and 2
                            item = table.item(row, col)
                            if item and search_text in item.text():
                                match_positions.append((row, col))
                    current_match_index = 0  # Restart from the first match.

                # If no matches were found, optionally show feedback (here we just return).
                if not match_positions:
                    return

                # If we've cycled past the last match, loop back to the first.
                if current_match_index >= len(match_positions):
                    current_match_index = 0

                # Get the match position and scroll to it.
                row, col = match_positions[current_match_index]
                item = table.item(row, col)
                if item:
                    table.scrollToItem(item, QAbstractItemView.ScrollHint.PositionAtCenter)
                    table.setCurrentItem(item)
                current_match_index += 1
            fund_button = QPushButton("Find next", dialog)
            layout.addWidget(fund_button)
            # Connect both the button click and the Enter key press to trigger finding next.
            fund_button.clicked.connect(find_text_in_table)
            find_edit.textChanged.connect(find_text_in_table)

            # Add "Replace" button.
            replace_button = QPushButton("Replace", dialog)
            layout.addWidget(replace_button)
            def replace_text_in_table():
                search_text = find_edit.text()
                new_text_val = replace_edit.text()
                # Replace occurs only in the text columns (0 and 1).
                for row in range(table.rowCount()):
                    for col in range(2):
                        item = table.item(row, col)
                        if item and search_text in item.text():
                            updated = item.text().replace(search_text, new_text_val)
                            table.setItem(row, col, QTableWidgetItem(updated))
            replace_button.clicked.connect(replace_text_in_table)
            buttonBox = QDialogButtonBox(
                QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel,
                parent=dialog
            )
            layout.addWidget(buttonBox)
            buttonBox.accepted.connect(dialog.accept)
            buttonBox.rejected.connect(dialog.reject)
            result = dialog.exec()
            new_text_list = []
            if result == QDialog.DialogCode.Accepted:
                for row in range(table.rowCount()):
                    orig_item = table.item(row,1)
                    repl_item = table.item(row, 0)
                    orig_text = orig_item.text() if orig_item else ""
                    repl_text = repl_item.text() if repl_item else ""
                    new_text_list.append(f"{orig_text}`{repl_text}")
            else:
                new_text_list = text_list
            text_to_save = "\n".join(new_text_list)

            with open(replace2, 'w', encoding="utf-8") as f:
                f.write(text_to_save)
            self.open_4()
            return new_text_list

        except Exception as e:
            self.info3(str(e))
            return text_list

    def learning(self, text_collection):
        try:
            text_list = text_collection
            dialog = QDialog(self.tab_widget)
            dialog.setWindowTitle(f"لیست ترجمه {self.source} به {self.target} . {self.M133}")
            dialog.setGeometry(66, 66, 1000, 600)
            layout = QVBoxLayout(dialog)
            toggle_button = QPushButton("Full Screen", dialog)
            layout.addWidget(toggle_button)

            def toggle_fullscreen():
                    # Check if the dialog is currently in full-screen mode.
                    if dialog.windowState() & Qt.WindowState.WindowFullScreen:
                        # Currently full screen; revert to normal mode.
                        dialog.showNormal()
                        toggle_button.setText("Full Screen")
                    else:
                        # Not full screen; switch to full screen.
                        dialog.showFullScreen()
                        toggle_button.setText("Normal Screen")
                
                
                    return dialog
            toggle_button.clicked.connect(toggle_fullscreen)

            # Add Find/Replace controls.
            find_label = QLabel("Find:")
            find_edit = QLineEdit()
            replace_label = QLabel("Replace:")
            replace_edit = QLineEdit()
            layout.addWidget(find_label)
            layout.addWidget(find_edit)
            layout.addWidget(replace_label)
            layout.addWidget(replace_edit)

            table = QTableWidget(dialog)
            table.setRowCount(len(text_list))
            table.setColumnCount(4)
            table.setHorizontalHeaderLabels([
                "Translate",
                "Original text in " + self.source,
                "Translated text in " + self.target,
                "Delete"
            ])

            # Helper function to add Translate and Delete buttons for a given row.
            def add_buttons_for_row(row):
                # --- Translate Button ---
                translate_button = QPushButton("Translate", dialog)
                def on_translate_clicked():
                    btn = dialog.sender()  # Get the button that was clicked.
                    target_row = -1
                    # Find the row that contains this translate button.
                    for r in range(table.rowCount()):
                        if table.cellWidget(r, 0) is btn:
                            target_row = r
                            break
                    if target_row == -1:
                        return
                    # Retrieve the original text from column 1.
                    original_item = table.item(target_row, 1)
                    original_text = original_item.text() if original_item else ""
                    # Call your translation method.
                    translated_full = self.transe_learn(original_text)
                    # Always set the translated text in column 2.
                    if table.item(target_row, 2):
                        table.item(target_row, 2).setText(translated_full)
                    else:
                        table.setItem(target_row, 2, QTableWidgetItem(translated_full))
                translate_button.clicked.connect(on_translate_clicked)
                table.setCellWidget(row, 0, translate_button)
              #  table.verticalScrollBar().setValue(table.verticalScrollBar().maximum())

                # --- Delete Button ---
                delete_button = QPushButton("Delete", dialog)
                def on_delete_clicked():
                    btn = dialog.sender()  # Identify the button that was clicked.
                    target_row = -1
                    for r in range(table.rowCount()):
                        if table.cellWidget(r, 3) is btn:
                            target_row = r
                            break
                    if target_row != -1:
                        table.removeRow(target_row)
                delete_button.clicked.connect(on_delete_clicked)
                table.setCellWidget(row, 3, delete_button)

            # Populate the table with rows from text_list.
            for row, line in enumerate(text_list):
                parts = line.split("`", 1)  # Expect exactly one backtick.
                if len(parts) == 2:
                    original_text, translated_text = parts[0], parts[1]
                    table.setItem(row, 2, QTableWidgetItem(original_text))
                    table.setItem(row, 1, QTableWidgetItem(translated_text))
                    add_buttons_for_row(row)

            table.setWordWrap(True)
            table.resizeColumnsToContents()
            table.resizeRowsToContents()
            table.horizontalHeader().setStretchLastSection(True)
            layout.addWidget(table)

            # Automatically scroll the table to the last row if any row exists.
            if table.rowCount() > 0:
                # Option 1: Scroll to the item in the first column of the last row.
                last_item = table.item(table.rowCount()-1, 0)
                if last_item:
                    table.scrollToItem(last_item, QTableWidget.PositionAtBottom)
                # Option 2: or using the scrollbar directly:
            table.verticalScrollBar().setValue(table.verticalScrollBar().maximum())

            # Add Find and Replace buttons.
            find_button = QPushButton("Find next", dialog)
            layout.addWidget(find_button)
            last_search_text = ""
            match_positions = []  # List of tuples: (row, col)
            current_match_index = 0

            def find_text_in_table():
                nonlocal last_search_text, match_positions, current_match_index
                # Get and trim the search string.
                search_text = find_edit.text().strip()
                
                # If the search text is empty, exit.
                if not search_text:
                    return

                # If this is a new search string, recalc all match positions.
                if search_text != last_search_text:
                    last_search_text = search_text
                    match_positions = []
                    # Look in both columns (columns 1 and 2).
                    for row in range(table.rowCount()):
                        for col in range(1, 3):  # columns 1 and 2
                            item = table.item(row, col)
                            if item and search_text in item.text():
                                match_positions.append((row, col))
                    current_match_index = 0  # Restart from the first match.

                # If no matches were found, optionally show feedback (here we just return).
                if not match_positions:
                    return

                # If we've cycled past the last match, loop back to the first.
                if current_match_index >= len(match_positions):
                    current_match_index = 0

                # Get the match position and scroll to it.
                row, col = match_positions[current_match_index]
                item = table.item(row, col)
                if item:
                    table.scrollToItem(item, QAbstractItemView.ScrollHint.PositionAtCenter)
                    table.setCurrentItem(item)
                current_match_index += 1

            # Connect both the button click and the Enter key press to trigger finding next.
            find_button.clicked.connect(find_text_in_table)
            find_edit.textChanged.connect(find_text_in_table)

            replace_button = QPushButton("Replace", dialog)
            layout.addWidget(replace_button)
            def replace_text_in_table():
                search_text = find_edit.text()
                new_text = replace_edit.text()
                # Operate on text columns (columns 1 and 2).
                for row in range(table.rowCount()):
                    for col in range(1, 3):
                        item = table.item(row, col)
                        if item and search_text in item.text():
                            updated = item.text().replace(search_text, new_text)
                            table.setItem(row, col, QTableWidgetItem(updated))
            replace_button.clicked.connect(replace_text_in_table)

            # Add OK and Cancel buttons.
            buttonBox = QDialogButtonBox(
                QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel,
                parent=dialog
            )
            layout.addWidget(buttonBox)
            buttonBox.accepted.connect(dialog.accept)
            buttonBox.rejected.connect(dialog.reject)

            # Execute the dialog modally.
            result = dialog.exec()
            new_text_list = []
            if result == QDialog.DialogCode.Accepted:
                # Rebuild the text list in the format: original_text`translated_text.
                for row in range(table.rowCount()):
                    original_item = table.item(row, 2)
                    translated_item = table.item(row, 1)
                    original_text = original_item.text() if original_item else ""
                    translated_text = translated_item.text() if translated_item else ""
                    new_text_list.append(f"{original_text}`{translated_text}")
            else:
                new_text_list = text_list

            return new_text_list

        except Exception as e:
            self.info3(str(e))
            return text_collection

    def farsi_semi_auto_correction(self, word, text):
        try:
            size = int(self.size_box.value())
            fo = self.font_box.currentText()
            font = QFont(fo)
            font.setPointSize(size)
            synonyms_str = Sequence.get_close_matches(word, self.fa_words, n=19, cutoff=0.7)
            win = QDialog(self.tab_widget)
            win.setGeometry(66, 66, 769, 600)
            win.setWindowTitle(self.M133)
            label = QTextEdit(win)
            label.setFont(font)
            label.setText(text)
            self.highlight_words_curser(label, word)
            label.move(100, 20)
            label.resize(600, 110)
            self.Farhang = QTextEdit(win)
            self.Farhang.setFont(font)
            self.Farhang.setReadOnly(True)
            self.Farhang.move(100, 250)
            self.Farhang.resize(600, 110)
            self.button_Farhang = QPushButton(win)
            self.words_list = QPushButton(win)
            entry = QLineEdit(win)
            entry.setFont(font)
            entry.setText(word)
            label.setReadOnly(True)
            entry.move(100, 166)
            entry.resize(600, 35)
            label2 = QLabel(win)
            label2.setText(self.M131)
            label2.move(10, 135)
            label2.resize(600, 35)
            combo = QComboBox(win)
            combo.insertItem(0, self.M135)
            combo.addItems(synonyms_str)
            combo.setCurrentIndex(0)
            combo.move(100, 365)
            combo.resize(600, 35)
            hand_word = None

            def on_combobox_changed(text):
                nonlocal hand_word
                hand_word = text
                entry.setText(hand_word)

            combo.currentTextChanged.connect(on_combobox_changed)
            button = QPushButton(win)
            button.setText(self.M132)
            button.move(100, 500)
            button.resize(600, 55)
            button.clicked.connect(win.accept)

            def on_checkbox_state_changed(state):
                self.semi_correct.setCheckState(Qt.CheckState.Unchecked)
                self.semi_corrections = False
            checkbox = QCheckBox(self.M136, win)
            checkbox.stateChanged.connect(on_checkbox_state_changed)
            checkbox.move(100, 460)
            checkbox.resize(600, 35)
            checkbox_2 = QCheckBox(self.M142, win)
            checkbox_2.move(313, 460)
            checkbox_2.resize(600, 35)
            self.button_Farhang.move(100, 205)
            self.button_Farhang.resize(600, 35)
            self.words_list.resize(600, 35)
            self.words_list.move(100, 420)
            self.hand_word = word
            self.button_Farhang.clicked.connect(self.searching_farhang)
            self.button_Farhang.setText(self.M246)
            self.words_list.setText(self.M247)
            self.words_list.clicked.connect(self.learning2)
            win.show()
            win.exec()
            if entry.text():
                hand_word = entry.text()
            elif hand_word is None:
                hand_word = word
            self.hand_word = hand_word
            confirmation = self.q.StandardButton.No
            if not checkbox_2.checkState() == Qt.CheckState.Unchecked:
                self.not_save_active = True
            if checkbox_2.checkState() == Qt.CheckState.Unchecked and not self.not_save_active:
                if  hand_word not in self.fa_words and hand_word != word:
                    self.saveReplacedWords(word, hand_word)
                    if not re.findall(r' ', hand_word):
                        self.saveNewWord(hand_word)                       
                elif hand_word == word:
                    if hand_word not in self.fa_w_not_found and hand_word not in self.fa_words:
                        if not re.findall(r' ', hand_word):
                            confirmation = self.qq(None, self.M134, f"{self.M137} ' {hand_word} '", self.q.StandardButton.Yes | self.q.StandardButton.No)
                            if confirmation == self.q.StandardButton.Yes:
                                self.saveNewWord(hand_word)
            else:
                if hand_word != word:
                    confirmation = self.qq(None, self.M245, f"{self.M245} {word} ← {hand_word}", self.q.StandardButton.Yes | self.q.StandardButton.No)
                    if confirmation == self.q.StandardButton.Yes:
                        self.saveReplacedWords(word, hand_word)
                        confirmation = self.q.StandardButton.No
                elif hand_word not in self.fa_w_not_found and hand_word not in self.fa_words:
                    if not re.findall(r' ', hand_word):
                        confirmation = self.qq(None, self.M134, f"{self.M137} ' {hand_word} '", self.q.StandardButton.Yes | self.q.StandardButton.No)
                        if confirmation == self.q.StandardButton.Yes:
                            self.saveNewWord(hand_word)
            return hand_word
        except Exception as e:
            self.info3(str(e))
            return word
    def searching_farhang (self):
        try:
            self.clicked=False
            self.clicked_trueWords=False
            self.words_list.setText(self.M249)
            lines=[]
            self.skip_save=True
            similar_word = self.farsi_auto_correction(self.hand_word)
            if similar_word is not None:
                self.corrected_word= similar_word
                self.button_Farhang.setText(self.M250+" "+self.corrected_word)
            else:
                return self.button_Farhang.setText("کلمه صحیح یافت نشد")
            if len(self.farhang) < 100:
                with open(os.path.join(self.bookdir, 'fa_dic.json'), 'r', encoding='utf-8') as f:
                    self.farhang = f.read().splitlines()

            for line in self.farhang:
                words = words = re.split(r'(\W+)', line)
                if self.corrected_word in words:
                    lines.append(line+'\n')
                else:continue
            if not len(lines)>0:
                FARHANG_RESULTS=self.corrected_word
            else:
                FARHANG_RESULTS=" ".join(lines)
            self.Farhang.setText(FARHANG_RESULTS)
            self.highlight_words( self.Farhang,self.corrected_word)
            return lines
        except:
            pass
    def words_list_F(self):
        try:
            if  self.clicked_trueWords==False :
                if not self.clicked==True  :
                    self.clicked=True
                    with open(self.replace, 'r', encoding='utf-8') as f:
                        self.replaced_=f.read()
                    self.Farhang.clear()
                    self.listReplaced=len(self.replaced_)
                    self.Farhang.setText(self.replaced_)
                    self.Farhang.setReadOnly(False)
                    self.words_list.setText(f"{self.M248}  {self.M249} ")
                else:
                    if  len(self.Farhang.toPlainText()) >14:
                        with open(self.replace, 'w' , encoding="utf-8") as f:
                            f.write(self.Farhang.toPlainText())
                        with open(self.replace, 'r', encoding='utf-8') as f:
                            self.replaced_lines =f.read().splitlines()
                    self.Farhang.setReadOnly(True)
                    self.clicked=False
                    self.words_list.setText(self.M247)
                    self.clicked_trueWords=True
            else:
                if not self.clicked==True:
                    self.clicked=True
                    with open(self.not_found_file, 'r', encoding='utf-8') as f:
                        self.not_found_words=f.read()
                    self.list_newWords=len(self.replaced_)
                    self.Farhang.clear()
                    self.Farhang.setText(self.not_found_words)
                    self.Farhang.setReadOnly(False)
                    self.words_list.setText(f"{self.M248}  {self.M247} ")
                else:
                    if  len(self.Farhang.toPlainText()) >5 :
                        with open(self.not_found_file, 'w' , encoding="utf-8") as f:
                            f.write(self.Farhang.toPlainText())
                        with open(self.not_found_file, 'r', encoding='utf-8') as f:
                            self.fa_w_not_found =f.read().splitlines()
                    self.Farhang.setReadOnly(True)
                    self.clicked=False
                    self.words_list.setText(self.M249)
                    self.clicked_trueWords=False
        except:
            self.info0(self.M254)
            self.Farhang.setReadOnly(True)
    def highlight_words(self, text_edit, word):
        if  self.stopBook==True:
            return
        try:
            format =QTextCharFormat()
            format.setForeground(QBrush(QColor("blue")))
            cursor = text_edit.textCursor()
            cursor.setPosition(0)
            while True:
                if  self.stopBook==True:
                    return
                cursor = text_edit.document().find(word, cursor.position())
                if not cursor.isNull():
                    cursor.mergeCharFormat(format)
                else:
                    break
        except:
            pass
    def highlight_words_curser(self, text_edit, word):
        try:
            format =QTextCharFormat()
            format.setForeground(QBrush(QColor("blue")))
            cursor = text_edit.textCursor()
            cursor.setPosition(0)
            while True:
                cursor = text_edit.document().find(word, cursor.position())
                if not cursor.isNull():
                    cursor.mergeCharFormat(format)

                    text_edit.setTextCursor(cursor)
                    text_edit.ensureCursorVisible()
                else:
                    break
        except:
            pass
    def al110(self,state):
        self.internal=True
        if self.dgh==False :
            self.deghat=True
            self.dgh=True
        else:
            self.deghat=False
            self.dgh=False
    def on_button_clicked_2(self):
        try:
            QApplication.exit()
        except:
            pass
    def on_button_clicked(self):
        try:
            if self.entery_s:
               self.entery_s.clear()
            if self.separate_search==True:
               self.separate_search=False
            self.win2.accept
        except:
            pass
    def save_file_in_qt5(self):
        self.export_docx_=True
        try:
            doc = dox.Document()
            size = int(self.size_box.value())
            new_font = self.font_box.currentText()
            if self.color_code[0]:
                rgb_color = tuple(int(self.color_code[0][i]) for i in range(3))
            else:
                rgb_color = (0, 0, 0)
            text=self.entry2.toPlainText().split('\n')
            if  self.pdf_convert==True or len(text)>2 :
                
                self.file_pat, _ = QFileDialog.getSaveFileName(self, "Export Results as docx file", "",
                                                "Word files (*.docx);;All Files (*)")
                if not self.file_pat:
                    return
                for para in text :
                    paragraph = doc.add_paragraph(para)
                    for run in paragraph.runs:
                        run.font.size = dox.shared.Pt(size)  # Set the font size
                        run.font.name = new_font  # Set the font style
                        run.font.color.rgb =dox.shared.RGBColor(*rgb_color)  # Set the font color
                doc.save(self.file_pat)
                os.startfile(self.file_pat)
                self.process_finish()
                self.rev=False
            else: self.update_notification(self.M197)
        except Exception as e:
            self.process_finish()
            self.rev=False
            self.q.information(self,self.M23, f"{str(e)}")
    def paste_from_clipboard_qt5(self):
            try:
                self.entry.setText( pyperclip.paste())
            except :
                pass
    def copy_to_clipboard_qt5(self):
        try:
            selected_text = self.entry2.toPlainText()
            pyperclip.copy(selected_text)
            self.update_notification("متن کپی شد")
        except Exception:
            self.update_notification("کپی نشد")
    def update_total_para(self):
        if len (self.total.text())>0:
          try:
           self.total_para = int(self.total.text())
          except:
              self.total_para=5
    def research(self):
        self.stopBook=False
        self.total.setText(str(self.total_para))
        self.cunter=0
        self.entry2.clear()
        self.stops = False
        try:
            lines=[]
            word= f"{self.entry.toPlainText()}"
            if len(word)<2:
                return None
            if self.dict in [self.M207] :
                results=self.abjad(word)
                self.entry2.clear()
                self.entry2.setText(results)
                self.separate_search=False
                return self.entry.toPlainText()
            elif self.search_process==self.M207:
                results=self.abjad(word)
                self.entry2.clear()
                self.entry2.setText(results)
                self.separate_search=False
                return self.entry.toPlainText()
            else:
                self.total_para=int(self.total.text())
                if not self.separate_search==True:
                    if  word in self.persian_conjunctions:
                        return self.entry.toPlainText()
                for line in self.ketab_:
                    if self.stopBook==True:
                        self.ketab_=""
                        lines=[]
                        return None
                    if self.end==True:
                        break
                    if  self.deghat:
                        if self.ktb in [self.M194,self.M211,self.M174]:
                            word=self.replace_characters(word,True)
                            line2=self.clean_text_arabic(line)
                            line1=self.replace_characters(line2,True,False)
                            words = re.split(self.special_chars2, line1)
                            word=self.clean_text_arabic(word)
                        else:
                            words = re.split(r'(\W+)', line)
                            word=self.replace_characters(word,False)
                        if words:
                            if word in words:
                                lines.append(line+'\n')
                                self.cunter+=1
                                if len(lines)+1>self.total_para or self.stops==True:
                                    break
                                if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                    else:
                        if self.ktb in [self.M194,self.M211,self.M174]:
                            word=self.replace_characters(word,True)
                            line2=self.clean_text_arabic(line)
                            line1=self.replace_characters(line2,True,False)
                            word=self.clean_text_arabic(word)
                            if word in line1:
                                lines.append(line+'\n')
                                self.cunter+=1
                                if len(lines)+1>self.total_para or self.stops==True :
                                    break
                                if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                        else:
                            word=self.replace_characters(word,False)
                            if word in line:
                                lines.append(line+'\n')
                                self.cunter+=1
                                if len(lines)+1>self.total_para or self.stops==True:
                                    break
                                if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                self.entry2.clear()
                if  self.stopBook==True:
                        self.ketab_=""
                        lines=[]
                        return None
                self.total_cunt.setText(str(self.cunter))
                if not lines:
                    self.entry2.setText(f"Not found ... {word}  \n Searching possible only with arabic or persian languages  ")
                    self.separate_search=False
                    return None
                else:
                    results = '\n'.join(lines)
                    self.entry2.setText(results)
                    self.highlight_words(self.entry2, word)
                    self.separate_search=False
                    return self.entry.toPlainText()
        except:
                self.separate_search=False
                return None

    def book_search(self, text, word, results):
        self.end = False
        try:
            self.dgh = False
            self.internal = False
            self.win2 = QDialog(self.tab_widget)
            self.win2.setWindowTitle(f"                                                                                                          {self.search_process}   ")
            self.win2.setGeometry(560, 40, 786, 660)
            try:
                self.win2.setStyleSheet(open(self.them1).read())
            except Exception as e:
                print(f"Failed to load stylesheet: {str(e)}")
            label = QTextEdit(self.win2)
            label.setText(f"{text}")
            self.highlight_words(label,word)
            label.setReadOnly(True)
            label.move(14, 20)
            label.resize(765, 76)
            self.entry = QTextEdit(self.win2)
            size = 14
            font = QFont()
            font.setPointSize(size)
            self.entry2 = QTextEdit(self.win2)
            self.total = QLineEdit(self.win2)
            self.total_cunt = QLineEdit(self.win2)
            label2 = QLabel(self.win2)
            self.totalable = QLabel(self.win2)
            self.total_cuntlable = QLabel(self.win2)
            button = QPushButton(self.win2)
            button_search = QPushButton(self.win2)
            button_copy = QPushButton(self.win2)
            button_paste = QPushButton(self.win2)
            button_save = QPushButton(self.win2)
            checkbox = QCheckBox(self.M136, self.win2)
            checkbox_2 = QCheckBox(self.M178, self.win2)
            layout = QVBoxLayout(self.win2)
            
            label2.setText(self.M158)
            self.entry.setFont(font)
            self.entry.move(150, 140)
            self.entry.resize(520, 45)
            label2.move(300, 100)
            label2.resize(410, 45)
            
            self.entry2.setFont(font)
            self.entry2.setAlignment(Qt.AlignmentFlag.AlignRight)
            self.entry2.setReadOnly(True)
            self.entry2.resize(765, 360)
            self.entry2.move(14, 185)
            
            self.totalable.setFont(font)
            self.totalable.resize(270, 35)
            self.totalable.move(400, 550)
            self.totalable.setText(" حداکثر تعداد یافتن را وارد فرمایید")
            
            self.total.setFont(font)
            self.total.resize(66, 35)
            self.total.move(320, 550)
            self.total.setText(str(self.total_para))
            self.total.textChanged.connect(self.update_total_para)
            
            self.total_cuntlable.setFont(font)
            self.total_cuntlable.resize(270, 35)
            self.total_cuntlable.move(500, 610)
            self.total_cuntlable.setText("تعداد یافته ها")
            
            self.total_cunt.setFont(font)
            self.total_cunt.resize(66, 35)
            self.total_cunt.move(600, 610)
            self.total_cunt.setReadOnly(True)
            
            button.setText(self.M132)
            button.move(14, 610)
            button.resize(110, 35)
            self.entry.setText(word)
            self.entry2.setText(results)
            self.total_cunt.setText(str(self.i))
            self.highlight_words(self.entry2, word)
            
            button_search.setText(self.M198)
            button_search.move(14, 140)
            button_search.resize(135, 35)
            
            button_paste.setText(self.M21)
            button_paste.move(669, 140)
            button_paste.resize(110, 35)
            
            button_copy.setText(self.M20)
            button_copy.move(135, 610)
            button_copy.resize(110, 35)
            
            button_save.setText(self.M248)
            button_save.move(280, 610)
            button_save.resize(250, 35)
            button_search.clicked.connect(self.research)
            self.entry.textChanged.connect(self.research)
            self.entery_s = self.entry2
            button.clicked.connect(self.win2.accept)
            button_paste.clicked.connect(self.paste_from_clipboard_qt5)
            button_copy.clicked.connect(self.copy_to_clipboard_qt5)
            button_save.clicked.connect(self.save_file_in_qt5)

            def on_checkbox_state_changed(state):
                self.off_ketab()

            checkbox.stateChanged.connect(on_checkbox_state_changed)
            checkbox.move(14, 566)
            checkbox.resize(236, 35)
            
            checkbox_2.stateChanged.connect(self.al110)
            checkbox_2.move(14, 100)
            checkbox_2.resize(236, 35)
            
            self.win2.setLayout(layout)
            self.win2.exec()

            if self.stopBook:
                return word

            if self.entry.toPlainText():
                hand_word = self.entry.toPlainText()
            elif self.research():
                hand_word = self.research()
            else:
                hand_word = word

            return hand_word

        except Exception as e:
            print(f"Error: {str(e)}")
            if self.separate_search:
                self.separate_search = False
            return word
    def abjad(self,word):
        try:
            word = word.replace(' ', '')
            self.search_process=self.M207
            input_string=word
            abjad_dict = { 'ا': 1, 'ب': 2, 'پ': 2, 'ج': 3, 'چ': 3, 'د': 4, 'ه': 5, 'و': 6, 'ز': 7,
                        'ژ': 7, 'ح': 8, 'خ': 8, 'ط': 9, 'ی': 10, 'ک': 20, 'گ': 20, 'ل': 30, 'م': 40, 'ن': 50,
                            'س': 60, 'ع': 70, 'ف': 80, 'ص': 90, 'ق': 100, 'ر': 200, 'ش': 300, 'ت': 400, 'ث': 500,
                            'ذ': 700, 'ض': 800, 'ظ': 900, 'غ': 1000 }
            return str(sum(abjad_dict.get(char, 0) for char in input_string if char.isalpha()))
        except:
            return word

    def search(self, text):
        self.cunter=0
        self.end=False
        self.stops = False
        self.stopBook=False
        try:
            new_text = []
            lines=[]
            self.dgh=False
            self.internal=False
            if self.from_code in('fa','ar','ur') and self.ktb in [self.M194,self.M172]:
                text=self.replace_characters(text,True)
            else:
                if self.from_code in('fa','ar','ur'):
                   text=self.replace_characters(text,False)
            if self.sjmle:
                for_chunks =re.split(r'(\W+)', text)
                words = [''.join(for_chunks[i:i+3]) for i in range(0, len(for_chunks), 3)]
            else:
                words=re.split(r'(\W+)', text)
            for word in words:
                self.cunter=0
                if  len(word)<3 or self.from_code in('fa','ar','ur') and (self.ktb not in [self.M166,self.M167,self.M173,self.M188,] and not re.search(r'[\u0600-\u06FF]', word)):
                    new_text.append(word)
                    continue
                if self.end==True:
                    new_text.append(word)
                    continue
                if not self.sjmle:
                    if len(word)<3:
                        new_text.append(word)
                        continue
                if self.search_Active==False:
                    new_text.append(word)
                    continue
                elif self.dict.currentText()==self.M207:
                    results=self.abjad(word)
                    new_word=self.book_search(text,word,results)
                    new_text.append(word)
                else:
                    for line in self.ketab_:
                        if self.stopBook==True:
                           self.ketab_=""
                           return None
                        if  self.deghat:
                            if self.ktb in [self.M194,self.M211,self.M174] or self.deepgf:
                                line_celear=self.clean_text_arabic(line)
                                line_celear=self.replace_characters(line_celear,True,False)
                                word=self.clean_text_arabic(word)
                                words = re.split(r'(\W+)', line_celear)
                            elif self.ktb ==self.M172:
                                line_celear=self.replace_characters(line,True,False)
                                word=self.replace_characters(word,True)
                                words = re.split(r'(\W+)', line_celear)
                            else:
                                words = re.split(r'(\W+)', line)
                            if words:
                                if  (words[0] == word or words[-1] == word )or (len(words)>3 and (words[+1] == word or words[+2] == word)):
                                    lines.append(line+'\n')
                                    self.cunter+=1
                                    if len(lines)+1>self.total_para or self.stops==True:
                                        break
                                    if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                                else:
                                    if word in words:
                                        lines.append(line+'\n')
                                        self.cunter+=1
                                    if len(lines)+1>self.total_para or self.stops==True:
                                            break
                                    if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                        else:
                            if self.ktb in [self.M194,self.M211,self.M174] or self.deepgf:
                                line_celear=self.clean_text_arabic(line)
                                line_celear=self.replace_characters(line,True)
                                word=self.clean_text_arabic(word)
                                if word in line_celear:
                                    lines.append(line+'\n')
                                    self.cunter+=1
                                    if len(lines)+1>self.total_para or self.stops==True :
                                        break
                                if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                            elif self.ktb ==self.M172:
                                line_celear=line
                                word=self.replace_characters(word,False)
                                if word in line:
                                    lines.append(line+'\n')
                                    self.cunter+=1
                                    if len(lines)+1>self.total_para or self.stops==True:
                                        break
                                if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                            else:
                                word=self.replace_characters(word,False)
                                if word in line:
                                    lines.append(line+'\n')
                                    self.cunter+=1
                                    if len(lines)+1>self.total_para or self.stops==True:
                                        break
                                if  self.stopBook==True:
                                     self.ketab_=""
                                     lines=[]
                                     return None
                    if  self.stopBook==True:
                        self.ketab_=""
                        lines=[]
                        return None
                    if not lines:
                        new_text.append(word)
                    else:
                        results = '\n'.join(lines)
                        lines = []
                        new_word=self.book_search(text,word,results)
                        if word:
                            new_text.append(new_word)
                        else:
                            new_text.append(word)
            return ''.join(new_text)
        except Exception:
            return text
    def clean_text_arabic(self,text):
        normalized_text = unicodedata.normalize('NFD', text)
        cleaned_text = "".join(c for c in normalized_text if unicodedata.category(c) != 'Mn')
        return cleaned_text
    def replace_characters(self, text,research,special_chars=True):       
        words = re.split(r'(\W+)', text)
        try:
            for i in range(len(words)): 
                if self.stopBook==True or self.stops==True:    
                    return text
                words[i] = words[i].replace('  ', ' ')
                words[i] = words[i].replace('  ', ' ')
                if words[i] == "اهلل":
                    words[i]='الله'
                words[i] = words[i].replace("الل ه", "الله")
                words[i] = words[i].replace("اهلل", "الله")
        except Exception :
            words[i] = ""
        return ''.join(words) 
    def replace_characters1(self, text,research,special_chars=True):
        words = re.split(r'(\W+)', text)
        try:
            for i in range(len(words)):
                words[i] = words[i].replace('َ', '')
                words[i] = words[i].replace('ُ', '')
                words[i] = words[i].replace('هـ', 'ه')
                words[i] = words[i].replace('  ', ' ')
                words[i] = words[i].replace('ي', 'ی')
                words[i] = words[i].replace('ئ', 'ی')
                words[i] = words[i].replace('ك', 'ک')
                words[i] = words[i].replace('ﻚ', 'ک')
                words[i] = words[i].replace('ﺑ', 'ب')
                words[i] = words[i].replace('ﺎ', 'ا')
                words[i] = words[i].replace('ﻧ', 'ن')
                words[i] = words[i].replace('ﺧ', 'خ')
                words[i] = words[i].replace('إ', 'ا')
                words[i] = words[i].replace('ؤ', 'و')
                words[i] = words[i].replace('أ', 'ا')
                words[i] = words[i].replace('إ', 'ا')
                words[i] = words[i].replace('ۀ', 'ه')#"""
                words[i] = words[i].replace('أ', 'ا')
                words[i] = words[i].replace('%', '٪')
                words[i] = words[i].replace('ْ', '')
                words[i] = words[i].replace('ٍ', '')
                words[i] = words[i].replace('ْ', '')
                words[i] = words[i].replace('ٍ', '')
                words[i] = words[i].replace('ً', '')
                words[i] = words[i].replace('ٌ', '')
                words[i] = words[i].replace('ٍ', '')
                words[i] = words[i].replace('َ', '')
                words[i] = words[i].replace('ُ', '')
                words[i] = words[i].replace('', '')
                words[i] = words[i].replace(" ّ", "")  
                words[i] = words[i].replace(u"\u200F", "")  
                words[i] = words[i].replace(u"\u200E", "")  
                words[i] = words[i].replace(u"\u200D", "") 
                words[i] = words[i].replace(u"\u200B", "")  
                words[i] = words[i].replace(u"\u200A", "") 
                words[i] = words[i].replace(u"\u00A0", "")  
                words[i] = words[i].replace(u"\u0640", "")  
                words[i] = words[i].replace(u"\u0640", "") 
                words[i] = words[i].replace(" ْ", "") 
                words[i] = words[i].replace(u"\ufeb7","ش")  
                words[i] = words[i].replace(u"\ufeae","ر")  
                words[i] = words[i].replace(u"\ufe96","ت") 
                words[i] = words[i].replace(u"\ufedf","ل") 
                words[i] = words[i].replace(u"\ufeb0","ز")  
                words[i] = words[i].replace(u"\ufee8","ن") 
                words[i] = words[i].replace(u"\ufb93","گ") 
                words[i] = words[i].replace(u"\ufeb3","س")  
                words[i] = words[i].replace(u"\ufeec","ه") 
                words[i] = words[i].replace(u"\ufee3","م")  
                words[i] = words[i].replace(u"\ufecb","ع")
                words[i] = words[i].replace(' ِ', '')
                words[i] = words[i].replace(' ', '')
                words[i] = words[i].replace('  ', ' ')
        except Exception :
            words[i] = ""
        return ''.join(words)
    def auto_reverse_mix_text(self, text):
        grouped_words = []
        text = self.space_correction(text)
        words = re.split(r'(\W+)', text)
        arabic_pattern = re.compile(r'[\u0600-\u06FF]')
        length = len(words)-1
        for i in range(length):
            if all(i + j < length and not arabic_pattern.search(words[i + j]) for j in range(2)):
               grouped_word = words[i] + '__'
            else:
               grouped_word = words[i]+'""'
            grouped_words.append(grouped_word)
        new_text = "".join(grouped_words)
        new_words = new_text.split('""')
        new_words.reverse() 
        return str("".join(self.replace_reversed(new_words)))
    def replace_reversed(self, words):
        for i in range(len(words)): 
            words[i] = words[i].replace('__', ' ')  
            words[i] = words[i].replace('  ', ' ') 
            words[i] = words[i].replace('  ', ' ')                      
        return words
    def motaradef(self, text):
        new_text = []
        self.end = False
        try:
            words = re.split(r'(\W+)', text)
            win = QDialog(self.tab_widget)
            win.setGeometry(200, 200, 796, 500)
            win.setWindowTitle('مترادف ها')            
            label = QTextEdit(win)
            label.setText(text)
            label.setReadOnly(True)
            label.move(100, 20)
            label.resize(600, 96)
            entry = QLineEdit(win)
            label2 = QLabel(win)
            combo = QComboBox(win)  # Changed from self.qc to QComboBox
            button = QPushButton(win)  # Changed from self.qb to QPushButton
            checkbox = QCheckBox(self.M136, win)
            checkbox_2 = QCheckBox(self.M178, win)            
            entry.move(100, 166)
            entry.resize(600, 35)           
            label2.setText('لغت مترادف را بنویسید یا از لیست ریز انتخاب کنید')
            label2.move(100, 110)
            label2.resize(600, 35)           
            combo.insertItem(0, 'انتخاب مترادف')
            combo.setCurrentIndex(0)
            combo.move(100, 240)
            combo.resize(600, 35)            
            button.move(100, 300)
            button.resize(600, 35)           
            checkbox.move(100, 360)
            checkbox.resize(600, 35)           
            checkbox_2.move(100, 400)
            checkbox_2.resize(236, 35)
            layout = QVBoxLayout(win)  # Added layout
            win.setLayout(layout)
            if not self.S5:
                self.open_3()
            lin_word = []
            lines = ""
            word1 = ''
            self.dgh = False
            for word in words:
                if self.end or not re.search(r'[\u0600-\u06FF]', word):
                    new_text.append(word)
                    continue
                label.setText(text)
                self.highlight_words(label,word)
                adi = False
                for line in self.synonyms_str:
                    l = re.split(r'(\W+)', line)
                    if word in l:
                        lines += line
                    elif self.deghat and not word in line:
                        if len(word) > 2:
                            for suffix3 in self.suffixes:
                                if word.endswith(suffix3):
                                    new_word = word[:-len(suffix3)]
                                    if new_word in l:
                                        word1 = new_word
                                        adii = suffix3
                                        adi = True
                                        lines += line

                if not lines:
                    new_text.append(word)
                    continue
                else:
                    combo.clear()
                    lin_word = re.split(self.special_chars2, lines)
                    lines = ""
                    combo.addItems(lin_word)
                    lin_word = ""
                    hand_word = ""
                    def on_combobox_changed(text):
                        nonlocal hand_word
                        if not adi:
                            hand_word = text
                        else:
                            hand_word = text + adii
                        entry.setText(hand_word)
                    combo.currentTextChanged.connect(on_combobox_changed)
                    entry.setText(word)
                    button.setText(self.M132)
                    button.clicked.connect(win.close)
                    def on_checkbox_state_changed(state):
                        combo.clear()
                        self.off_ketab()
                    checkbox.stateChanged.connect(on_checkbox_state_changed)
                    checkbox_2.stateChanged.connect(self.al110)                  
                    win.show()
                    win.exec()
                    if entry.text():
                        hand_word = entry.text()
                    elif hand_word == "" and word:
                        if not adi:
                            hand_word = word
                        else:
                            hand_word = word1 + adii
                    new_text.append(hand_word)
                    combo.clear()
                    continue
        except Exception as e:
            self.update_notification(f"  {self.M23}  : {e}")        
        return "".join(new_text)

    def stop(self):
        self.excel=False
        self.stops=True
        self.thread_active = False
        self.skip_rtl=False
        self.pdf=False
        self.docx=False
        self.docx2=False
        self.console=False
        self.pdf_convert=False
        self.update_notification(self.M42)
    def stop_book(self):
        self.stopBook=True
    def process_finish(self):
        self.pack_downloaded = False
        self.excel=False
        self.thread_active = False
        self.skip_rtl=False
        self.pdf=False
        self.docx=False
        self.stops=False
        self.console=False
        self.pdf_convert=False
        self.fileMode=False
        gc.collect()
        self.update_notification(self.M42)
    def defultthem(self):
        try:
            self.defultthem()
        except :
            self.update_notification(self.M254)
            return
    def restart_program1(self):
        self.clear()
        self.defultthem()
        python = sys.executable
        os.execl(python, python, *sys.argv)
    def restart_program(self):
        self.clear()
        gc.collect()
        python = sys.executable
        os.execl(python, python, *sys.argv)
    def closeEvent(self, event):
        gc.collect()
        self.clear()
        sys.exit()
    def wait(self):
        QApplication.processEvents()
def main():
    app = QApplication(sys.argv)
    window = a()
    print("الحمد الله")
    print("اللّهم صل علي محمّد و آل محمّد و عجّل فرجهم و اهلک والعن اعدائهم")
    window.show()
    sys.exit(app.exec())
if __name__ == "__main__":
    main()


