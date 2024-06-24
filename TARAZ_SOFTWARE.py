# -*- coding: utf-8 -*-
#Email : mj.taraz@yahoo.com
# This software is the property of Taraz. 
# No part of this software may be reproduced in any form without the prior written permission of Taraz.
"""
Name: Taraz_Software.exe
Version: 1.0.0
Author: Apple
https://www.apple.com
Date Created: 2024
Description: text edit and translator and books search_internet etc .
"""
#بسم الله الرحمن الرحیم#
import sys
import os

import ttkthemes # type: ignore
script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, script_dir)
try:
    from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QLineEdit, QPushButton, QComboBox, QMessageBox,QTextEdit,QCheckBox
    from PyQt5.QtCore import Qt
    from PyQt5 import QtGui
    import re
    from bidi.algorithm import get_display
except:
    pass
import tkinter
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox,colorchooser
from tkinter.simpledialog import askstring


import time
class LazyImport:
    def __init__(self, module_name):
        self.module_name = module_name
        self.module = None
    def __getattr__(self, name):
        if self.module is None:
            self.module = __import__(self.module_name, fromlist=[name])
        return getattr(self.module, name)
Sequence = LazyImport("difflib")
collec = LazyImport("collections")
langdetec = LazyImport("langdetect")
rapidfuz = LazyImport("rapidfuzz")
pdf2= LazyImport("pdf2docx")
docx2= LazyImport("docx2pdf")
dox = LazyImport("docx")
SpellChecke = LazyImport("spellchecker")
Translator2 = LazyImport("argostranslate")
Translator1 = LazyImport("translatepy")
Translator3 = LazyImport("googletrans")
Translator5 = LazyImport("deep_translator")
ggl = LazyImport("google_searching")
wikipedia = LazyImport("wikipedia")
threading = LazyImport("threading")
gc = LazyImport("gc")
shutil = LazyImport("shutil")
farsi_tool = LazyImport("farsi_tools")
Speller = LazyImport("autocorrect")

requests = LazyImport("requests")
urllib = LazyImport("urllib.parse")
hashlib = LazyImport("hashlib")
unicodedata = LazyImport("unicodedata")
ftfy = LazyImport("ftfy")
def on_close():
    if messagebox.askokcancel("Quit  \n بستن برنامه", "Do you want to quit? \n\n  برنامه بسته شود؟"):    
        os._exit(0)
def importer():
    import argostranslate.package
    import argostranslate.translate
    import tkinter
   # import awesometkinter as tkinter.ttk
    try:
        from difflib import SequenceMatcher
        from collections import Counter
        import re
        from langdetect import detect
        from rapidfuzz import process
    except:
        pass
    try:
        from pdf2docx import Converter
        from docx2pdf import convert 
        import docx2pdf
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from spellchecker import SpellChecker
    except:
        pass
    try:
        from translatepy import Translator as tr
        from translatepy.translators.yandex import YandexTranslate 
    except:
        pass
    try:
        from googletrans import Translator 
    except:
        pass
    try:
        from deep_translator import GoogleTranslator 
        from deep_translator import MyMemoryTranslator 
    except:
        pass
    try:
        import threading
        import time
        import requests
        import hashlib
        import gc
        import shutil
    except:
        pass
    try:
        from farsi_tools import replace_ascii_digits_with_farsi,stop_words
        from autocorrect import Speller
        from google_searching import ggl
        import wikipedia
    except:
        pass
class TranslationWindow:
    def __init__(self, root):
       super().__init__()
       
       self.nam="Taraz Software_313                                                                                                                    بِسْمِ اللهِ الرَّحْمنِ الرَّحِیم"
       self.myColor = '#dee9ef'  
       self.myColor2= '#d3ecfa'
       self.myColor3='#8de0d1'
       self.myColor4='#d3ecfa'
       self.myColor5='#a67b19'
       self.myColor6='#a67b19'
       try:                             
            self.root = root
            root.title(self.nam)

            with open('them.json', 'r') as f:
                self.num = f.read()
            root.geometry("1100x669")
            root.configure(bg=self.myColor3)  # Setting color of main window to self.myColor
            self.dir= os.getcwd()
            style = ttkthemes.ThemedStyle()
            try:
                style.theme_use(self.num)  # Replace with the desired theme (e.g., 'adapta') itft1 breeze plastik adapta  radiance clearlooks kroc smog
            except:
                style.theme_use('itft1')
            style.configure('TButton', font=('arial', 11,'bold'))
           # style.configure('TCombobox', font=('arial', 16,'bold'))
            #style.configure('TButton', background='green')

            book = tkinter.StringVar(value="کتابخانه")        
            books = ["قرآن",'نهج البلاغه',"اصول کافي","نهج الفصاحه","حافظ","سه دقیقه در قیامت"]
            self.book_box = tkinter.OptionMenu(root, book, *books)
            self.book_box.pack()
            self.book_box.config(font=('Arial', 11,"bold"),bg=self.myColor4,highlightthickness=1)
            self.book_box.bind("<<StringVarSelected>>", self.gift)
            self.var = tkinter.IntVar()

            book.trace("w", self.gift)
            font_var = tkinter.StringVar(value='Arial')        
            fonts = ["Arial",'Arial (Arabic)','Simplified Arabic Fixed',
                     'Courier New (Arabic)','Urdu Typesetting','Sakkal Majalla',
                     'Simplified Arabic','Traditional Arabic', ]
            self.font_box = tkinter.OptionMenu(root, font_var, *fonts)            
            self.font_box.pack()       
            self.font_box.cget("text")
            self.size_var = tkinter.StringVar(value="16")  # Create a StringVar
            self.size_box = tkinter.Spinbox(root, from_=0, to=100, textvariable=self.size_var,width=10, relief="sunken", repeatdelay=500, repeatinterval=100,
                     font=("Arial", 12), bg="lightgrey", fg="blue", command=self.update_font)
            #self.size_box.set(14)
            self.size_box.pack()

            self.text_font = ( 'Arial', 19)
            font_var.trace("w", self.update_font)
            self.notif_console = self.create_console("white", "red", 10, 487, 777, 30)
            self.notif_console.config(font='Arial')
            self.notif_console.config(state=tkinter.DISABLED)
            self.empty_1=False

            self.regui()
            self.start1()          
       except :
              pass   
    def regui(self):
        try:
            root.configure(bg=self.myColor)  # Setting color of main window to self.myColor
        except :
            pass
    def start1(self):
        try:                       
            self.dir= os.getcwd() 
            self.size_box.config(font=('Arial', 11,"bold"),bg=self.myColor4,highlightthickness=1)
            self.font_box.config(font=('Arial', 11,"bold"),bg=self.myColor4,highlightthickness=1)
            # File Button
            self.file_button = ttk.Button(root, command=self.select_file,)
            # Translate Button
            self.translate_button = ttk.Button(root, command=self.translate,)
            self.trans_file_button = ttk.Button(root, command=self.trans_file,)
            self.gift_button = ttk.Button(root,  command=self.pdf_converter,)
            # Export Button
            self.export_button = ttk.Button(root,  command=self.export_docx,)
            # Clear Button
            self.clear_button = ttk.Button(root,  command=self.clear,)
            self.Qt_translator = ttk.Button(root,  command=self.Qt_translator_,)
            self.Qt_translator.pack()
            self.download = ttk.Button(root,command=self.stop,)
            self.download.pack()
            # Create and place the copy_button
            self.copy_button = ttk.Button(root, command=self.copy_to_clipboard,)
            self.copy_button.pack(pady=10, padx=10)        # paste from
            self.paste_button = ttk.Button(root, command=self.paste_from_clipboard,)
            self.color_button = ttk.Button(root, command=self.restart_program,)
            self.color_button.pack()
            # Help Button
            self.help_button = ttk.Button(root, command=self.show_help_message,) 
            self.lang_button = ttk.Button(self.root, text="فارسی", command=self.Lang, compound="center",)
            self.lang2_button = ttk.Button(self.root, text="English", command=self.Lang2, compound="center",)  
            self.lang_button.pack()
            self.lang2_button.pack()  
            self.source_language_label = ttk.Label(root)
            self.source_language_label.config(font=('Arial', 10,"bold"))
            self.target_language_label = ttk.Label(root,  text="To language :")
            self.target_language_label.config(font=('Arial', 10,"bold"))
            with open('index_lang.json', 'r') as f:
                lang = f.read()
                if lang == 'en':
                    self.lang_en()
                else:
                    self.lang_fa()  
            self.space = tkinter.IntVar()
            self.space_corect = tkinter.Checkbutton(root, text=self.M11, variable=self.space, command=self.aktive_correct)
            self.space_corect.pack()
            self.text=""  
            self.language_codes_2 = {
                "en":self.M100, 
                "auto":self.M104,
                "fa":self.M102,
                "de":self.M101,
                "ar":self.M109,
                "fr":self.M108,
                "zh": self.M103,
                "es": self.M110,
                "ru": self.M111,
                "it": self.M106,
                "tr":self.M112,
                "pt":self.M113,
                "id": self.M114,
                "nl": self.M107,
                "hi": self.M115,
                "ja": self.M116,
                "ur":self.M117}
            self.language_codes = {
                self.M100: "en",
                self.M104:"auto",
                self.M102: "fa",
                self.M101: "de",
                self.M109: "ar",
                self.M108: "fr",
                self.M103: "zh-CN",
                self.M110: "es",
                self.M111: "ru",
                self.M106: "it",
                self.M112: "tr",
                self.M113: "pt",
                self.M114: "id",
                self.M107: "nl",
                self.M115: "hi",
                self.M116: "ja",
                self.M117: "ur",}
            self.correction = tkinter.Checkbutton(root, text=self.M84, variable=self.var, command=self.aktive_correction)
            self.correction.pack()
            self.translator_var = tkinter.StringVar()           
            self.translator_var_label = ttk.Label(root,  text=self.M124)           
            self.translator_var_label.config(font=('Arial', 11))
            self.translator_var.set(self.M122)  # Set default value
            self.translator_menu = ttk.Combobox(root,state="readonly", textvariable=self.translator_var,
                                                values=[self.M118, self.M119,self.M120,self.M121,self.M122,self.M201,self.M123])
            self.translator_menu.bind("<<ComboboxSelected>>", self.lang_code)
            self.translator_menu.pack()
            self.translator_menu.config(font=('Arial', 13))
            them=[
                'itft1','aquativo','plastik','radiance','clearlooks',
                'adapta','kroc','breeze','smog','alt',
                'classic','winnative','clam','default', 
                'vista', 'xpnative','arc','elegance' 
            ]
            self.them_num_var = tkinter.StringVar()  
            self.them_num_var.set(self.num)
            self.them_num = ttk.Combobox(root,state="readonly", textvariable=self.them_num_var,
                                                values=them)
            self.them_num.bind("<<ComboboxSelected>>", self.them)
            self.them_num.pack()
            self.them_num.config(font=('Arial', 13))


            document_mode = tkinter.StringVar(value=self.M96)        
            docu = [self.M96,self.M98,self.M99,self.M97, ] 
            self.document_mode_box = tkinter.OptionMenu(root, document_mode, *docu)
            self.document_mode_box.pack() 
            self.document_mode_box.config(font=('Arial', 11,"bold"),bg=self.myColor4,highlightthickness=1)

            fot = tkinter.StringVar(value=self.M12)        
            docfot = [self.M153,self.M12,self.M146,self.M152  ,self.M182,self.M196] 
            self.format_box= tkinter.OptionMenu(root, fot, *docfot)
            self.format_box.pack() 
            self.format_box.config(font=('Arial', 11,"bold"),bg=self.myColor4,highlightthickness=1)

            self.source_language_combo = ttk.Combobox(root, state="readonly",values=[self.M102,self.M100, self.M103,self.M110,self.M101,self.M104,
                                                                    self.M109,self.M113,self.M106,self.M107,
                                                                    self.M108,  self.M111,
                                                                    self.M112,self.M113,self.M114,self.M115,
                                                                    self.M116,self.M117],cursor="hand2")
            self.source_language_combo.set(self.M102)           
            # self.target Language Label and Combo
            self.target_language_combo = ttk.Combobox(root,state="readonly",values=[self.M102,self.M100,self.M103,self.M113, self.M101,
                                                                    self.M109,self.M106,self.M107,
                                                                    self.M108, self.M110, self.M111,
                                                                    self.M112,self.M113,self.M114,self.M115,
                                                                    self.M116,self.M117],cursor="hand2")
            self.source_language_combo.pack()
            self.target_language_combo.pack()
            self.target_language_combo.set(self.M102)
            self.target_language_combo.config(font=('Arial', 12,"bold"))
            self.source_language_combo.config(font=('Arial', 12,"bold"))
            self.source_language_combo.bind("<<ComboboxSelected>>", self.lang_code)
            self.target_language_combo.bind("<<ComboboxSelected>>", self.lang_code)
            self.output_console = self.create_console("white", "blue", 550, 100, 520, 380)
            self.output_console.configure(state=tkinter.DISABLED)
            self.input_console = self.create_console("white", "black", 10, 100, 520, 380)
            self.input_console.bind('<Return>', self.translate_aout) 
            self.info_console = self.create_console("white", "red", 800, 485, 290, 30)
            self.info2_console = self.create_console("white", "green", 800, 530, 290, 30)
            self.info2_console.config(font='Arial')
            self.info3_console = self.create_console("white", "green", 800, 575, 290, 66)
            self.info3_console.config(font='Arial')
            self.info_console.config(font='Arial')
            self.info_console.config(state=tkinter.DISABLED)
            self.info2_console.config(state=tkinter.DISABLED)
            self.info3_console.config(state=tkinter.DISABLED)
            # Add Scrollbar to Output Console
            input_scrollbar = tkinter.Scrollbar(self.root, command=self.input_console.yview)
            output_scrollbar = tkinter.Scrollbar(self.root, command=self.output_console.yview)
            output_scrollbar.place(x=1070, y=110, height=360)
            input_scrollbar.place(x=530, y=110, height=360)
            self.output_console.config(yscrollcommand=output_scrollbar.set)
            self.input_console.config(yscrollcommand=input_scrollbar.set)
            self.translate_pay=False
            self.arg_para=False
            self.start2()
            self.start()
        except :
            pass
    def start2(self):
        # Copy Button
        #==========Check boxes====================
        self.font_box.place(x=570, y=38,width=210, height=25)
        self.correction.place(x=790, y=38,width=95, height=25)
        self.space_corect.place(x=900, y=38,width=180, height=25)
        self.size_box.place(x=500, y=38,width=66, height=25)
        self.document_mode_box.place(x=14, y=38,width=178, height=25)
        self.format_box.place(x=200, y=38,width=210, height=25)
        self.book_box.place(x=420, y=38,width=78, height=25)
        #===========combo=========================
        self.source_language_label.place(x=15, y=10,width=210, height=25)
        self.source_language_combo.place(x=118, y=10,width=170, height=25)
        self.them_num.place(x=700, y=10,width=110, height=25)
        self.target_language_label.place(x=370, y=10,width=210, height=25)
        self.target_language_combo.place(x=470, y=10,width=170, height=25)
        self.translator_var_label.place(x=835, y=10,width=77, height=25)
        self.translator_menu.place(x=910, y=10,width=170, height=25)

       # self.show_search.place(x=120, y=555, width=270, height=30)
        
        #===========================================   
        # Translate Button     
        self.Qt_translator.place(x=400, y=530, width=140, height=50)  
        self.download.place(x=560, y=530, width=140, height=50)

        self.trans_file_button.place(x=120, y=555, width=270, height=30)
        self.translate_button.place(x=120, y=525, width=270, height=30)
      #  self.jmt.place(x=120, y=520, width=270, height=36)       
        # paste Button
        self.paste_button.place(x=10, y=590, width=96, height=50)
        self.file_button.place(x=10, y=530, width=96, height=50)
        # Gift Button 
        self.gift_button .place(x=110, y=590, width=96, height=40)
        #lang_button
        self.lang_button.place(x=310, y=590, width=96, height=40)
        #lang2_button
        self.lang2_button.place(x=210, y=590, width=96, height=40)
        # Clear Button
        self.clear_button.place(x=410 ,y=590, width=96, height=40)
        # Help Button
        self.help_button.place(x=510, y=590, width=96, height=40)
        #color
        self.color_button.place(x=610, y=590,width=96, height=40)
        self.copy_button.place(x=710, y=590, width=85, height=50)
        self.export_button.place(x=710, y=530, width=85, height=50)   
        self.internet_aktive=False
        self.book_aktive=False
        self.dict_aktive=False
        self.coorrect_aktive=True
        self.search_Active=False
        self.separate_search=True
        self.offline_installed=False
        self.ketab_=""
        self.file_path=''
        self.tranc_err=False
        self.word_office_not_installed=False
        self.search_process=''
        self.save_word=False
        self.clicked=False
        self.clicked_trueWords=False
        self.pdf=False
        self.excel=False 
        self.words2=''
        self.i=0
        self.update_font() 
        self.anti_crash()
        self.books_handle()
    def jamal(self):
        self.info2("https://www.iranlawclinic.com")
        try:
            os.startfile('کلینیک حقوقی ایران.html')
        except:
             self.info("\n https://www.iranlawclinic.com")

    def them(self,event):
        t=self.them_num.get()
        with open('them.json', 'w') as f:
                f.write(t)
        result = messagebox.askquestion(self.M258,self.M87)       
        if result == 'yes':  
           self.restart_program()
        return
    def anti_crash(self):
        try:
            self.perian_num=tkinter.BooleanVar()
            self.Farsi_text_edit=tkinter.BooleanVar() 
            self.virast=tkinter.BooleanVar()            
            self.space_word__disabled=tkinter.BooleanVar()
            self.using_orginal_text_enable=tkinter.BooleanVar()
            self.rtl_format_true=tkinter.BooleanVar()
            self.text_edit=tkinter.BooleanVar()
            self.nazar=tkinter.BooleanVar()
            self.jomlesazi=tkinter.BooleanVar()
            self.deghat=tkinter.BooleanVar()
            self.sjmle=tkinter.BooleanVar()
            self.gf=tkinter.BooleanVar()
            menubar = tkinter.Menu(self.root)
            self.root.config(menu=menubar)
            self.file_menu = tkinter.Menu(menubar, tearoff=0)
            self.edit_menu = tkinter.Menu(menubar, tearoff=0)
            self.rakb_menu = tkinter.Menu(menubar, tearoff=0)
            self.frg = tkinter.Menu(menubar, tearoff=0)
            self.format = tkinter.Menu(menubar, tearoff=0)
            menubar.add_cascade(label=self.M1, menu=self.file_menu)
            menubar.add_cascade(label=self.M8, menu=self.edit_menu)
            menubar.add_cascade(label=self.M150, menu=self.rakb_menu)
            menubar.add_cascade(label=self.M212, menu=self.frg)
            self.file_menu.add_command(label=self.M6, command=on_close) 
            self.file_menu.add_command(label=self.M2, command=self.new_file)
            self.file_menu.add_command(label=self.M3, command=self.open_file)
            self.file_menu.add_command(label=self.M4, command= self.save_file)
            self.file_menu.add_command(label=self.M411, command= self.save_file_in)
            self.file_menu.add_separator()
            self.edit_menu.add_checkbutton(label=self.M94, variable=self.virast)
            self.edit_menu.add_checkbutton(label=self.M127,variable=self.jomlesazi)
            self.edit_menu.add_checkbutton(label=self.M14, variable=self.Farsi_text_edit)           
            self.edit_menu.add_checkbutton(label=self.M11, variable=self.space_word__disabled)
            self.edit_menu.add_checkbutton(label=self.M16, variable=self.perian_num)
            self.edit_menu.add_checkbutton(label=self.M17, variable=self.rtl_format_true)
            self.edit_menu.add_checkbutton(label=self.M18, variable=self.using_orginal_text_enable)          
            self.rakb_menu.add_command(label=self.M10, command=self.choose_color)
            self.rakb_menu.add_command(label=self.M19, command=self.cut_text)
            self.rakb_menu.add_command(label=self.M20, command=self.copy_text)
            self.rakb_menu.add_command(label=self.M21, command=self.paste_text)
            def sent_degh ():
                self.sjmle.set(0)
                self.gf.set(0)
            def sent_degh_2 ():
                self.deghat.set(0)
                self.gf.set(0)
            def sent_degh_3():
                self.deghat.set(0)
                self.sjmle.set(0)                   
            self.frg.add_checkbutton(label=self.M178, variable=self.deghat,command=sent_degh)
            self.frg.add_checkbutton(label=self.M180, variable=self.sjmle,command=sent_degh_2)
            self.frg.add_checkbutton(label=self.M213, variable=self.gf,command=sent_degh_3)
            self.frg.add_command(label=self.M255,command=self.show)
        except:
            pass
    def books_handle(self):
        try:
            self.rabt_words=farsi_tool.stop_words()
            dic_labels= [self.M125, self.M154,self.M181, self.M159, self.M160, self.M166,self.M207,
                                self.M167, self.M168, self.M169, self.M173,self.M188,self.M213]
            ketab_labels = [self.M172,self.M194,
                    self.M170,
                    self.M208, self.M209, self.M210, self.M211,
                    self.M174, self.M175, self.M179,self.M244,
                    self.M185, self.M186, self.M187,self.M183,
                    self.M190, self.M192,self.M193,self.M213
                    ]
            int_labels= [self.M163, self.M164,self.M238,self.M73,self.M213]
            sina_labels= [self.M216,self.M217,self.M218,self.M219,self.M220,self.M221,
                            self.M222,self.M223,self.M224,self.M225,self.M226,
                            ]
            sher_labels = [self.M171,self.M189,self.M191,self.M228,self.M229,self.M230,
                            self.M231,self.M232,self.M233,self.M234,self.M235,
                            self.M213
                    ]
            
            teb_labels = [self.M185, self.M199,self.M242 ,self.M213 ]
            self.dict= ttk.Combobox(root, state="readonly",values=dic_labels,cursor="hand2")
            self.dict.set(self.M177)           
            self.dict.pack()
            self.dict.config(font=('Arial', 11))
            self.dict.bind("<<ComboboxSelected>>", self.on_ketab_1)

            self.sher= ttk.Combobox(root, state="readonly",values=sher_labels,cursor="hand2")
            self.sher.set(self.M227)           
            self.sher.pack()
            self.sher.config(font=('Arial', 11))
            self.sher.bind("<<ComboboxSelected>>", self.on_ketab_6)

            self.books= ttk.Combobox(root, state="readonly",values=ketab_labels,cursor="hand2")
            self.books.set(self.M176)           
            self.books.pack()
            self.books.config(font=('Arial', 11))
            self.books.bind("<<ComboboxSelected>>", self.on_ketab_2)
            self.internet= ttk.Combobox(root, state="readonly",values=int_labels,cursor="hand2")
            self.internet.set(self.M165)           
            self.internet.pack()
            self.internet.config(font=('Arial', 10))
            self.internet.bind("<<ComboboxSelected>>", self.on_ketab_3)
            self.sina= ttk.Combobox(root, state="readonly",values=sina_labels,cursor="hand2")
            self.sina.set(self.M203)           
            self.sina.pack()
            self.sina.config(font=('Arial', 11))
            self.sina.bind("<<ComboboxSelected>>", self.on_ketab_4)

            self.teb= ttk.Combobox(root, state="readonly",values=teb_labels,cursor="hand2")
            self.teb.set(self.M241)           
            self.teb.pack()
            self.teb.config(font=('Arial', 11))
            self.teb.bind("<<ComboboxSelected>>", self.on_ketab_7)
            self.correct= ttk.Combobox(root, state="readonly",values=[self.M128,self.M15,self.M213,],cursor="hand2")
            self.correct.set(self.M84)   
            self.correct.pack()
            self.correct.config(font=('Arial', 11))
            self.correct.bind("<<ComboboxSelected>>", self.on_ketab_5)
            #=============books================================================
            self.correct.place(x=15, y=66,width=178, height=25)
            self.dict.place(x=195, y=66,width=160, height=25)
            self.books.place(x=360, y=66,width=160, height=25)
            self.sina.place(x=521, y=66,width=140, height=25)
            self.sher.place(x=664, y=66,width=135, height=25)
            self.internet.place(x=807, y=66,width=135, height=25)
            self.teb.place(x=945, y=66,width=135, height=25)
            self.info2(self.correct.get())
        except:
                pass
    def off_ketab(self):
        self.end=True
        self.search_Active=False
        self.separate_search=False
        self.internet_aktive=False
        self.book_aktive=False
        self.dict_aktive=False
        
        self.internet.set(self.M165) 
        self.dict.set(self.M177) 
        self.sher.set(self.M227) 
        self.sina.set(self.M203) 
        self.books.set(self.M176)
        self.teb.set(self.M241) 
       # self.coorrect_aktive=False
        self.info2('') 
        return
    def on_ketab_1(self,event):
        QApplication.exit()
        self.info2('') 
        self.on_button_clicked_2()
        self.internet.set(self.M165) 
        self.sina.set(self.M203) 
        self.books.set(self.M176) 
        self.sher.set(self.M227) 
        self.teb.set(self.M241)
        if not  self.dict.get() in [self.M213,self.M177]:
            self.spm=False
            self.separate_search=False
            self.search_Active=True
            self.internet_aktive=False
            self.book_aktive=False
            self.dict_aktive=True
            if  self.dict.get() != self.M125:
                    self.info2(self.dict.get())               
                    try: 
                        if  self.thread_active==False:
                            self.separate_search=True
                            src_book_dic=self.searching_book_options()
                            if self.spm==True or src_book_dic==True :
                                self.run_book_search()
                            elif  self.dict.get()in [self.M207] :
                                self.search_process=self.dict.get()
                                self.run_book_search()
                    except:
                        pass                
            else: self.info2(self.M125)
        else:
            self.off_ketab()
    def on_ketab_2(self,event):
        QApplication.exit()
        self.info2('') 
        self.on_button_clicked_2()
        self.internet.set(self.M165) 
        self.dict.set(self.M177) 
        self.sina.set(self.M203) 
        self.sher.set(self.M227) 
        self.teb.set(self.M241)
        if not  self.books.get() in [self.M213,self.M176]:    
            self.separate_search=False
            self.spm=False
            self.search_Active=True
            self.internet_aktive=False
            self.dict_aktive=False
            self.book_aktive=True
            try:
                if  self.thread_active==False:
                    self.separate_search=True
                    src_book_dic=self.searching_book_options()
                    if self.spm==True or src_book_dic==True :  
                        self.run_book_search()
            except:
                pass    
        else:
            self.off_ketab()     
    def on_ketab_3(self,event):
        QApplication.exit()
        self.info2('') 
        self.on_button_clicked_2()
        self.dict.set(self.M177) 
        self.sina.set(self.M203) 
        self.books.set(self.M176)
        self.sher.set(self.M227) 
        self.teb.set(self.M241)
        if self.internet.get() in [self.M238]:
            self.jamal()
            self.internet.set(self.M165)
            self.internet_aktive=False
            self.update_notification("")
            return
        if self.internet.get() in [self.M73]:
            self.selected_url = ""
            self.url = ""
            self.download_manager()
            self.internet.set(self.M165)
            self.internet_aktive=False
            self.update_notification("")
            return
        if not  self.internet.get() in [self.M213,self.M165]:
            self.search_Active=True
            self.book_aktive=False
            self.dict_aktive=False
            self.separate_search=False
            self.internet_aktive=True
            self.info2(self.internet.get())  
            try:
                if  self.thread_active==False:
                    self.separate_search=True
                    self.search_process=self.internet.get()
                    self.run_book_search()
            except:
                pass
        else:
            self.off_ketab()
    def on_ketab_4(self,event):
        QApplication.exit()
        self.info2('') 
        self.on_button_clicked_2()
        self.internet.set(self.M165) 
        self.dict.set(self.M177) 
        self.books.set(self.M176)
        self.sher.set(self.M227) 
        self.teb.set(self.M241)
        if not self.sina.get() in [self.M213,self.M203]:
            self.info2(self.sina.get())
            self.spm=False
            self.separate_search=False
            self.search_Active=True
            self.internet_aktive=False
            self.dict_aktive=False
            self.book_aktive=True
            try:
                if  self.thread_active==False:
                    self.separate_search=True
                    src_book_dic=self.searching_book_options()
                    if self.spm==True or src_book_dic==True :   
                        self.run_book_search()
                        
            except:
                pass 
        else:
            self.off_ketab()
    def on_ketab_5(self,event):
        self.coorrect_aktive=True
        self.search_Active=True
        self.info2(self.correct.get())
        if self.correct.get()==self.M213:
           self.coorrect_aktive=False
    def aktive_correction(self):

        if self.var.get()==1:
            self.coorrect_aktive=True
            self.search_Active=True
            self.info2(self.M128)
            self.correct.set(self.M128)
        else:
            self.correct.set(self.M133)
            self.info2(" ")
            self.coorrect_aktive=False
    def aktive_correct(self):
        if self.space.get()==1:
            self.space_word__disabled.set(1)
        else:
            self.space_word__disabled.set(0)
    def on_ketab_6(self,event):
        QApplication.exit()
        self.info2('') 
        self.on_button_clicked_2()
        self.internet.set(self.M165) 
        self.dict.set(self.M177) 
        self.sina.set(self.M203) 
        self.books.set(self.M176)
        self.teb.set(self.M241)
        if not  self.sher.get() in [self.M213,self.M227]:    
            self.separate_search=False
            self.spm=False
            self.search_Active=True
            self.internet_aktive=False
            self.dict_aktive=False
            self.book_aktive=True
            self.info2(self.sher.get())   
            try:
                if  self.thread_active==False:
                    self.separate_search=True
                    src_book_dic=self.searching_book_options()
                    if self.spm==True or src_book_dic==True :   
                        self.run_book_search()
            except:
                pass 
        else:
            self.off_ketab()
    def on_ketab_7(self,event):
        QApplication.exit()
        self.info2('') 
        self.on_button_clicked_2()
        self.internet.set(self.M165) 
        self.dict.set(self.M177) 
        self.sina.set(self.M203) 
        self.sher.set(self.M227) 
        self.books.set(self.M176)
        if not  self.teb.get() in [self.M213,self.M241]:    
            self.separate_search=False
            self.spm=False
            self.search_Active=True
            self.internet_aktive=False
            self.dict_aktive=False
            self.book_aktive=True
            try:
                if  self.thread_active==False:
                    self.separate_search=True
                    src_book_dic=self.searching_book_options()
                    if self.spm==True or src_book_dic==True : 
                        self.run_book_search()
            except:
                pass   
        else:
            self.off_ketab()
    def run_book_search(self):
        self.i=0
        threading.Thread(target=self.book_search, args=("", "", "")).start()

    def start(self):
        self.special_chars = ['$', '#', '%', '&', '(', ')', '*', '+', ',', '-', '.', '/', ':', ';','«','»' ,
                              '<', '=', '>', '?', '@', '[', '\\', ']', '^', '_', '`', '{', '|', '}','،']
        self.url_downloads='url_downloads.json'
        self.sin = ['~','!','@','#','$','%','^','&','*','(',')','_','-','=','.','/','','+','<',
                    '>','{','}','?','؟','|','"',"'",
                    ':',';',',','حح"', '"حح', 'حححح',]           
        self.invalid_languages = {'', 'Exception', 'No features in text', 'id', 'ur', 'ch', 'af', 'sl', 'se', 'sr', 'sk', 'su',
                                'hy', 'as', 'av', 'ay', 'bn', 'bg', 'ch', 'cv', 'cr', 'cs', 'dv', 'so', 't1', 'ca','lv','lt','pl',
                                'et', 'ee', 'no', 'ro', 'fy', 'gu', 'ha', 'kn', 'kk', 'km', 'fi','sv','da','hr','t1','tl',
                                'kj', 'ko', 'ms', 'nn', 'uk', 'nb', 'pa', 'rn', 'tkinter', 'xh','uz','bo','cy','vi','ro','sw'}
        self.time= 0.23
        self.info3_console.config(state=tkinter.NORMAL)
        self.info3_console.insert(tkinter.END, self.M126)
        self.info3_console.config(state=tkinter.DISABLED)
        self.MainWindow()
        self.create_context_menu(self.input_console)
        self.create_context_menu(self.output_console)
        self.create_context_menu(self.notif_console)
        self.patterns = [
            r'\b[a-zA-Z0-9]+\b|\W'
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email
            r'\b\d{5}(?:-\d{4})?\b',  # Zip code
            r'\b\d{1,16}\b',  # ID
            r'\b\d+\b',  # ID
            r'\b00\d{2}\d{7,12}\b',  
            r'\+\d{2}\d{7,12}\b',  # Phone number starting 
            r'\b0\d{2}\d{7,12}\b',  # Phone number starting with '0' 
            r'\b00\d{2}\d{7,10}\b',
            r'\b0\d{2}\d{7,10}\b',
            r'^[a-zA-Z]',
            r'([a-zA-Z]+)',
            r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b'  # IP address
        ]
        self.email = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email
    ]
        self.patterns_latin = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email
        r'\b\d{5}(?:-\d{4})?\b',  # Zip code
        r'\b\d{1,16}\b',  # ID
        r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b',  # IP address
        r'\b\d{3}[-.\s]??\d{3}[-.\s]??\d{4}\b',  # Phone number
        r'\+?\d{1,4}?[-.\s]?\(?(?:\d{1,3}?\)?[-.\s]?)?\d{1,4}[-.\s]?\d{1,9}',
        r'\b\d{1,16}\b',  # ID
        r'\b00\d{2}\d{7,12}\b',  # Phone number starting with '00' followed by 2 digits and 7 to 12 more digits
        r'\b0\d{2}\d{7,12}\b',  # Phone number starting with '0' followed by 2 digits and 7 to 12 more digits
        r'\b00\d{2}\d{7,10}\b',
        r'\b0\d{2}\d{7,10}\b']
        self.start_with=['ب', 'ن', 'می', 'نمی', 'بی', 'نا','ا','ال','الا','ول','پر','داد','آ','جا','محمد','میر',
                            'با','بال','و','هم','این','به','نی','ل','خو','نیک','پاک','وا','بر','دار','هر','نم','چون',
                            'در','سر','م','پیش','پس','ان','علی','خان','خواجه','نی','نیا','بار','خوش','بد','الی','فر',
                            'یک','دوی','سی','چهار','پان','پنج','چهل','پنجاه','شش','شصت','هفت','هفتاد','زشت','والا',
                            'هشت','هشتاد','نه','نود','هزار','میلیون','ملیارد','ترلیون','بهر','بهره','زیبا','من','ز',
                            ]
        self.start_with.sort(key=len, reverse=True)
        self.start_with=set(self.start_with)
        
        self.suffixes = [ 'یمایشان', 'هایمان', 'هایشان', 'یمایش', 'ستان', 'طور', 'طوری', 'گاری', 'گذار',
                            'ستانی','خو','سر','وند',  'انی', 'یمان', 'هایم', 'هایش', 'های', 'آلات','ریز', 'که',
                            'علی','خان','خواجه','گی','بیک','بیگ','اژ','خور','ار','زاد','راد','سیرت','بار','فر',
                            'یک','دوی','سی','چهار','پان','پنج','چهل','پنجاه','شش','شصت','هفت','سرشت','یار',
                            'هفتاد','هشت','هشتاد','نه','نود','هزار','میلیون','ملیارد','ترلیون','شاه','رت',
                            'دار', 'زا', 'هایت', 'گار', 'مین','بر','فت','افت','فند','تان', 'سرا', 'های','داد', 
                            'یند', 'نامه', 'آموزی', 'آموز', 'وار', 'کار', 'مند', 'گرا','ست','یافت','پیش',
                            'پس''داری', 'گیری', 'آور', 'ستان', 'گری', 'گاه', 'بین', 'زاده', 'واری', 'منش',
                            'یان','ییان','دا','بود','خوار', 'آوازه', 'بند', 'بندی', 'نواز', 'انه', 'پذیر', 
                            'ترین', 'پسین', 'یه', 'چه','ک','وه','مان','در','ات', 'یم', 'گر', 'یت', 'یش', 
                            'بندی', 'بند', 'ان', 'ای','نامه', 'بان', 'بانی','ند','اند','شان','سیرت','ین',
                            'ید', 'نی', 'می','تر' ,'ها', 'ی', 'م', 'ت', 'ه', 'د', 'ا','ش','را','انداز',
                        ]
        self.persian_conjunctions = [
        # حروف ربط ساده
        "و", "یا", "پس", "اگر", "نه", "را", "چون", "چه", "تا", "اما", "باری", "خواه", "زیرا", "که", "لیکن", "نیز", "ولی", "هم",
        # حروف ربط مرکب
        "بالعکس", "ولو", "به جز", "سپس", "از این گذشته", "همچنین", "چون‌ که", "چندان ‌که", "زیرا که", "همان‌ که", "بلکه", "چنانچه", "تا این که", "تا آن که", "آنگاه که", "از آن‌جا که", "از این ‌رو", "از بس", "از بهر آن که", "اکنون که", "اگرچه", "اگر", "مگر این که", "با این حال", "با این‌ که", "با وجود این", "بس که", "به شرط آن که",
        # حروف ربط گسسته ربطی
        "خواه", "چه", "یا", "نه", "هم"
    ]     
        self.suffixes.sort(key=len, reverse=True)
        self.suffixes=set(self.suffixes)
        self.char_groups = [ ['ز', 'ظ'], ['ز', 'ض'], ['ذ', 'ز'], ['ض', 'ظ'], ['ظ', 'ض'], ['ذ', 'ظ'],
                                ['ذ', 'ض'], ['ز', 'ذ'], ['ظ', 'ز'], ['ض', 'ز'], ['ط', 'ت'], ['ت', 'ط'],
                                ['ر', 'ز'], ['ز', 'ر'], ['د', 'ذ'], ['ذ', 'د'], ['ح', 'ه'], ['ه', 'ح'],
                                    ['ج', 'چ'], ['چ', 'ج'], ['ح', 'خ'], ['خ', 'ح'], ['ح', 'خ'], ['ح', 'ج'], 
                                    ['ج', 'ح'], ['ص', 'س'], ['س', 'ص'], ['س', 'ث'], ['ص', 'ث'], ['س', 'ص'], 
                                    ['ش', 'س'], ['ص', 'ث'], ['ص', 'س'], ['ث', 'ص'], ['ث', 'س'], ['ص', 'س'], 
                                    ['س', 'ش'], ['ع', 'ا'], ['ا', 'ع'], ['ا', 'ع'],['ا', 'ع'],
                                    ['ق', 'غ'], ['غ', 'ق'], ['ق', 'ف'], ['ف', 'ق'],
                                    ['ک', 'گ'], ['یی', 'ت'], ['یی', 'ی'], ['یی', ''], ['ک', 'گ'], ['گ', 'ک'], ['خوا', 'خا'], ['خا', 'خوا'] ]
        self.chars = set(['','آ','ا','ا','ب', 'پ', 'ت', 'ث', 'ج', 'چ', 'ح', 'خ', 'د', 'ذ', 'ر',
                        'ز', 'ژ', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ک', 'گ', 
                        'ل', 'م', 'ن', 'و', 'ه', 'ی'])
        self.single=set(['ب', 'پ', 'ت', 'ث', 'ج', 'چ', 'ح', 'خ', 'د', 'ذ', 'ر', 'ز', 'ژ', 'س', 'ش', 'ص', 'ض', 'ط', 'ظ', 'ع', 'غ', 'ف', 'ق', 'ک', 'گ', 'ل', 'م', 'ن','ی'])
        self.on_off()
    def on_off(self):
        self.filter = ['Download', 'Herunterladen', 'تحميل', 'Télécharger','اکستروژن','فلز','game','بازی','Punch-tera','موسیقی','Crush',
                        '&quot;','Descargar', 'Загружать', 'Scaricare', 'İndirmek', 'Baixar','دانلود','کامپیوتری','Britney' ,'Bastard','Punch']
       # self.argpattern = '|'.join(self.filter)
        self.dir_path = os.path.join(os.path.expanduser('~'), '.local', 'cache', 'argos-translate', 'downloads')
        self.base_path = os.path.join(os.path.expanduser('~'), '.local', 'share', 'argos-translate', 'packages')
        self.destination_path = os.path.join(os.path.expanduser('~'), '.local', 'cache', 'argos-translate')
        self.source_path = os.path.join(self.dir, 'index.json')
        self.file_name = os.path.basename(self.source_path)
        self.destination_file_path = os.path.join(self.destination_path, self.file_name)
        self.color_code = (0, 0, 0)
        self.search_opt='select'
        self.file_content = {}
        self.reerror_pakages = False
        self.reinstalled = False
        self.thread_active=False
        self.pack_install = False
        self.argose_err = False
        self.pack_downloaded=False
        self.rtl_mode = False
        self.packerror=False
        self.er=False
        self.last_detected_languages ="en"
        self.tra=False
        self.reerror_pakages=False
        self.reinstalled=False
        self.new_word=False  
        self.dfiscancel=False
        self.rev=False
        self.export_docx_=False
        self.sp1=False
        self.sp2=False
        self.sp3=False
        self.sp4=False
        self.S5=False
        self.S6=False
        self.S7=False
        self.S8=False
        self.S9=False
        self.S10=False
        self.S11=False
        self.S12=False
        self.S13=False
        self.S14=False
        self.S15=False
        self.S16=False
        self.S17=False
        self.S18=False
        self.S19=False
        self.S20=False
        self.S21=False
        self.S22=False
        self.S23=False
        self.S24=False
        self.S25=False
        self.S26=False
        self.S27=False
        self.S28=False
        self.S29=False
        self.S30=False
        self.S31=False
        self.spm=False
        self.S32=False
        self.S33=False
        self.S34=False
        self.S35=False
        self.S36=False
        self.S37=False
        self.S38=False
        self.S40=False
        self.S41=False
        self.S42=False
        self.S43=False
        self.S44=False
        self.S45=False
        self.S46=False
        self.S47=False
        self.S48=False
        self.S49=False
        self.S50=False
        self.S51=False
        self.S52=False
        self.S53=False
        self.S54=False
        self.S55=False
        self.S56=False
        self.S57=False
        self.S58=False
        self.S59=False
        self.S60=False
        self.S61=False
        self.S62=False
        self.farhang=[]
        self.sp63=False
        self.S63=False
        self.S64=False
        self.S65=False
        self.stops=False
        self.thread_active=False
        self.skip_rtl=False
        self.pdf_convert=False
        self.console=False
        self.skip_save=False
        self.total_para=5
        self.lang_code("")
    def open_3(self): 
        self.update_notification(self.M162)
        with open('mjt.json', 'r', encoding='utf-8') as f:
            self.synonyms_str = set(f.read().splitlines())
            self.S5=True
            return
    def open_4 (self):    
        try:
            self.not_found_file = 'not_found_fa.json'
            txt_per_comp ='Book_fa.json'
            self.replace='replace.json'
            with open(txt_per_comp, 'r', encoding='utf-8') as f:
                self.fa_words =set(f.read().splitlines())
            with open(self.not_found_file, 'r', encoding='utf-8') as f:
                self.fa_w_not_found =f.read().splitlines()
            with open(self.replace, 'r', encoding='utf-8') as f:
                self.replaced_lines =f.read().splitlines()
            with open(self.replace, 'r', encoding='utf-8') as f:
                self.replaced_words=f.read().split()
            self.S6=True      
        except:
            if not os.path.exists('not_found_fa.json'):
                with open('not_found_fa.json', 'w') as f:
                        f.write('')
            if not os.path.exists('found_fa.json'):
                with open('found_fa.json', 'w') as f:
                        f.write('')
    def open_5(self):
        try:
            if self.S7==True and self.S8==True:
                return
            if self.from_code=='ar':
                if self.S8==True:
                   return
                self.update_notification(self.M162)
                self.not_found_file_ar ='not_found_ar.json'
                with open(self.not_found_file_ar, 'r', encoding='utf-8') as f:
                    self.fa_w_not_found_ar=f.read().splitlines() 
                self.fa_w_not_found_ar = list(self.fa_w_not_found_ar)
                self.S8=True
            else:
                if self.S7==True:
                    return
                self.update_notification(self.M162)
                self.not_found_file_en ='not_found_en.json'
                with open(self.not_found_file_en, 'r') as f:
                    self.fa_w_not_found_en =f.read().splitlines() 
                self.fa_w_not_found_en = list(self.fa_w_not_found_en)
                self.S7=True
        except:
                if not os.path.exists('not_found_ar.json'):
                   with open('not_found_ar.json', 'w') as f:
                        f.write('') 
                if not os.path.exists('not_found_en.json'):
                   with open('not_found_en.json', 'w') as f:
                        f.write('')
    def create_context_menu(self,widget):
        menu = tkinter.Menu(widget, tearoff=0)
        menu.add_command(label='Copy', command=lambda: (root.clipboard_clear(), root.clipboard_append(widget.selection_get())))
        menu.add_command(label='Cut', command=lambda: (root.clipboard_clear(), root.clipboard_append(widget.selection_get()), widget.delete('sel.first', 'sel.last')))
        menu.add_command(label='Paste', command=lambda: widget.insert(tkinter.INSERT, root.clipboard_get()))
        def show_menu(event):
                    menu.post(event.x_root, event.y_root)
        widget.bind('<Button-3>',show_menu)

    def create_console(self, bg_color, fg_color, x, y, width, height):
            console = tkinter.Text(self.root,width=1,tabstyle='wordprocessor',autoseparators=False,exportselection=False,undo=True,startline='', cursor = "xterm",wrap='word', font=self.text_font)
            console.config(bg = "light cyan", fg=fg_color)

            console.place(x=x, y=y, width=width, height=height)            
            return console
    def update_notification(self, message): 
        self.notif_console.config(state=tkinter.NORMAL)
        self.notif_console.delete("1.0", tkinter.END)
        self.notif_console.insert(tkinter.END, message)
        self.notif_console.config(state=tkinter.DISABLED)
    def update_input(self, text):        
        self.input_console.insert(tkinter.END, text)
        self.input_console.see(tkinter.END)
    def update_output(self, text):
        self.output_console.config(state='normal')
        self.output_console.insert(tkinter.END, text)
        self.output_console.see(tkinter.END)
        self.output_console.config(state='disabled')
    def show(self):
        # Remove any previous 'found' tags
        self.input_console.tag_remove('found', '1.0', tkinter.END)
        
        # Get the search query from the user
        search_query = askstring("Find", "Enter your text")
        
        if search_query:
            idx = '1.0'
            while True:
                # Search for the next occurrence of the query
                idx = self.input_console.search(search_query, idx, nocase=True, stopindex=tkinter.END)
                if not idx:
                    break
                
                # Calculate the end index of the found word
                last_idx = f"{idx}+{len(search_query)}c"
                
                # Add the 'found' tag to highlight the word
                self.input_console.tag_add('found', idx, last_idx)
                
                # Scroll to the found word
                self.input_console.see(last_idx)
                
                # Update the index for the next search iteration
                idx = last_idx
        
        # Configure the 'found' tag color
        self.input_console.tag_configure('found', foreground='blue')

    def Lang(self):
        with open('index_lang.json', 'r', encoding='utf-8') as f:
            lang = f.read()
            if lang == 'en':
                    result = messagebox.askquestion('بازنشانی برنامه برای ایجاد تغییر زبان فارسی','برای ایجاد تغییرات زبان برنامه باید بازنشانی شود آیا تایید می فرمایید؟')       
                    if result == 'yes':
                        with open('index_lang.json', 'w', encoding='utf-8') as f:
                                f.write('')  
                        self.restart_program()
            else:
                self.update_notification('  زبان برنامه فارسی است      ' )
    def info(self,text):
        self.info_console.config(state=tkinter.NORMAL)
        self.info_console.delete(1.0, tkinter.END)
        self.info_console.insert(tkinter.END, text)
        self.info_console.config(state=tkinter.DISABLED)
    def info2(self,text):
        self.info2_console.config(state=tkinter.NORMAL)
        self.info2_console.delete(1.0, tkinter.END)
        self.info2_console.insert(tkinter.END, text)
        self.info2_console.config(state=tkinter.DISABLED)
    def info3(self,text):
        self.info3_console.config(state=tkinter.NORMAL)
        self.info3_console.delete(1.0, tkinter.END)
        self.info3_console.insert(tkinter.END, text)
        self.info3_console.config(state=tkinter.DISABLED)
    def lang_fa(self):           
        self.source_language_label.configure(text='زبان مبدا')
        self.target_language_label.configure(text='زبان مقصد')
        self.file_button.configure(text="انتخاب فایل")
        self.translate_button.configure(text= "ویرایش و ترجمه متن")
        self.trans_file_button.configure(text="ویرایش و ترجمه فایل")
        self.gift_button.configure(text='تبدیل PDF')                
        self.export_button.configure(text='ذخیره ')
        self.clear_button.configure(text='پاک کردن')
        self.paste_button.configure(text=' جایگذاری')
        self.help_button.configure(text='راهنما')
        self.color_button.configure(text='بازنشانی')
        self.copy_button.configure(text='کپی')
        self.Qt_translator.configure(text='تایپ فارسی 313')
        self.download.configure(text="توقف پردازش")
        self.M1='فایل'
        self.M2='پاک کردن متن ورودی برای متن جدید'
        self.M3='وارد نمودن فایل با فرمت  تی ایکس تی'
        self.M4='ذخیره متن با فرمت  تی ایکس تی '
        self.M411='ذخیره متن ورودی با فرمت  تی ایکس تی '
        self.M6='خروج از برنامه'
        self.M8='ویرایش'
        self.M9='تغییر جهت راست یا چپ متن '
        self.M10='انتخاب رنگ متن'
        self.M11= 'عدم اصلاح فاصله '
        self.M12='1 تغییر چیدمان حروف'
        self.M13='اصلاح خودکار لغت غير فارسي'
        self.M14='ویرایش متن فارسی '
        self.M15='اصلاح خودکار لغت  '
        self.M16='عدم تبدیل اعداد از انگلیسی به  فارسی '
        self.M17='استفاده از فونت اولیه برای فایل های پی دی اف و ورد'
        self.M18='اضافه نمودن متن اصلی به متن ترجمه شده در فایل های پی دی اف و ورد '
        self.M19='بریدن'
        self.M20='رونوشت'
        self.M21=' جایگذاری'
        self.M22='خطای دانلود'
        self.M23=' خطا '
        self.M24='لطفا صبر کنید تا پردازش قبلی کامل شود'
        self.M25='هیچ فایلی  انتخاب نشده است'
        self.M26='فایل وارد شد'
        self.M27='خطا در انتخاب فایل'
        self.M28='متن از فایل پی دی اف وارد شد'
        self.M29='خطا در پردازش پی دی اف'
        self.M30='استخراج جدول از پی دی اف شروع خواهد شد.'
        self.M31=' خطا در استخراج متن از پی دی اف'
        self.M32='هیچ فایلی  برای ذخیره ایجاد یا انتخاب نشد.'
        self.M33='پایان پردازش'
        self.M34=' حذف فایل'
        self.M35='خطا در تبدیل پردازش پی دی اف'
        self.M36=' پردازش متن  پی دی اف'
        self.M37=' پردازش   پی دی اف'
        self.M38='  در حال پردازش متن'
        self.M39='اتمام پردازش '
        self.M40='  در حال پردازش ورد'
        self.M41='ترجمه و ذخیره فایل'
        self.M42='  پایان پردازش  '
        self.M43=' خطا در  پردازش ورد'
        self.M44='   خطا   ' 
        self.M45='پردازش پی دی اف و استخراج جدول'
        self.M46='لغو پردازش پی دی اف و استخراج جدول'
        self.M47='لغونصب برنامه کمکی استخراج جدول'
        self.M48='در حال پردازش پی دی اف و استخراج جدول'
        self.M49='جدول اسخراج شده '
        self.M50=' اسخراج جدول تکمیل شد'
        self.M51=' فایل  برای ذخیره ایجاد یا انتخاب نشد'
        self.M52='ذخیره خروجی در فایل '
        self.M53='جدول اسخراج شده '
        self.M54='لطفا یک کتاب ار لیست بالا انتخاب فرمایید'
        self.M55='کتاب مورد نظر یافت نشد '
        self.M56='خطا در ترجمه'
        self.M57='در حال ترجمه'
        self.M58='لطفا صبر فرمایید'
        self.M59='خطا در ترجمه یا پردازش متن '
        self.M60=' زبان مورد نظر یافت نشد '
        self.M61='خطا در ترجمه آنلاین لطفا اینترنت را چک یا مترجم آفلاین را امتحان کنید '
        self.M62='خطا در مترجم دوم گوگل'
        self.M63='مترجم ارگوس'
        self.M64='در حال بررسی بسته زبان مورد نظر'
        self.M65='خطا در نرجمه آفلاين'
        self.M66='بسته زبان مورد نظر یافت نشد'
        self.M67='خطا در نصب بسته زبان دوباره دانلود می شود'
        self.M68='خطا در نصب بسته زبان ' 
        self.M69='یک خطا رخ داد  بسته زبان دوباره نصب و برنامه بازنشانی گردد؟  '
        self.M70='دانلود مجدد'
        self.M71=' زبان درخواست شده برای مترجم آفلاین در دسترس نیست'
        self.M72='یک مترجم دیگر را از لیست انتخاب فرمایید '
        self.M73='دانلود  '
        self.M74='لغو گردید لطفا یک مترجم دیگر را از لیست انتخاب فرمایید '
        self.M75='در حال دانلود'
        self.M76='بسته زبان'
        self.M77='دانلود  پایان یافت  '
        self.M78='خطا در دانلود'
        self.M79='لفا دوباره سعی فرمایید'
        self.M80='2 مترجم گوگل'
        self.M81='خطای مترجم مای مموری'
        self.M82='خطای مترجم دیپ'
        self.M83='خطای مترجم یاندکس'
        self.M84='اصلاح لغت'
        self.M85=' فایل  اجرا شود ؟'
        self.M86=' اطلاع رسانی'
        self.M87='نیاز به بازنشانی سیستم است تایید می فرمایید ؟'
        self.M88='برای استخراج جدول از فایل پی دی اف نرم افزار گوییست اسکریپ باید نصب و سسیستم بازنشانی گردد تایید می فرمایید ؟'
        self.M89='نصب بسته زبان تکمیل شد برنامه بازنشانی  می شود' 
        self.M90='بسته زبان مورد نظر یافت نشد بسته  دانلود شود ؟'
        self.M91='در حال آماده سازی فایل لطفا صبر فرمایید:'
        self.M92='نوع فایل انتخاب شده برای تبدیل با قالب مشابه مناسب نیست آیا تمایل دارید فقط متن از فایل استخراج شود ؟'
        self.M93='حذف فایل موقت تبدیل شده از پی دی اف به ورد آفیس  آیا موافق هستید؟'
        self.M94=' حالت ویرایش  بدون ترجمه'
        self.M95='ویرایش متن'
        self.M96='تبدیل فایل و ترجمه'
        self.M97='استخراج جداول از فایل پی دی اف و ترجمه'
        self.M98='استخراج  متن از فایل و ترجمه'
        self.M99='استخراج  متن از فایل   '
        self.M100='انگلیسی'
        self.M101='آلمانی'
        self.M102='فارسی'
        self.M103='چینی'
        self.M104='شناسایی خودکار زبان '
        self.M105='لطفا از تشابه زبان فایل انتخابی با زبان مبدا اطمینان حاصل فرمایید '
        self.M106='ایتالیایی'
        self.M107='هلندی'
        self.M108='فرانسوی'
        self.M109='عربی'
        self.M110='اسپانیولی'
        self.M111='روسی'
        self.M112='ترکی'
        self.M113='پرتقالی'
        self.M114='اندونزی'
        self.M115='هندی'
        self.M116='ژاپنی'
        self.M117='اردو'
        self.M118='مترجم آفلاین'
        self.M119='مترجم گوگل  1' 
        self.M120='مترجم گوگل  2'
        self.M121='مترجم مای مموری'
        self.M122='مترجم دیپ'
        self.M123='مترجم یاندکس'
        self.M124= 'انتخاب مترجم'
        self.M125=' کلمات متضاد و مترادف'
        self.M126='mj.taraz@yahoo.com'
        self.M127=' وارد نمودن متن و اصلاح جملات'
        self.M128='اصلاح نيمه خودکار کلمات '
        self.M129='اصلاح نيمه خودکار کلمات  غير فارسي'        
        self.M130='لغت درست را انتخاب کنيد'
        self.M131='لغت صحيح را انتخاب و يا در پنجره زير بنويسيد '
        self.M132='تاييد '
        self.M133="اصلاح لغت "
        self.M134="ذخيره لغت جديد در مرجع لغت !!!"
        self.M137="از اين به بعد جزو لغت صحيح باشد ؟" 
        self.M135="ادامه اصلاح لغت بصورت نيمه خودکار براي جملات بعدي"
        self.M136='عدم نمایش این پنجره '  
        self.M138="تایید"
        self.M139= "بنویسید"
        self.M140= "زبان مبدا و مقصد يکي است  عدم ترجمه  (حالت ويرايش متن )"
        self.M141='خطا در تبدیل فایل به پی دی اف اگر آفیس نصب نیست آن را نصب فرمایید'   
        self.M142='عدم ذخیره خودکار کلمه جایگزین'
        self.M143='مترجم  313 راه اندازی نشد'
        self.M144="لغو ذخیره فایل ترجمه شده؟"
        self.M145='در حال ذخیره فایل'
        self.M146="2 تغییر چیدمان حروف"
        self.M150="رنگ متن"
        self.M151="اصلاح لغت"
        self.M152="3  تغییر چیدمان حروف"
        self.M153=" عدم تغییر چیدمان متن و حروف"
        self.M154=" فرهنگ لغت فارسی"
        self.M155="براي اصلاح لغت غير فارسي بايد يک زبان غير فارسي را انتخاب نماييد"
        self.M156=" زبان مبدا  با زبان دیکشنری یا کتاب  یکی نیست "
        self.M157="متاسفانه کلمات مترادف فقط برای فارسی موجود است"
        self.M158= 'لغت یا متن مورد نظر  را بنویسید و  یا از متن زیر کپی  کنید'
        self.M159= 'فرهنگ لغت ابجد به فارسی'
        self.M160= 'فرهنگ لغت عربی به فارسی'
        self.M162= "بارگذاری فایل از شکیبایی شما سپاسگزاریم "
        self.M163= "گوگل"
        self.M164= "ویکی پدیا"
        self.M165= 'جستجو و دانلود از اینترنت'
        self.M166= 'فرهنگ لغت انگلیسی به فارسی'
        self.M167= 'فرهنگ لغت فارسی به انگلیسی'
        self.M168= 'فرهنگ لغت فارسی به عربی'
        self.M169= 'فرهنگ لغت عربی'
        self.M170= " ترجمه قرآن"
        self.M171= ' حافظ'
        self.M172= ' قرآن'
        self.M173= "فرهنگ لغت انگلیسی"
        self.M174= " نهج البلاغه"
        self.M175= " نهج الفصاحه"
        self.M176= " جستجو در "
        self.M177= "فرهنگ لغت"
        self.M178= "جستجوی دقیق"
        self.M179= "صحیفه سجادیه"
        self.M180="جستجوی جمله"
        self.M181="فرهنگ لغت دهخدا"
        self.M182="4  تغییر چیدمان حروف"
        self.M183="بحارالانوار"
        self.M184="قانون در طب"
        self.M185="طب سنتی"
        self.M186="علم رجال"
        self.M187="آیین دادرسی"
        self.M188="فرهنگ لغت آلمانی به فارسی"
        self.M189="شاهنامه"
        self.M190="اصول کافی"
        self.M191="دیوان سعدی"
        self.M192="قطره"
        self.M193="قطره فارسی"
        self.M194="قرآن با اعراب"
        self.M195="صفحه جستجو"
        self.M196="تغییر چیدمان کلمات متن در فایل خروجی"
        self.M197=" هیچ متنی پیدا نشد یا متن ورودی خالی است"
        self.M198="جستجو"
        self.M199="قانون طب ابوعلی سینا"
        self.M200=" آشنایی با طب سینا"
        self.M201="ترانسلیت کام"
        self.M202=" داروهای سنتی "
        self.M203="کتب مهندسی"
        self.M204="درمان بیماری داخلی"
        self.M205="درمان بیماری  اندام"
        self.M207="محاسبه ابجد"
        self.M208="المیزان"
        self.M209="مفردات راغب"
        self.M210="مفردات راغب با ترجمه"
        self.M211="ترجمه المیزان"
        self.M212="گزینه جستجو"
        self.M213="غیر فعال"
        self.M214="نمایش صفحه جستجو جداگانه برای "
        self.M215="فایل با فرمت پی دی اف ذخیره نشد اگر آفیس نصب نیست آن را نصب و دوباره سعی کنید در حال حاضر فایل با فرمت نات پد ذخیره می شود اگر متن درست نیست دوباره سعی فرمایید"
        self.M216="SMD  کدهای قطعات الکترونیکی"
        self.M217='استاندارد آلمان DIN بخش یک'
        self.M218='استاندارد آلمان DIN بخش دو'
        self.M219='استاندارد آلمان DIN بخش سه'
        self.M220='استاندارد آلمان DIN بخش چهار'
        self.M221='استاندارد آلمان DIN بخش پنج'
        self.M222='مهندسی مکانیک'
        self.M223='مهندسی برق'
        self.M224='مهندسی الکترونیک'
        self.M225='مهندسی عمران'
        self.M226='مهندسی کامپیوتر'
        self.M227='کتب شعر'
        self.M228='شهریار'
        self.M229='نیما یوشیج'
        self.M230='مولوی'
        self.M231='سهراب سپهری'
        self.M232='خیام'
        self.M233='بابا طاهر'
        self.M234='عنصری بلخی'
        self.M235=' رودکی'
        self.M236='مترجم آفلاین نصب نیست اگر فایل نصب آفلاین را ندارید لطفا برای دریافت مترجم آفلاین ایمیل بزنید mj.taraz@yahoo.com'
        self.M237='این کتاب  نصب نیست اگر فایل نصب  را ندارید لطفا برای دریافت   ایمیل بزنید mj.taraz@yahoo.com'
        self.M238='کلنیک حقوقی ایران'
        self.M239="دانلود شروع شود ؟"
        self.M240= ' لینک دانلود را وارد یا جایکذاری و دکمه دانلود را بزنید \n (هشدار از اعتبار لینک جهت دانلود اطمینان حاصل فرمایید که حاوی ویروس یا تروجان نباشد)'
        self.M241="جسنجو در طب سنتی"
        self.M242="طب الرضا"
        self.M243="رساله"
        self.M244="عیون الرضا"
        self.M245="اصلاح همه کلمات مشابه"
        self.M246="یافتن کلمه صحیح"
        self.M247="لیست کلمات  جدید"
        self.M248=" !  ذخیره"
        self.M249="لیست کلمات جایگزین شده"
        self.M250="کلمه جایگزین پیشنهادی :"
        self.M251="پردازش  پی دی اف برای زبانهای عربی فارسی و اردو درست انجام نمی شود و ممکن است متن درست استخراج نشود آیا ادامه می دهید ؟"   
        self.M252="بارگذاری کامل شد"
        self.M253="بازنشانی نرم افزار" 
        self.M254="بایدو"
        self.M255="جستجو در متن ورودی" 
        self.M256="در حال استخراج جدول از پی دی اف"
        self.M257 ="جایگزین"
        self.M258 ="اجرای تغییرات"
    def Lang2(self):
        with open('index_lang.json', 'r') as f:
                lang = f.read()
                if lang != 'en':
                   result= messagebox.askquestion("Program need restart to take effect for  English", "Do you want to restart program for take effect ?")
                if result == 'yes':
                    with open('index_lang.json', 'w', encoding="utf-8") as f:
                            f.write('en')  
                    # self.start1()
                    self.restart_program()    
                else:
                    self.update_notification("Program language already English")
    def lang_en(self):
        self.source_language_label.configure(text="From Language:")
        self.target_language_label.configure(text="To Language:")
        self.file_button.configure(text="Select")
        self.translate_button.configure(text="Translate text")
        self.gift_button.configure(text="PDFconvert")
        self.export_button.configure(text="Save")
        self.clear_button.configure(text="Clear")
        self.paste_button.configure(text="Paste")
        self.help_button.configure(text="Help")
        self.color_button.configure(text="Restart")
        self.copy_button.configure(text="Copy")
        self.Qt_translator.configure(text="Run Mode 313")
        self.download.configure(text="Stop")
        self.trans_file_button.configure(text='Translate File')   
        self.M1="File"
        self.M2="New"
        self.M3="Open Text File"
        self.M4="Save as TXT File"
        self.M411="Save Input Text as TXT File"
        self.M6="Exit"
        self.M7="Reverse Word (2)"
        self.M8="Edit"
        self.M9="Change Text Direction"
        self.M10="Change Text Color"
        self.M11=" Skip Space Correction"
        self.M12="Reverse Word (1)"
        self.M13="Non Farsi Word Correction (Auto) "
        self.M14="Farsi Text Edition "
        self.M15="Farsi Word Correction (Auto) "
        self.M16="Not Farsi Number convert "
        self.M17="Farsi Default Font and Size for "
        self.M18="Add Source Text in Translated Text "
        self.M19="Cut"
        self.M20="Copy"
        self.M21="Paste"
        self.M22="GUI Error"
        self.M23="An Error Occurred"
        self.M24="Please Wait Until the Other Part of the Translator Finishes Work."
        self.M25="No File Selected."
        self.M26="File Insert"
        self.M27="Select File Error"
        self.M28="Text from PDF Inserted to the Entry "
        self.M29="PDF RE Process Error"
        self.M30="Table Extract Enabled from Edit Menu Extracting and Translation Table from PDF Will Start."
        self.M31="PDF Text Error"
        self.M32="No File Selected for Save Translated."
        self.M33="Process is Finished . "
        self.M34="Removed "
        self.M35="PDF Process Error"
        self.M36="PDF re Processing in Text Mode Only ..."
        self.M37="PDF Process "
        self.M38="Document Processing ..."
        self.M39="Document Process Finished..."
        self.M40="Error process_docx"
        self.M41="Translation and  Save File"
        self.M42="Successful!"
        self.M43="Docx rapidfuz.process Error"
        self.M44="Error:"
        self.M45="PDF Process : Extracting Tables..."
        self.M46="PDF Process : Extracting Tables Canceled"
        self.M47="PDF Process : Installation Was Canceled"
        self.M48="Extracting Tables from PDF Please Wait ..."
        self.M49="Table Extracted : "
        self.M50="Tables Extracted Successfully!"
        self.M51="No Save Directory Selected for Translation"
        self.M52="Output File Saved To"
        self.M53="PDF Table Extract:"
        self.M54="Please Select Book from Box on Top"
        self.M55="Book Not Found:"
        self.M56="Translate Error:{str(e)}"
        self.M57="Translation "
        self.M58="Please Wait..."
        self.M59="translate_doc Error : "
        self.M60="Translation Error ! Please  Make Sure Requested Packages Downloaded or Using an Online Translator "
        self.M61="Translation Error !Please Check the Internet Connection or Using Offline Translator "
        self.M62="Google 02:Internet Connection"
        self.M63="Argos Translator"
        self.M64="Checking Language Package if Necessary."
        self.M65="Unknown Error with Argos Offline Translator !."
        self.M66=" The  Language Package Not Found "
        self.M67="Installation Error Will Downloading :"
        self.M68="Installation Language Error " 
        self.M69="Reinstall Language Packages?"
        self.M70="Retry Downloading..."
        self.M71=" Not Available Language Support for Argos:"
        self.M72="Change Translator and Retry "
        self.M73="Download "
        self.M74="Canceled Try Translate with Selecting Other Translators "
        self.M75="Downloading " 
        self.M76="Language Package..."
        self.M77="Download Finished."
        self.M78="Download Error"
        self.M79=".Please  Retry"
        self.M80="Google 2 "
        self.M81="MyMemory  Error"
        self.M82="Deep Error"
        self.M83="Yandex Error"
        self.M84="Corrections"
        self.M85="Would You Like to Open"
        self.M86="Confirmation"
        self.M87="Gs Installed System Must Restart to Take Effect ,  Would You Like to Restart ?"
        self.M88="For Extract Table Ghostscript Must Install in (Default Directory )and Restart System Would You Like to Install Ghostscript ?"
        self.M89="Completed Program will be  Restart to Take Effect" 
        self.M90="Requested Package Not Installed ! Would You Like to  Download Language Package "
        self.M91="Working on File Please Wait :"
        self.M92="Selected PDF File Not Supported or Some Unknown Error for docx2.convert with the Same Structure ,  Would You Like to Extract  Text Only ?"
        self.M93="Remove Temp Converted PDF to Word Office ?!!"
        self.M94="Correcting Without Translation Mode"
        self.M95="Correcting Text"
        self.M96="Convert and Translate"
        self.M97='Table Extract and Translate(For PDF Files Only)'
        self.M98="Text Extract from File  and Translate "
        self.M99=" Text Extract from File "
        self.M100="English"
        self.M101="German"
        self.M102="Farsi"
        self.M103="Chinese"
        self.M104="Auto_Detect_Language"
        self.M105="Please Make Sure the Source Language is the Same as Your PDF File"
        self.M106="Italian"
        self.M107="Dutch"
        self.M108="French"
        self.M109="Arabic"
        self.M110="Spanish"
        self.M111="Russian"
        self.M112="Turkish"
        self.M113="Portuguese"
        self.M114="Indonesian"
        self.M115="Hindi"
        self.M116="Japanese"
        self.M117="Urdu"
        self.M118="Argos"
        self.M119="Google 1"
        self.M120="Google 2"
        self.M121="MyMemory"
        self.M122="Deep Translator"
        self.M123="Yandex"
        self.M124="Select Translator:"
        self.M125="Synonymous Word (Farsi) "
        self.M126="mj.taraz@yahoo.com  \n call +989914604366 "
        self.M127="Entry Text and Correction"
        self.M128="Word Correction(Semi Auto )"
        self.M129="None Farsi Word Correction (Semi Auto)"
        self.M130='Select Correct Word'
        self.M131='Enter or Select Correct Word '
        self.M132="OK"
        self.M133="Word Correction"
        self.M134="Save New Word in Words Reference ! "
        self.M137="From Now This Word Will Be a Correct Word ?  ." 
        self.M135="Continue Word Correction for Next Text "
        self.M136='Do not show this window'
        self.M138="Confirm"
        self.M139="Write Text"
        self.M140= "The Source and Target Language are the Same No Translations (Text Correction Mode)"
        self.M141="Microsoft Office Not Installed for docx2.convert Word Office to PDF File"
        self.M142='َُSkip Auto Save Replaced  Words'
        self.M143='QT_Taraz_313 Not Started'
        self.M144='Really Cancel Save docx2.convert and Translated File ?'
        self.M145='Saving File'
        self.M146="Reverse Word (2)"
        self.M150="Text Color"
        self.M151="Word Correction"
        self.M152="Reverse Word (3)"
        self.M153="No Change Text and Char"
        self.M154="Farsi Dictionary"
        self.M155="For Latin Word Correction Must Select Other Languages Not Farsi"
        self.M156="Please Set Source Language with Dictionary or Book "
        self.M157="Synonymous Words Only Work for Farsi. Select Farsi from Language."    
        self.M158="Write Word or Copy from Text."  
        self.M159= "Abjad to Farsi DICTIONARY."
        self.M160= "Arabic to Farsi DICTIONARY."
        self.M162= "Loading File..."
        self.M163= "Google."
        self.M164= "Wikipedia."
        self.M165= 'SEARCH and Download INTERNET.'
        self.M166= 'English to Farsi Dictionary.'
        self.M167= 'Farsi to English Dictionary.'
        self.M168= 'Farsi to Arabic Dictionary.'
        self.M169= 'Arabic Dictionary.'
        self.M170= "Quran Farsi."
        self.M171= "HAFEZ."
        self.M172= "Quran Arabic."
        self.M173= 'English Dictionary.'
        self.M174= "Nahj_albalaqah."
        self.M175= "Nahj_alfasaha."
        self.M176= "SEARCH IN BOOK."
        self.M177= "DICTIONARY."
        self.M178= "SEARCHING THE SAME WORD."
        self.M179= "Sahife sajjadieh."
        self.M180="SEARCH BY SENTENCE."
        self.M181='Dehkhoda Dictionary.'
        self.M182="Reverse Word (4)."
        self.M183="beharolanvar."
        self.M184="Ghanoon_teb."
        self.M185="Teb_sonati."
        self.M186="Elm-rijal."
        self.M187="Ayin-dadrasi."
        self.M188="German_to_Farsi_Dictionary."
        self.M189="Shahname_Ferdosi."
        self.M190="Osul_KAfi."
        self.M191="saadi."
        self.M192="َAL_Qatrah."
        self.M193="َAL_Qatrah_farsi."
        self.M194="Quran with erab."
        self.M195="SEARCH WINDOW."
        self.M196="Reverse_Text."
        self.M197="Not Found or Text is Empty."
        self.M198="SEARCH."
        self.M199="Qanoun fe teb aboAliSina."
        self.M200="Medical Education."
        self.M201="translatorCom"
        self.M202="Making Medicinal Plants."
        self.M203="Engineering Books."
        self.M204="Diseases and Treatment 1."
        self.M205="Diseases and Treatment 2."
        self.M207="ABJAD Calc."
        self.M208="AL_MIZAN."
        self.M209="AL_Mofradat."
        self.M210="AL_Mofradat_Farsi."
        self.M211="AL_MIZAN_Farsi."
        self.M212="SEARCH OPTIONS."
        self.M213="INACTIVE."
        self.M214="SHOW SEARCH WINDOW IN SEPARATE WINDOW?  "
        self.M215="File Not Saved as PDF. If Microsoft Office Not Installed Please Install and Try Again. In a Moment File Will Save as Notepad. If Format is Wrong Try Again."
        self.M216="SMD Electronics Code."
        self.M217='DIN STANDARD PART 1.'
        self.M218='DIN STANDARD PART 2.'
        self.M219='DIN STANDARD PART 3.'
        self.M220='DIN STANDARD PART 4.'
        self.M221='DIN STANDARD PART 5.'
        self.M222='MECHANICAL ENG.'
        self.M223='ELECTRICAL ENG.'
        self.M224='ELECTRONIC ENG.'
        self.M225='CONSTRUCTION ENG.'
        self.M226='COMPUTER ENG.'
        self.M227='Poetry Books.'
        self.M228='Shahriar.'
        self.M229='Nima youshij.'
        self.M230='Molana.'
        self.M231='Sohrab sepehri.'
        self.M232='Khayyam.'
        self.M233='Baba Taher.'
        self.M234='Onsori balkhi.'
        self.M235='Roudaki.'
        self.M236='Offline Translator is Not Installed. If You Have No Install File Please Send Email to mj.taraz@yahoo.com to Receive Offline Translator. Thank You.'
        self.M237='This Book is Not Installed. If You Have No Install File Please Send Email to mj.taraz@yahoo.com to Receive Offline Translator. Thank You.'
        self.M238='IranClinicLaw.'
        self.M239="Would You Like to Download?"
        self.M240='Please Insert Download Link or Paste from Clipboard and Select Download Button. \n (Note: Please Ensure Link is Safe and There is No Virus.)'
        self.M241="Search in Teb."
        self.M242="Teb Reza."
        self.M243="Resaleh."
        self.M244="Oyoun Areza."
        self.M245="Replace All."
        self.M246="Results Corrected Word May:"
        self.M247="New Words List."
        self.M248="Save!"
        self.M249="Saved Replaced Words List."
        self.M250="Order Correct Word is:" 
        self.M251="PDF Process for Source Language Not Supported. It Can Be Incorrect Extracting Text. Would You Like Continue?"
        self.M252="Loading complete"
        self.M253="Restarting" 
        self.M254="baidu"
        self.M255="search in input text" 
        self.M256="Table extract enabled from edit menue Extracting and translation table from PDF will start"
        self.M257="Replace"
        self.M258="Make changes "
    def choose_color(self):
        self.color_code = colorchooser.askcolor(title ="Choose color")
        self.input_console.config(fg=self.color_code[1])
        self.output_console.config(fg=self.color_code[1]) 
    def update_font(self, *args):   
           size = self.size_box.get()
           self.new_font = (self.font_box.cget("text"), size)           
           self.output_console.config(font=self.new_font)
           self.input_console.config(font=self.new_font)
    def new_file(self):
        self.input_console.delete(1.0, tkinter.END)
        self.current_file = None
    def open_file(self):
        try:
            options = {
                'filetypes': [('Text files', '*.txt')],
                'initialdir': os.getcwd(),
                'title': 'Select Text file'
            }
            self.file = filedialog.askopenfilename(**options)
            if self.file:
                with open(self.file, "r", encoding="utf-8") as file:
                    content = file.read()                
                    self.update_input( content)
                    self.current_file = self.file
                self.input_console.tag_configure("rtl")
        except :
            pass
    def save_file(self):
        text_to_save = self.output_console.get(1.0, tkinter.END)
        self.file_p = filedialog.asksaveasfilename(
            title="Save File",
            filetypes=[("Text files", "*.txt")]
        )
        if self.file_p:
            try:
                with open(self.file_p, "w", encoding="utf-8") as file:
                    file.write(text_to_save)
                self.current_file = self.file_p
                tkinter.messagebox.showinfo(self.M42, "File saved successfully.")
                os.startfile(self.file_p)
            except Exception as e:
                tkinter.messagebox.showerror("Error", f"Saving the file error: {str(e)}")
    def save_file_in(self):
        text_to_save = self.input_console.get(1.0, tkinter.END)
        self.file_p = filedialog.asksaveasfilename(
            title="Save File",
            filetypes=[("Text files", "*.txt")]
        )
        if self.file_p:
            try:
                with open(self.file_p, "w", encoding="utf-8") as file:
                    file.write(text_to_save)
                self.current_file = self.file_p
                tkinter.messagebox.showinfo(self.M42, "File saved successfully.")
                os.startfile(self.file_p)
            except Exception as e:
                tkinter.messagebox.showerror("Error", f"Saving the file error: {str(e)}")
    def cut_text(self):
        try:
            selected_text = self.input_console.get(1.0, tkinter.END)        
            self.root.clipboard_clear()
            self.root.clipboard_append(selected_text)
            self.input_console.delete(1.0,tkinter.INSERT)
        except Exception as e:
            pass
    def copy_text(self):
        try:
            selected_text = self.input_console.get(1.0, tkinter.END)
            self.root.clipboard_clear()
            self.root.clipboard_append(selected_text)
        except:
            pass
    def paste_text(self):
        if self.thread_active==False:
            try:
                clipboard_text = self.root.clipboard_get()
                self.input_console.insert(tkinter.INSERT, clipboard_text)
            except :
                pass
        else:self.update_notification(self.M24)
    def run(self):
        self.root.mainloop()
    def MainWindow(self):
        from tkinter import PhotoImage
        try :
            self.root.resizable(width = 0, height = 0)
            icon_path = os.path.join(self.dir, 'icon.png')
            icon = PhotoImage(file=icon_path)
            self.root.iconphoto(False, icon)
        except Exception as e:
            self.update_notification(f"  Icon path: {str(e)}")
    def show_help_message(self): 
        if  self.thread_active==False:
            try:
                text_path = os.path.join(self.dir, 'Help.json')
                text_path2 = os.path.join(self.dir, 'Help2.json')
                with open(text_path, 'r', encoding='utf-8') as file:
                    text_content = file.read()
                with open(text_path2, 'r', encoding='utf-8') as file:
                    text_content2 = file.read()   
                self.update_input( text_content)
                messagebox.showinfo("Help", text_content2)
                os.startfile('help.pdf') 
            except :
                pass
    def copy_to_clipboard(self):
        try:
            clipboard_text = self.output_console.get(1.0, tkinter.END)
            self.root.clipboard_clear()
            self.root.clipboard_append(clipboard_text)
            self.root.update()
            self.update_notification(" متن کپی شد")
        except :
            self.update_notification("کپی نشد")
    def paste_from_clipboard(self):
        try:
            if self.thread_active==False:
                self.input_console.delete(1.0, tkinter.END)
                clipboard_text = self.root.clipboard_get()
                self.input_console.insert(tkinter.INSERT, clipboard_text)
                self.root.update()
            else:self.update_notification(self.M24)
        except :
            self.update_notification(f"  هیچ متنی کپی نشده")   

        

    def select_file(self):
        if self.thread_active==True:
           self.update_notification(self.M24)
           return
        self.pdf = False
        try:
            filetypes = [
                ('Supported Files', '*.pdf *.docx *.xlsx'),
                ('PDF files', '*.pdf'),
                ('Word files', '*.docx'),
                ('Excel files', '*.xlsx')
            ]
            self.file_path = filedialog.askopenfilename(filetypes=filetypes)
            if not self.file_path:
                self.update_notification(self.M25)
                self.process_finish()
                return
            if not os.path.exists(self.file_path):
                self.process_finish()
                messagebox.showinfo("Error", f"The selected file does not exist: {self.file_path}")
                return
            file_name = os.path.basename(self.file_path)
            file_size = os.path.getsize(self.file_path) / 1000024  # Convert to KB
            self.info3(f"{ file_name }  '\n'   {file_size:.2f} MB")
            if self.thread_active ==False:
                if self.file_path:
                    self.thread_active=True
                    self.tranc_err=False 
                    self.thrad=threading.Thread(target=self.run_document).start() 
                else:
                    self.update_notification(self.M25)
            else:
                self.update_notification(self.M24)
        except Exception as e:
            self.process_finish()
            self.update_notification(f"  {self.M27}   : {str(e)}")
    def trans_file(self):
        self.stops=False 
        if self.thread_active ==False:
            if self.file_path:
                self.thread_active=True
                self.tranc_err=False 
                threading.Thread(target=self.run_document).start() 
            else:
                self.update_notification(self.M25)
        else:
            self.update_notification(self.M24)
    def run_document(self):
        self.stops=False 
        try:
            self.clear()
            self.thread_active=True
            if self.file_path.endswith(".pdf") or self.pdf_convert==True:
                if self.from_code in ['fa','ar','ur']:
                    confirmation = messagebox.askyesno(self.M86,f"{self.M251} {self.source}")                
                    if not confirmation:
                        self.stop()
                        return
                self.pdf = True
                self.process_pdf_file()
            elif self.file_path.endswith(".docx"):
                self.docx=True
                self.skip_rtl = True
                self.process_docx_file()
            elif self.file_path.endswith(".xlsx"):
                self.skip_rtl = True
                self.docx=True
                self.process_xlsx_file()
        except:
            self.update_notification(self.M25)
    def process_pdf_file(self):
        self.stops=False 
        try:
           
            if self.document_mode_box.cget('text') in [self.M96, self.M97]  or self.pdf_convert==True:
               self.process_pdf2()
            elif self.document_mode_box.cget('text') ==self.M98:
                self.skip_rtl = True
                self.PDF_Text_RTL()
            else:
                self.skip_rtl = True
                self.process_pdf()
            
            self.update_notification(self.M42) 
        except Exception as e:
            self.skip_rtl = False
          
            self.update_notification(f"  {self.M44}   : {str(e)}")
    def process_docx_file(self):
        self.thread_active = True
        if self.document_mode_box.cget('text') in [self.M96, self.M97]:
           self.process_docx2()
           self.rev=True
        else:
           self.process_docx()
        
        self.update_notification(self.M42) 
    def process_xlsx_file(self):
             
        self.thread_active = True
        if self.document_mode_box.cget('text') in [self.M96, self.M97]:
           self.process_xlsx2()
        else:
           self.process_xlsx()
        
        self.update_notification(self.M42) 
    def pdf_converter(self):
        try:
            self.pdf_convert=True
            if self.from_code in ['fa','ar','ur']:
                confirmation = messagebox.askyesno(self.M86,f"{self.M251} {self.source}")                
                if not confirmation:
                    self.stop()
                    return
            self.select_file()          
        except Exception as e:
            self.update_notification(f"  {self.M44}  : {e}")
    def process_pdf(self):
        from pypdf import PdfReader
        self.pdf=True
        try:
            self.update_notification(self.M37)
            if self.document_mode_box.cget('text')==self.M97 :
               self.update_notification(self.M30)
               self.extract_tables()
               return
            else:
                reader = PdfReader(self.file_path)
                full_text = ""
                for page in reader.pages:                     
                    full_text += page.extract_text() + "\n"
                    if self.stops==True:                        
                        break 
                    self.update_input(full_text+'\n')
                if self.pdf_convert==True:
                    self.full_text=full_text
                    self.pdf_convert=True
                    self.thread_active = False
                    self.export_docx()
                    return               
                self.process_finish()                               
        except Exception as e:
                self.update_notification(f"  {self.M31}: {str(e)}")
                self.process_finish()    
    def process_pdf2(self):
        self.pdf==True
        if self.document_mode_box.cget('text')==self.M97 :
           self.update_notification(self.M256)
           self.extract_tables()
           return
        else:
            try:
                file_info = (f"{self.M91}   {  os.path.basename(self.file_path)}")
                self.update_notification(file_info)
                directory = os.path.dirname(self.file_path)
                file_name = os.path.splitext(os.path.basename(self.file_path))[0]
                docx_file = os.path.join(directory, file_name+'_converted.docx')   
                cv = pdf2.Converter(self.file_path)                
                cv.convert(docx_file, start=0, end=None)                
                cv.close()
                doc = dox.Document(docx_file)
                translated_texts = {}
                processed_runs = set()
                for para in doc.paragraphs: 
                    self.lang_code("")    
                    if self.stops==True:
                        break              
                    if para.text.strip():  
                        para.text1=self.space_correction(para.text)    
                        if  self.from_code in ["fa", "ar","ur"] :   
                            if self.format_box.cget('text') in [self.M146]:
                                para.text1=self.reverse_words(para.text1)
                            elif self.format_box.cget('text') in [self.M153,self.M12,self.M152,self.M182]:                               
                                para.text1=self.format_text(para.text1)                                    
                            if not self.to_code in ["fa", "ar","ur"]:
                                para_properties = para._element.get_or_add_pPr()
                                para.alignment = 0                            
                                para_properties.rtl = False                                        
                        elif not self.from_code in ["fa", "ar","ur"] and  self.to_code in ["fa", "ar","ur"]:
                                para_properties = para._element.get_or_add_pPr()
                                para.alignment = 2                             
                                para_properties.rtl = True
                                para.text1=para.text  
                        else:para.text1=para.text                   
                        self.update_input( para.text1 + '\n')
                        
                        if para.text1 not in translated_texts:         
                            translated_texts[para.text1]=self.translate_text(para.text1)                   
                        translated_text = translated_texts[para.text1]
                        # Flag to check if translation has been applied
                        translation_applied = False                                                 
                        for run in para.runs:
                            if self.stops==True:
                                break  
                            if run in processed_runs:
                                continue  
                            if not 'graphicData'  in run._r.xml:
                                if run.text.strip():
                                    if run.font.size:
                                        font_size = run.font.size.pt
                                        run.font.size = dox.shared.Pt(font_size) 
                                    else:
                                        font_size = 12 
                                    run.font.size = dox.shared.Pt(font_size)                                    
                                    if not self.using_orginal_text_enable.get():
                                       run.clear()                                    
                                    if not translation_applied:    
                                        if  self.format_box.cget('text') in [self.M153]:
                                            if  (not self.from_code in ["fa", "ar","ur"] and  self.to_code in ["fa", "ar","ur"]):
                                                run.font.rtl = True                                
                                        run.add_text(translated_text)                                        
                                        translation_applied = True                                    
                        processed_runs.add(run)  # Mark this run as processed
                        if self.using_orginal_text_enable.get():
                           run.clear() 
                docx_file2 = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("PDF Document", "*.docx")])
                if not docx_file2:                                
                    self.process_finish()                           
                    os.remove(docx_file)
                    return 
                self.update_notification(self.M145)                      
                translated_docx_file = os.path.splitext(docx_file2)[0] + '.docx'
                translated_pdf_file = os.path.splitext(docx_file2)[0] + '.pdf'
                doc.save(translated_docx_file) 
                docx2.convert(translated_docx_file,translated_pdf_file)
                file_name_2 = os.path.basename(translated_pdf_file)    
                confirmation = messagebox.askyesno(self.M86,f"{self.M42} {self.M85}  {file_name_2}")                
                if confirmation:
                    os.startfile(translated_docx_file)  
                    os.remove(docx_file)                  
                self.process_finish()
                self.update_notification(f"  {self.M41} {self.M42}") 
            except Exception as e:               
                self.process_finish()                                                 
                messagebox.showinfo(self.M35,f"{self.M44} : {str(e)}")
                self.update_notification( f"{self.M35} : {str(e)}")                                      

    def PDF_Text_RTL(self):
        from pypdf import PdfReader
        self.tranc_err=False
        try:
            doc = dox.Document()
            pdf = PdfReader(self.file_path)
            for page in pdf.pages:
                self.text1 = page.extract_text('') 
                self.text2=self.space_correction(self.text1)
                self.update_input(self.text2 +'\n') 
                self.thread_active = True
                if self.stops==True:                       
                    break 
                translated_text = self.translate_text(self.text2 )
                para = doc.add_paragraph()
                if self.target in [self.M102,self.M109]:
                    para.alignment = dox.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
                #translated_text="\n".join(translated_text)
                para.add_run(translated_text)  # Add translated text to the paragraph
            docx_file2 = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Document", "*.pdf")])           
            if not docx_file2:
                self.process_finish()                                                
                self.update_notification(self.M32)
                return
            translated_docx_file = os.path.splitext(docx_file2)[0] + '_ta.docx'
            translated_pdf_file = os.path.splitext(docx_file2)[0] + '_ta.pdf'
            doc.save(translated_docx_file)  
            docx2.convert(translated_docx_file,translated_pdf_file)
            self.process_finish()           
            confirmation = messagebox.askyesno(self.M86,f"{self.M85} {translated_pdf_file}")                
            if confirmation:
                os.startfile(translated_pdf_file)                     
        except Exception as e:
            self.process_finish()            
            self.update_notification(f"  {self.M37} :{str(e)}")                     
    def process_docx(self ):
        self.update_notification(self.M38)
        self.skip_rtl=True
        try:            
            doc = dox.Document(self.file_path)            
            for para in doc.paragraphs:
                if self.stops==True:                    
                    break 
                para.text2=self.space_correction(para.text)      
              #  para.text3=self.format_text_2(para.text2)
                self.update_input( para.text2 + '\n')
                if self.document_mode_box.cget('text') ==self.M98:
                    self.translate_text(para.text2)
            self.thread_active =False
            if self.document_mode_box.cget('text') ==self.M98: 
                self.export_docx()  
            self.process_finish()                    
            self.update_notification(self.M39)
            return
        except Exception as e:            
            self.update_notification(f"  {self.M40}  {str(e)}")
            self.process_finish()                                    
    def process_docx2(self):       
            self.skip_rtl=True
            try:
                file_info = (f"{self.M91}  {os.path.basename(self.file_path)}")
                self.update_notification(file_info)
                doc = dox.Document(self.file_path)
                translated_texts = {}
                processed_runs = set()
                for para in doc.paragraphs:  
                    if self.stops==True:
                        break                    
                    if para.text.strip():  
                        para.text1=self.space_correction(para.text)    
                        if  self.from_code in ["fa", "ar","ur"] :                                    
                            if not self.to_code in ["fa", "ar","ur"]:
                                para_properties = para._element.get_or_add_pPr()
                                para.alignment = 0                            
                                para_properties.rtl = False 
                                        
                        elif not self.from_code in ["fa", "ar","ur"] and  self.to_code in ["fa", "ar","ur"]:
                                para_properties = para._element.get_or_add_pPr()
                                para.alignment = 2                             
                                para_properties.rtl = True
                                para.text1=para.text                    
                        self.update_input( para.text1 + '\n')
                        
                        if para.text not in translated_texts:                            
                            translated_texts[para.text1] = self.translate_text(para.text1)                            
                        translated_text = translated_texts[para.text1]
                        # Flag to check if translation has been applied
                        translation_applied = False                                                 
                        for run in para.runs:
                            if self.stops==True:
                                break   
                            if run in processed_runs:
                                continue  
                            if not 'graphicData'  in run._r.xml:
                                if run.text.strip():
                                    if run.font.size:
                                        font_size = run.font.size.pt
                                        run.font.size = dox.shared.Pt(font_size) 
                                    else:
                                        font_size = 12 
                                    run.font.size = dox.shared.Pt(font_size)                                    
                                    if not self.using_orginal_text_enable.get():
                                       run.clear()                                    
                                    if not translation_applied:    
                                        if  self.format_box.cget('text') in [self.M153]:
                                            if  (not self.from_code in ["fa", "ar","ur"] and  self.to_code in ["fa", "ar","ur"]):
                                                run.font.rtl = True                                
                                        run.add_text(translated_text)                                        
                                        translation_applied = True                                    
                        processed_runs.add(run)  # Mark this run as processed
                        if self.using_orginal_text_enable.get():
                           run.clear() 
                docx_file2 = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("PDF Document", "*.docx")])
                if not docx_file2:                                
                    self.process_finish()                           
                    return 
                self.update_notification(self.M145)                      
                translated_docx_file = os.path.splitext(docx_file2)[0] + '.docx'
                doc.save(translated_docx_file)   
                confirmation = messagebox.askyesno(self.M86,f"{self.M42} {self.M85}  {translated_docx_file}")                
                if confirmation:
                    os.startfile(translated_docx_file)   
                self.update_notification(f"  {self.M41} {self.M42}")
                self.process_finish()              
            except Exception as e:                        
                self.process_finish()                              
                messagebox.showinfo(self.M43,f"{self.M44} : {str(e)}")
                self.update_notification(f"  {self.M43} : {str(e)}")
    def process_xlsx2(self):
        from openpyxl import load_workbook
        from openpyxl import Workbook
        self.excel=True
        try:
            file_info = (f"{self.M91}  {os.path.basename(self.file_path)}")
            self.update_notification(file_info)
            wb = load_workbook(filename=self.file_path)
            translated_file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Document", "*.xlsx")])
            if not translated_file:
                self.process_finish()
                self.update_notification(self.M32)
                return
            new_wb = Workbook()
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                new_ws = new_wb.create_sheet(title=sheet)
                if self.stops==True:
                    break
                for row in ws.iter_rows(values_only=True):
                    translated_row = []
                    if self.stops==True:
                        break
                    for cell in row:
                        try:
                            if self.stops==True:
                                break
                            if cell is None or len(str(cell)) < 1:
                                break
                            cell = self.space_correction(str(cell))
                            self.thread_active = True 
                            self.update_input(cell+'\n')
                            translated_text = self.translate_text(cell)
                            if not isinstance(cell, float):
                                if isinstance(cell, int):
                                    cell = str(cell)
                        except Exception as e:
                            self.update_notification( f"{self.M44} : {str(e)}")
                            translated_text = cell
                        translated_row.append(translated_text)
                    new_ws.append(translated_row)
                if self.stops==True:
                    break
            messagebox.showinfo({self.M42}, {self.M33})
            self.process_finish()
            new_wb.save(translated_file)
            os.startfile(translated_file)
            return
        except Exception as e:
            self.process_finish()
            messagebox.showinfo(self.M23,f"{self.M44} : {str(e)}")
            self.update_notification( f"{self.M23} : {str(e)}")

    def extract_tables(self):
        self.update_notification(self.M45)
        try:
            import camelot
            from openpyxl import Workbook
            source_path = os.path.join(self.dir, 'gs10012w64.exe')
            if self.find_folders_2()==False:
              if messagebox.askyesno(self.M86,f"{self.M88}"):
                    os.startfile(source_path)
                    #time.sleep(12)
                    if self.find_folders_2()==True:                         
                        if messagebox.askyesno(self.M86,f"{self.M87}"  ):
                            os.system("shutdown /r /t 3")
                    else:
                        self.update_notification(self.M45)
                        self.process_finish()
                        return
        except Exception as e:
            self.update_notification(f"  {self.M47}  {str(e)}")
            self.process_finish()
            return        
        try:
            self.update_notification(self.M48)
            destination =filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel Document", "*.xlsx")])
            if not destination:
                self.update_notification(self.M47)
                self.process_finish()
                return
            tables = camelot.read_pdf(self.file_path, flavor='lattice', pages='all')
            wb = Workbook()
            wb.remove(wb.active)
            for i, table in enumerate(tables):
                if self.stops==True:                       
                    break 
                df = table.df
                ws = wb.create_sheet(title=f'Table {i}')
                for index, row in df.iterrows():
                    ws.append(row.tolist())
                self.update_notification(  f" {self.M49} {i}.")
            wb.save(destination)
            messagebox.showinfo(self.M42, f"{len(tables)} {self.M50} !")
            self.process_finish()
            self.update_notification(self.M50)
            if self.pdf_convert==True:
                self.update_notification(self.M33)
                os.startfile(destination)
                self.process_finish()
                return
            if not destination:
                self.update_notification(self.M51)
                self.process_finish()                                             
                return
            self.update_notification( f" {self.M52} {destination}")
            self.file_path=destination
            self.process_xlsx2()
            self.process_finish()   
               
            self.update_notification(self.M42)                                          
            return
        except Exception as e:
            self.process_finish()                                                               
            messagebox.showinfo("PDF Table extract", f"{self.M23}: {e}")
            self.update_notification( f" {self.M53}: {e}")

    def find_folders_2(self):
        base_path = r"C:\Program Files\gs\gs10.01.2"
        base_path_2 = r"C:\Program Files (x86)\gs\gs10.01.2"
        if os.path.exists(base_path) or os.path.exists(base_path_2):
            return True
        else:
            return False                        
    def process_xlsx(self):
        import openpyxl       
        self.output_console.delete("1.0", tkinter.END)
        try:            
            workbook = openpyxl.load_workbook(self.file_path, data_only=True)  
            for sheet in workbook.sheetnames:
                #self.update_input( f"Sheet: {sheet}\n")
                worksheet = workbook[sheet]
                #row_text)
                row_text=self.space_correction(sheet)
                if self.stops==True:                    
                    break 
                for row in worksheet.iter_rows(values_only=True):
                    row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                   # row_text_formatted=self.format_text(row_text)
                    self.update_input( row_text + '\n')            
            self.process_finish()                                                          
        except Exception as e:            
            self.process_finish()                                      
            self.update_input( f"{self.M23}: {e}\n")
            return
    def export_docx(self):
        if self.thread_active==False:
            self.export_docx_=True
            try:
                doc = dox.Document()
                size = int(self.size_box.get())
                new_font = self.font_box.cget("text")
                if self.color_code[0]:  # If a color is selected
                    rgb_color = tuple(int(self.color_code[0][i]) for i in range(3))  # Get the RGB color
                else:  # If no color is selected, use a default color
                    rgb_color = (0, 0, 0)
                if self.pdf_convert==True:
                    text=self.full_text.split('\n') 
                else:
                    text=self.output_console.get(1.0, tkinter.END).split('\n')                
                if  self.pdf_convert==True or len(text)>2 :
                    self.file_pat = filedialog.asksaveasfilename(
                        title="Export Translated Text",
                        filetypes=[("Word files", "*.docx"), ("PDF files", "*.pdf")],
                        defaultextension=".docx" )             
                    if not self.file_pat:
                        return 
                    for para in text :                                
                        para1=self.space_correction(para)
                        if self.target in [self.M102,self.M109,self.M117]:
                            if self.file_pat.endswith('.pdf'):
                                if (not self.format_box.cget('text') in [self.M146] and self.rev==True) :
                                    para1=self.format_text_2(para1)
                        paragraph = doc.add_paragraph(para1)
                        if self.target in [self.M102, self.M109,self.M117]: 
                            if self.file_pat.endswith('.pdf'):
                                paragraph.alignment = dox.enum.text.WD_PARAGRAPH_ALIGNMENT.RIGHT
                        for run in paragraph.runs:
                            run.font.size = dox.shared.Pt(size)  # Set the font size
                            run.font.name = new_font  # Set the font style
                            run.font.color.rgb =dox.shared.RGBColor(*rgb_color)  # Set the font color                    
                    if self.file_pat.endswith('.pdf'):
                        doc.save('wordFile.docx')
                        docx2.convert('wordFile.docx', self.file_pat)
                        os.remove('wordFile.docx')
                    else:
                        doc.save(self.file_pat)
                    os.startfile(self.file_pat)
                    self.process_finish()  
                    self.rev=False     
                else: self.update_notification(self.M197)                                                  
            except Exception as e:               
                self.process_finish()
                self.rev=False
                messagebox.showinfo(self.M44, f"{str(e)}")
        else:
            self.update_notification(self.M24)
        return
    def Qt_translator_(self):
        try:          
            file_path=os.path.join(self.dir, 'Qt_tr_313')
            os.startfile(file_path)
        except:
            self.update_notification(self.M143)       
    def gift(self,*args):
        try:
            selected_book = self.book_box.cget('text')
            if selected_book =='قرآن':
               file_path=os.path.join(self.dir, 'FarsiQuran_Vista')
               os.startfile(file_path)
               return
            if selected_book =='اصول کافي' :   
               file_path_2=os.path.join(self.dir, 'kafi_j1.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='نهج الفصاحه' :   
               file_path_2=os.path.join(self.dir, '108-fa-nahjol-fasahe.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='نهج البلاغه' :   
               file_path_2=os.path.join(self.dir, '11070-fa-nahjulbalaghe ba tarjeme farsi ravan.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='حافظ' :   
               file_path_2=os.path.join(self.dir, '16294-fa-shogh-mahdi.pdf')
               os.startfile(file_path_2)
               return
            if selected_book =='سه دقیقه در قیامت' :   
               file_path_2=os.path.join(self.dir, '3daghighe-dar-ghiamat.pdf')
               os.startfile(file_path_2)
               return
            else:
               self.update_notification(self.M54)
               return
        except Exception as e:
            self.update_notification(f"   {self.M55}:{e}")       

    def load_book(self, book_name, book_var, Svar):
        if not Svar:
            with open(book_name, 'r', encoding="utf-8") as f:
                book_var = f.read().splitlines()
            Svar = True
        self.ketab_ = book_var
        return True
    def searching_book_options(self):
        try:
            
            book_options = {
                self.M170: ("Quran_fa.json", "Quran_fa", "S12"),
                self.M171: ("Hafez.json", "Hafez", "S13"),
                self.M172: ("Quran.json", "Quran", "S14"),
                self.M168: ("fa_ar_dic.json", "fa_ar_dic", "S10"),
                self.M167: ("en_fa.json", "en_fa", "S9"),
                self.M159: ("ab_dic.json", "ab_dic", "sp3"),
                self.M160: ("ar_fa_dic.json", "ar_fa_dic", "sp4"),
                self.M154: ("fa_dic.json", "fa_dic", "sp2"),
                self.M169: ("ar_dic.json", "ar_dic", "S11"),
                self.M166: ("en_fa_dic.json", "en_fa_dic", "S11"),
                self.M173: ("en_dic.json", "en_dic", "S15"),
                self.M174: ("Nahj_albalaqah.json", "Nahj_albalaqah", "S16"),
                self.M175: ("Nahj_alfasaha.json", "Nahj_alfasaha", "S17"),
                self.M179: ("sahifeh_sajjadieh.json", "sahifeh_sajjadieh", "S19"),
                self.M181: ("dehkhoda.json","dehkhoda","S20"),
                self.M183: ("beharolanvar.json", "beharolanvar", "S21"),
                self.M184: ("Ghanoon_teb.json", "Ghanoon_teb", "S22"),
                self.M185: ("Teb_sonati.json", "Teb_sonati", "S23"),
                self.M186: ("Elm_rijal.json", "Elm_rijal", "S24"),
                self.M187: ("Ayin_dadrasi.json", "Ayin_dadrasi", "S25"),
                self.M188: ("gr_to_Fa_dic.json", "gr_to_Fa_dic", "S26"),
                self.M189: ("Shahname.json", "Shahname", "S27"),
                self.M190: ("Osul_KAfi.json", "Osul_KAfi", "S28"),
                self.M191: ("saadi.json", "saadi", "S29"),
                self.M192: ("qatreh.json", "َAL_Qatrah", "S30"),
                self.M193: ("qatreh_tar.json", "َAL_Qatrah_farsi", "S31"),
                self.M199: ("Qanoun_Teb.json", "P_M_P", "S32"),
                self.M200: ("amoozesh.json", "M_E", "S33"),
                self.M202: ("sakht_daroo.json", "M_M_P", "S34"),
                self.M204: ("bimari_va_ellat_1.json", "D_A_T_1","S35"),
                self.M205: ("bimari_va_ellat_2.json", "D_A_T_2","S36"),
                self.M194: ("Quran_erab.json", "Quran_erab","S37"),
                self.M208: ("almizan.json", "almizan_","S40"),
                self.M209: ("mofradat.json", "mofradat_","S41"),
                self.M210: ("almizan_tr.json", "almizan_tr_","S42"),
                self.M211: ("mofradat_tr.json", "mofradat_tr_","S43"),
                self.M216: ("smd.json", "smd","S44"),
                self.M217: ("din1.json", "din1","S45"),
                self.M218: ("din2.json", "din2","S46"),
                self.M219: ("din3.json", "din3","S47"),
                self.M220: ("din4.json", "din4","S48"),
                self.M221: ("din5.json", "din5","S49"),
                self.M222: ("mechanic.json", "mechanic","S50"),
                self.M223: ("electric.json", "electric","S51"),
                self.M224: ("electronic.json", "electronic","S52"),
                self.M225: ("construction.json", "construction","S53"),
                self.M226: ("computer.json", "computer","S54"),
                self.M228: ("shahriari.json", "shahriar","S55"),
                self.M229: ("nima.json", "parvin","S56"),
                self.M230: ("molavi.json", "molavi","S57"),
                self.M231: ("sohrab.json", "sohrab","S58"),
                self.M232: ("khayyam.json", "khayyam","S59"),
                self.M233: ("babataher.json", "babataher","S60"),
                self.M234: ("onsori.json", "onsori","S61"),
                self.M235: ("roudaki.json", "roudaki","S62"),
                self.M242: ("tebreza.json", "reza","S63"),
                self.M243: ("resaleh.json", "khamenei","S64"),
                self.M244: ("ouonoreza.json", "Oyoun","S65"),
                }              
            if not self.books.get()==self.M176:
                variant =self.books.get()  
            elif not self.dict.get()==self.M177:
                variant =self.dict.get() 
            elif not self.sina.get()==self.M203:
                variant =self.sina.get() 
            elif not self.sher.get()==self.M227:
                variant=self.sher.get()
            elif not self.teb.get()==self.M241:
                variant=self.teb.get()
            else :return False
            for book, (book_name, book_attr, Sattr) in book_options.items():
                if self.stops==True:                   
                    break 
                if book == variant:
                    self.ktb=variant
                    if book_name in ["en_dic.json","en_fa_dic.json","en_fa.json"] and not self.from_code in ["en","fa" ]:
                        messagebox.showinfo(self.M44,self.M156)
                        self.off_ketab()
                        #time.sleep(6)
                        return
                    if book_name not in ["en_dic.json","en_fa_dic.json","en_fa.json","gr_to_Fa_dic.json"] and not  self.from_code in ["ar","fa" ]:
                        messagebox.showinfo(self.M44,self.M156)
                        self.off_ketab()
                        #time.sleep(6)
                        return
                    if book_name in ["gr_to_Fa_dic.json"] and not self.from_code in ["fa","gr"]:
                        messagebox.showinfo(self.M44,self.M156)
                        self.off_ketab()
                        #time.sleep(6)
                        return
                    if not self.from_code in ["fa","ar","en","gr"]:
                        messagebox.showinfo(self.M44,self.M156)
                        self.off_ketab()
                        #time.sleep(6)
                        return
                    self.path_book = book_name
                    self.search_process = variant
                    if getattr(self, Sattr) == True :
                        self.spm=True
                        self.ketab_ = set(getattr(self, book_attr))
                        return True
                    split0001=[self.M184,self.M185,self.M186,self.M187
                               ,self.M190,self.M192,
                               self.M199,self.M242,self.M243,self.M244,
                                self.M174,self.M175,self.M193,self.M205,self.M200,self.M202,self.M204]
                    with open(self.path_book, 'r', encoding="utf-8") as f:
                        if  variant in split0001 :
                            text = f.read()
                            parts = text.split('۰۰۰۱')
                            setattr(self, book_attr, parts)
                        else:
                            setattr(self, book_attr, f.read().splitlines())
                        setattr(self, Sattr, True)
                    self.ketab_ = set(getattr(self, book_attr))
                    return True
            return False          
        except Exception as e:
            self.update_notification(f"   {self.M44}:{str(e)}")
            messagebox.showinfo(self.M44,self.M237)
            return False
    def replace_ascii_digits_with_farsi(self,text):
        words = text.split()
        for i in range(len(words)): 
            words[i] = words[i].replace('0', '۰')
            words[i] = words[i].replace('1', '۱')
            words[i] = words[i].replace('2', '۲')
            words[i] = words[i].replace('3', '۳')
            words[i] = words[i].replace('4', '۴')
            words[i] = words[i].replace('5', '۵')
            words[i] = words[i].replace('6', '۶')
            words[i] = words[i].replace('7', '۷')
            words[i] = words[i].replace('8', '۸')
            words[i] = words[i].replace('9', '۹')
        return ' '.join(words) 
    def text_process(self,text):
        try:
            if  self.correct.get() in [self.M13 , self.M129]:
                if self.from_code in ["fa"]:
                    messagebox.showinfo(self.M44,self.M155)
            if self.jomlesazi.get():
                text=self.jomlesaz(text)   
            if self.from_code in ["fa","ur"] :
                text=self.correction_Farsi_text(text)
            else:
                text=self.correction_english_text(text)  
            if self.to_code in ["fa"]:   
                if not self.perian_num.get():     
                   text=self.replace_ascii_digits_with_farsi(text)
            if self.book_aktive==True:
                src_book_dic=self.searching_book_options()
                if  self.spm==True or src_book_dic==True :  
                    text=self.ketab(text) 
            elif self.internet_aktive==True:
                text=self.ketab(text) 
            elif self.dict_aktive==True:
                if self.dict.get()==self.M125: 
                    if self.from_code in ['fa']:
                       text=self.motaradef(text) 
                    else:
                        messagebox.showinfo(self.M44,self.M157)
                elif self.dict.get()==self.M207:
                     text=self.ketab(text) 
                else:
                    src_book_dic=self.searching_book_options()
                    if self.spm==True or src_book_dic==True :   
                        text=self.ketab(text) 
            return text
        except:
            return text
    def translate_aout(self,event):
        if  self.thread_active==False:
            self.thread_active = True
            self.stops=False 
            self.tranc_err=False 
            threading.Thread(target=self.translate_starter).start()
        else:
            self.update_notification(self.M24) 

    def translate(self):
        if  self.thread_active==False:
            self.thread_active = True
            self.stops=False 
            self.tranc_err=False 
            threading.Thread(target=self.translate_starter).start()
        else:
            self.update_notification(self.M24)  
    def translate_starter(self): 
        self.rev=False
        try: 
            self.output_console.configure(state=tkinter.NORMAL)                
            self.output_console.delete(1.0, tkinter.END)
            input_text=self.input_console.get(1.0, tkinter.END)              
            text_parts = input_text.split('\n')
            text_groups = [[part] for part in text_parts]
            for group in text_groups:
                self.console=True
                group_text = ' '.join(group)
                if self.stops==True:
                    break 
                self.thread_active = True   
                group_text=self.space_correction(group_text)
                self.translate_text(group_text)                           
            self.process_finish()
            self.update_notification(self.M42) 
        except Exception as e: 
            self.process_finish()  
            self.update_notification(f"   {self.M56}:{str(e)}") 
    def translate_text(self,text):
        self.text=text
        if self.source == self.M104 :
            self.detect_language()
        try:
            if  (self.virast.get() or self.from_code==self.to_code) and not self.source == self.M104  :
                self.info(self.M95)
                self.update_notification(self.M140) 
                self.text=self.text_process(self.text)  
                self.update_output(self.text+'\n')
                   
                return self.text+'\n' 
            text=self.text_process(text)
            if not self.text or len(self.text)<2 or  isinstance(self.text, float) or self.text==" "  :
                self.update_output( self.text+'\n')
                 
                return self.text+"\n" 
            self.thread_active = True
            translators = {
                self.M119:self.google1,
                self.M118:self.argos,
                self.M120:self.googl2,
                self.M121:self.mymemory,
                self.M122:self.deep,
                self.M123:self.yandex, 
                self.M201:self.translatorCom,          
            }
            selected_lang=self.source
            if self.source==self.M104:
                selected_lang=self.language_codes_2.get(self.from_code)
            self.update_notification(f"  {self.M57}  {selected_lang}  ...   {self.target}    ")
            translator_name = self.translator_var.get()
            translate_func = translators.get(translator_name)
            if translate_func:    

                translated_text = translate_func()     
                if not translated_text:  
                    translated_text=text  
                self.update_output( translated_text)                              
            return translated_text                            
        except Exception as e:
            
            self.update_notification(f"  {self.M59} : {str(e)}")
            self.process_finish()   
            if self.translator_var.get() == self.M118:
                self.update_notification(self.M60)
            else:
                self.packerror=True
                self.update_notification(self.M61)
    def google1(self):
        try:
            text_parts = self.text.split('\n')
            PARAGRAPH_GROUP_SIZE = 66
            text_groups = [text_parts[i:i + PARAGRAPH_GROUP_SIZE] for i in range(0, len(text_parts), PARAGRAPH_GROUP_SIZE)]
            translated_text = ""
            for paragraph_group in text_groups:              
                #time.sleep(self.time)
                if self.source == self.M104 :
                    self.detect_language()
                group_text = '\n'.join(paragraph_group)                                        
                self.translated_group =Translator1.Translator().translate(group_text, source_language=self.source_language_code,destination_language=self.target_language_code).result
                translated_text += self.translated_group + '\n'                  
            #self.process_finish()                    
            return translated_text
        except Exception as e:
                if self.tranc_err==False:
                      self.tranc_err=True
                      tkinter.messagebox.showerror("Error", f"{self.M62}  {self.M61} : {str(e)}")
                else:
                    self.info2(f"{self.M62}  {self.M61} : {str(e)}")
                      #time.sleep(6)
                self.process_finish()
                return #text + '\n' 
    def argos(self):
        translated_text = ""
        if self.to_code=='zh-CN':
            self.to_code='zh'
        if self.from_code=='zh-CN':
                self.from_code='zh' 
        try:
            paragraphs = self.text.split('\n')
            for paragraph in paragraphs:
                translated_par=Translator2.translate.translate(paragraph, self.from_code, self.to_code)
                for filter in  translated_par.split():
                    if filter in self.filter :
                        if self.translate_pay==False:
                            from translatepy import Translator 
                            self.trans=Translator
                            self.translate_pay=True
                        try:
                            translated_par=self.trans().translate(paragraph, source_language='auto',destination_language=self.target_language_code).result
                        except:
                            translated_par=paragraph
                    break
                translated_text += translated_par + '\n'
            self.argose_err = False
            return translated_text            
        except Exception as e:
            if "NoneType" in str(e) or 'opening' in str(e):
                if self.argose_err == True:
                    self.argose_err = False
                else:
                    self.argose_err = True
                #self.pack_install=False                   
                self.info2(f"{self.M60} : {str(e)}")
                self.process_finish()
                self.argos_pak()
            else:
                if self.argose_err == True:
                    self.argose_err = False
                    tkinter.messagebox.showerror(self.M59, e)
                else:
                    self.info2(f"{self.M63} : {str(e)}")
            return #text + '\n'                                        
    def googl2(self):
        try:
            translated_text=""
            text_parts = self.text.split('\n')
            text_groups = [[part] for part in text_parts]
            for group in text_groups:
                group_text = '\n'.join(group)
                #time.sleep(self.time)
                if group_text:
                    if self.source == self.M104 :
                        self.detect_language()                    
                    translated_group = Translator3.Translator().translate(group_text, src=self.source_language_code,
                                                          dest=self.target_language_code).text
                    translated_text += translated_group + '\n'                                
            return translated_text              
        except Exception as e:
            if self.tranc_err==False:
                self.tranc_err=True
                tkinter.messagebox.showerror("Error", f"{self.M80}  {self.M61} : {str(e)}")
            else:
                self.info2(f"{self.M80}  {self.M61} : {str(e)}")
            self.process_finish()
            return #text + '\n'
    def mymemory(self):
        self.source = self.from_code
        self.target = self.to_code
        language_codes = {
                            "en": 'en-US',
                            "fa": 'fa-IR',
                            "de": 'de-DE',
                            "ar": 'ar-SA',
                            "fr": 'fr-FR',
                            "zh": 'zh-CN',
                            "es": 'es-ES', 
                            "ru": 'ru-RU',
                            "it":  'it-IT',
                            "tr": 'tr-TR',
                            "pt": 'pt-PT',
                            "id":  'id-ID', 
                            "nl":  'nl-NL', 
                            "hi": 'hi-IN',
                            "ja":  'ja-JP',
                            "ur":'ur-PK',   
                        }                     
        self.source_language_code = language_codes[self.source]
        self.target_language_code = language_codes[self.target]
        try:
            translated_text = ""
            text_parts = self.text.split('\n')
            text_groups = [[part] for part in text_parts]
            for group in text_groups:
                group_text = '\n'.join(group)
                if self.source == self.M104 :
                    self.detect_language()
                self.translated_group = Translator5.MyMemoryTranslator(source=self.source_language_code, target=self.target_language_code).translate(group_text)                
                translated_text += self.translated_group + '\n'            
            return translated_text
        except Exception as e:
                if self.tranc_err==False:
                      self.tranc_err=True
                      tkinter.messagebox.showerror("Error", f"{self.M81}  {self.M61} : {str(e)}")
                else:
                    self.info2(f"{self.M81}  {self.M61} : {str(e)}")
                self.process_finish()
                return #text + '\n'   
    def deep(self):
        try:
            translated_text = ""
            PARAGRAPH_GROUP_SIZE = 12
            text_parts = self.text.split('\n')
            text_groups = [text_parts[i:i + PARAGRAPH_GROUP_SIZE] for i in range(0, len(text_parts), PARAGRAPH_GROUP_SIZE)]
            for group in text_groups:
                group_text = '\n'.join(group)
                if self.source == self.M104 :
                    self.detect_language()
                translated_paragraph = Translator5.GoogleTranslator(source=self.source_language_code, target=self.target_language_code).translate(group_text)
                translated_text += translated_paragraph + '\n'           
            return translated_text            
        except :
            if self.tranc_err==False:
                self.tranc_err=True
                tkinter.messagebox.showerror("Error", f"{self.M82}  {self.M61} ")
            else:
                self.info2(f"{self.M82}  {self.M61} ")
            self.process_finish()
            return #text + '\n' 
    def yandex(self):
        try:
            text_parts = self.text.split('\n')
            PARAGRAPH_GROUP_SIZE = 66  # Define your paragraph group size
            text_groups = [text_parts[i:i + PARAGRAPH_GROUP_SIZE] for i in range(0, len(text_parts), PARAGRAPH_GROUP_SIZE)]
            translated_text = ""
            for paragraph_group in text_groups:
                group_text = '\n'.join(paragraph_group)
                if self.source == self.M104 :
                    self.detect_language()
                self.translated_group = Translator1.translators.yandex.YandexTranslate().translate( group_text, source_language=self.source_language_code,
                destination_language=self.target_language_code).result                    
                translated_text += self.translated_group + '\n'               
            return translated_text
        except Exception as e:
            self.process_finish()
            if self.tranc_err==False:
                    self.tranc_err=True
                    tkinter.messagebox.showerror("Error", f"{self.M83}  {self.M61} : {str(e)}")
            else:
                self.info2(f"{self.M83}  {self.M61} : {str(e)}")
            self.process_finish()
            return   
    def translatorCom(self):
        try:
            text_parts = self.text.split('\n')
            PARAGRAPH_GROUP_SIZE = 66  # Define your paragraph group size
            text_groups = [text_parts[i:i + PARAGRAPH_GROUP_SIZE] for i in range(0, len(text_parts), PARAGRAPH_GROUP_SIZE)]
            self.translated_text = ""
            for paragraph_group in text_groups:
                group_text = '\n'.join(paragraph_group)     
                self.translated_group = Translator1.translators.translatecom.TranslateComTranslate().translate( group_text, source_language=self.from_code,
                destination_language=self.to_code).result                    
                self.translated_text += self.translated_group + '\n'               
            return self.translated_text
        except Exception as e:  
            self.update_notification(f"{self.M56}  'translatoCom' : {str(e)}")              
            return  group_text        
    def lang_code(self,event):
        try: 
            self.source = self.source_language_combo.get()
            self.target = self.target_language_combo.get() 
            if self.source == self.M104 :
                self.detect_language()                
                return
            self.from_code = self.language_codes.get(self.source)
            self.source_language_code = self.language_codes[self.source]
            self.to_code = self.language_codes.get(self.target)
            self.target_language_code = self.language_codes[self.target]
            self.info(f"{self.source}   ...   {self.target}    {self.translator_var.get()} ")  
            return
        except Exception as e:
            self.process_finish()                                 
    def detect_language(self):
        self.source = self.source_language_combo.get()
        self.target = self.target_language_combo.get()
        self.from_code = self.last_detected_languages
        self.source_language_code = self.last_detected_languages 
        if not len(self.text)<2:
            try:
                if self.translator_var.get() in [self.M119, self.M120]:
                    self.from_code = 'auto'
                    self.source_language_code = 'auto'
                elif re.search(r'[\u0600-\u06FF]', self.text):
                    if self.last_detected_languages != 'ar':
                        self.last_detected_languages = 'fa'
                    else:
                        self.last_detected_languages = 'ar'
                    self.from_code = self.last_detected_languages
                    self.source_language_code = self.last_detected_languages       
                else:
                    detected_language = langdetec.detect(self.text)  # use langdetect for language detection
                    if len(detected_language)<3 and not detected_language in self.invalid_languages and   detected_language!=None :
                        self.last_detected_languages = detected_language
                        self.from_code = self.last_detected_languages
                        self.source_language_code = self.last_detected_languages
                self.to_code = self.language_codes.get(self.target)
                self.target_language_code = self.language_codes[self.target]
                self.info(f"{self.source}   ...   {self.target}    {self.translator_var.get()} ")  
            except :
                self.last_detected_languages = 'en'
        else:self.last_detected_languages = 'en'
    def argos_pak(self):
        if self.thread_active==False :
            self.thread_active = True
            self.update_notification(self.M64)   
            self.check_and_install_argos_pak()
        else:self.update_notification(self.M24)
    def check_and_install_argos_pak(self):   
        pattern = f"{self.from_code}_en" if self.from_code != "en" else f"{self.from_code}_{self.to_code}"
        pattern_2 = f"en_{self.to_code}" if self.to_code != "en" else f"{self.from_code}_{self.to_code}"
        os.makedirs(self.dir_path, exist_ok=True)
        os.makedirs(self.base_path, exist_ok=True)   
        if not os.path.exists(self.destination_file_path):
            shutil.copy(self.source_path, self.destination_path) 
        if self.find_folders(pattern) is not None and self.find_folders(pattern_2) is not None:
            self.update_notification(self.M65)
            if not self.reerror_pakages or self.reinstalled:
                self.reerror_pakages = True
                return
            else:
                self.reinstalled=True
                if  messagebox.askyesno(self.M65, self.M69):      
                    if self.find_folders(pattern) is not None:
                        paths = self.find_folders(pattern)                        
                        for path in paths:      
                                shutil.rmtree(path)
                    elif self.find_folders(pattern_2) is not None:
                         paths=self.find_folders(pattern_2)
                         for path_2 in paths:                      
                            shutil.rmtree(path_2)        
                    return
        if self.from_code == "en" or self.to_code == "en":
            self.ls = self.from_code
            self.t = self.to_code
        else:
           if not self.find_folders(pattern)!=None:
                  self.ls=self.from_code
                  self.t="en"
           else:
                  self.ls="en"
                  self.t=self.to_code
        available_package = self.get_available_package()
        self.thread_active = True
        self.stops=False 
        download_path=self.down_path()
        if available_package:
           self.download_url = available_package.links[0] 
           try:               
                if download_path:                                         
                        if self.pack_install == True and self.pack_downloaded == True and self.argose_err == True :
                            self.update_notification(f"  {self.M67} {available_package} :{self.M68}")
                            self.pack_install = False
                            try:
                               os.remove(download_path)
                            except:
                                pass 
                            self.download_lang_pack() 
                            return
                        try: 
                            if self.pack_install==False:
                                self.pack_install= True
                                self.extract_zip(download_path, self.base_path)
                                self.pack_install= False
                                self.pack_downloaded = False
                                messagebox.showinfo(self.M42,f"Installation  {available_package} {self.M89}")
                                self.restart_program()
                                return
                        except Exception as e:
                            self.pack_install= False
                            self.download_lang_pack()                    
                            return 
                else:
                    self.download_lang_pack()
                self.process_finish() 
           except Exception as e:
                messagebox.showinfo(self.M44,f"{self.M68} {str(e)}")          
                self.update_notification(f"  {self.M68}: {str(e)} {self.M70}...")
        else:
            self.update_notification(f"   { self.M71} {self.source}  {self.target} { self.M72} ")
    def extract_zip(self,download_path, base_path):
        import zipfile
        with zipfile.ZipFile(download_path, 'r') as zip_ref:
             zip_ref.extractall(base_path)
    def get_available_package(self):
        available_packages =Translator2.package.get_available_packages()
        return next((pkg for pkg in available_packages if pkg.from_code == self.ls and pkg.to_code == self.t), None)            
    def download_lang_pack(self):       
        parsed_url = urllib.urlparse(self.download_url)
        #self.pack_install=False
        if not all([parsed_url.scheme, parsed_url.netloc]):
            self.update_notification(f"  Download Error: Please check internet connection and Retry")
        save_path= self.down_path()
        fromdlang=self.language_codes_2.get(self.ls)
        tomdlang=self.language_codes_2.get(self.t)
        confirmation = messagebox.askyesno(f"{self.M86}", f"{self.M90}     {fromdlang}  ...  {tomdlang} ")
        if not confirmation:
            self.process_finish()
            self.update_notification(f"    {self.M73}   {fromdlang}   ...    {tomdlang}   {self.M74}  ")            
            if not self.dfiscancel==True:
               self.dfiscancel=True
            return       
        self.thread_active = True
        self.stops=False 
        self.pack_downloaded = True
        headers = {}
        downloaded = 0 
        if os.path.exists(save_path):
            downloaded = os.path.getsize(save_path)
            headers['Range'] = f'bytes={downloaded}-'
        try:
            response = requests.get(self.download_url, headers=headers, stream=True)
            total_size = int(response.headers.get('Content-Length', 0)) + downloaded
            if total_size > 200000000:  
                raise ValueError("File too large")
            hash_object = hashlib.sha256()  # Change to desired hash function if needed
            with open(save_path, 'ab') as file:
                for data in response.iter_content(chunk_size=1048576):  
                    if self.stops:
                        break                              
                    file.write(data)
                    hash_object.update(data)
                    downloaded += len(data)
                    total_size_MB = int(total_size / 1048576)
                    percent = (downloaded) / 1048576
                    self.info3(f"  {self.M75}  {self.M76} {fromdlang} {tomdlang} \n {percent:.0f}/{total_size_MB}  Mb")
            if downloaded >= total_size:
                self.check_and_install_argos_pak()
            self.process_finish()
        except requests.exceptions.RequestException as e:
            messagebox.showinfo("download pack Error", f"Internet connection {str(e)}")
            self.pack_downloaded = True
            self.process_finish()
            self.update_notification(f"  Download Error: {str(e)}. Please Retry")
    def down_path(self):
        try:                           
            download_path = os.path.join(
                os.path.expanduser("~"),
                ".local",
                "cache",
                "argos-translate",
                "downloads",
                f"translate-{self.ls}_{self.t}.argosmodel"
            )
            return download_path                        
        except Exception as e:
            self.update_output(f"Down_path: {str(e)}")    
    def find_folders(self, pattern):
        base_path = os.path.join(
            os.path.expanduser("~"),
            ".local",
            "share",
            "argos-translate",
            "packages"
        )
        folders = [] 
        for file_name in os.listdir(base_path):
            full_path = os.path.join(base_path, file_name)
            if os.path.isdir(full_path):
                if pattern in file_name:
                    folders.append(full_path)
        if not folders:
            return None
        return folders
    def find_matching_files(self, directory, pattern):
        try:
            matching_folders = []
            pattern_regex = re.compile(pattern)
            for folder_name in os.listdir(directory):
                if os.path.isdir(os.path.join(directory, folder_name)) and pattern_regex.search(folder_name):
                    matching_folders.append(folder_name)            
            return matching_folders
        except Exception as e:
            self.update_notification(f"  find_matching_folders Error: {str(e)} and directory: {directory}")
    def restart_program(self):
        python = sys.executable
        os.execl(python, python, *sys.argv)   
    def clear(self):
        self.output_console.configure(state=tkinter.NORMAL) 
        self.output_console.delete(1.0, tkinter.END)
        self.input_console.delete(1.0, tkinter.END)
        self.rev=False
        self.output_console.configure(state=tkinter.DISABLED) 
    def correction_araby_text(self, text):
        #text =self.replace_characters(text) 
        if not self.correct.get() in [self.M15 ,self.M128]:
           return text
        self.open_5()
        try:
            corrected_words=self.spellingcheck(text)
            return corrected_words
        except :
            return text
    def spellingcheck(self, text):
        if not text or isinstance(text, float) or text.strip() == "":
            return text + '\n'
        if not self.from_code in ['en','fr','ru','tr','it','es']:
            text=self.spellingcheck_de(text)
            return text
        for pattern in self.patterns_latin:
            if re.search(pattern, text):
                return text + '\n'
        words = text.split()
        corrected_words=[]
        if self.from_code == "en":
            spell = SpellChecke.SpellChecker()
        else:
            spell = SpellChecke.SpellChecker(language=self.from_code)
        if self.correct.get()==self.M129:
            for word in words:

                self.update_notification(f"  {self.M84}  : {word}")
                if (self.from_code in ['ar'] and word in self.df_not_found_ar )or (not self.from_code in ['ar'] and word in self.df_not_found_en)  or len(word)<2:
                    corrected_words.append(word)
                    continue              
                corrected_word=self.latin_semi_auto_correction(word,spell,text)

                if not self.from_code in ['ar']:
                    if not corrected_word in self.not_found_file_en:
                        confirmation = messagebox.askyesno(self.M134, f"{self.M137} ' {corrected_word} '")
                        if confirmation :
                            with open(self.not_found_file_en, 'a') as f:
                                f.write(corrected_word + '\n')
                            with open(self.not_found_file_en, 'r') as f: 
                                self.df_not_found_en =f.read().splitlines()
                else:
                    if not corrected_word in self.not_found_file_ar:
                        confirmation = messagebox.askyesno(self.M134, f"{self.M137} ' {corrected_word} '")
                        if confirmation :
                            with open(self.not_found_file_ar, 'a') as f:
                                f.write(corrected_word + '\n')
                            with open(self.not_found_file_en, 'r') as f: 
                                self.df_not_found_ar =f.read().splitlines()
                corrected_words.append(corrected_word)            
            return " ".join(corrected_words)
        spell = Speller.Speller(self.from_code)
        self.update_notification(f"  {self.M84}  : {text}")
        text=spell(text)  
        return text
    def spellingcheck_de(self, text):
        if not text or isinstance(text, float) or text.strip() == "":
            return text + '\n'
        for pattern in self.patterns_latin:
            if re.search(pattern, text):
                return text + '\n'
        words = text.split()
        corrected_words = []
        if self.from_code == "en":
            spell =SpellChecke.SpellChecker()
        else:
            spell =SpellChecke.SpellChecker(language=self.from_code)
        for word in words:
            corr_word=""
            unormal_word=""
            if self.from_code=="ar":
                correct_words=self.df_not_found_ar
            else:correct_words=self.df_not_found_en
            if word in correct_words  or len(word)<2:
                corrected_words.append(word)
                continue
            if self.correct.get()==self.M129:
                corrected_word=self.latin_semi_auto_correction(word,spell,text)
                if not corrected_word in correct_words:
                    confirmation = messagebox.askyesno(self.M134, f"{self.M137} ' {corrected_word} '")
                    if confirmation :
                        if self.from_code=="ar":
                           with open(self.not_found_file_ar, 'a', encoding="utf-8") as f:
                                f.write(corrected_word + '\n')
                           with open(self.not_found_file_ar, 'r', encoding="utf-8") as f: 
                                self.df_not_found_ar =f.read().splitlines() 
                        else:                 
                            with open(self.not_found_file_en, 'a') as f:
                                    f.write(corrected_word + '\n')
                            with open(self.not_found_file_en, 'r') as f:
                                    self.df_not_found_en =f.read().splitlines()   
            else:
                try:
                    corr_word = spell.correction(word)
                except Exception as e:
                   self.update_notification(f"  {self.M44}  : {e}")
                   unormal_word=word
                if corr_word :
                    corrected_word=corr_word
                    if corrected_word in correct_words:
                       pass
                    else:
                        if self.from_code=="ar":
                           with open(self.not_found_file_ar, 'a', encoding="utf-8") as f:
                                f.write(corrected_word + '\n')
                           with open(self.not_found_file_ar, 'r', encoding="utf-8") as f: 
                                self.df_not_found_ar =f.read().splitlines() 
                        else:  
                            try:               
                                with open(self.not_found_file_en, 'a', encoding="utf-8") as f:
                                        f.write(corrected_word + '\n')
                                with open(self.not_found_file_en, 'r', encoding="utf-8") as f:
                                        self.df_not_found_en =f.read().splitlines()  
                            except:
                                with open(self.not_found_file_en, 'a') as f:
                                        f.write(corrected_word + '\n')
                                with open(self.not_found_file_en, 'r') as f:
                                        self.df_not_found_en =f.read().splitlines()  
                else:
                    if unormal_word :
                       corrected_word=unormal_word
                    else:
                        corrected_word=word
            corrected_words.append(corrected_word)
        correct_text=' '.join(corrected_words)
        return correct_text
    def on_checkbox_state_changed(self,state):
            self.coorrect_aktive=False
    def latin_semi_auto_correction(self,word,spell,text): 
        if self.coorrect_aktive==False:
           return word    
        try:
            if spell.candidates(word) is not None:
               synonyms_str = list(spell.candidates(word))
            else:
                synonyms_str=[]
            app =QApplication(sys.argv)
            win = QMainWindow()
            win.setGeometry(200, 200, 790, 500)
            win.setWindowTitle("Correction Word")
            label = QTextEdit(win)
            label.resize(600, 66)
            label.move(100,20)
            label.setText(text)
            self.highlight_words_curser(label,word)
            entry = QLineEdit(win)
            entry.setText(word)
            entry.move(100, 160)
            entry.resize(600, 35)
            label2 = QLabel(win)
            label2.setText(self.M131)
            label2.move(100, 110)
            label2.resize(600, 35)
            combo = QComboBox(win)
            combo.insertItem(0, self.M130)
            combo.addItems(synonyms_str)
            combo.setCurrentIndex(0)
            combo.move(100, 240)
            combo.resize(600, 35)
            hand_word = [None]
            checkbox = QCheckBox(self.M136, win)
            checkbox.stateChanged.connect(self.on_checkbox_state_changed)
            checkbox.move(100, 320)
            checkbox.resize(600, 35)
            def on_combobox_changed(text):
                hand_word[0] = text
                entry.setText(text)
            combo.currentTextChanged.connect(on_combobox_changed)
            button = QPushButton(win)
            button.setText(self.M132)
            button.move(100, 366)
            button.resize(600, 35)
            button.clicked.connect(win.close)
            win.show()
            app.exec_()
            if entry.text():  # If a word is manually entered
                hand_word[0] = entry.text()
            else:
                hand_word[0] = word
            return hand_word[0]
        except Exception as e:
            #self.update_notification(f"  {self.M44}  : {e}")
            return word
    def open_5(self):
        try:
            if self.S7==True and self.S8==True:
                return
            if self.from_code=='ar':
                if self.S8==True:
                   return
                #self.update_notification(self.M162)
                self.not_found_file_ar ='not_found_ar.json'
                with open(self.not_found_file_ar, 'r', encoding='utf-8') as f:
                    self.df_not_found_ar=f.read().splitlines() 
                self.df_not_found_ar = list(self.df_not_found_ar)
                self.S8=True
            else:
                if self.S7==True:
                    return
                ##self.update_notification(self.M162)
                self.not_found_file_en ='not_found_en.json'
                with open(self.not_found_file_en, 'r') as f:
                    self.df_not_found_en =f.read().splitlines() 
                self.df_not_found_en = list(self.df_not_found_en)
                self.S7=True
        except:
                if not os.path.exists('not_found_ar.json'):
                   with open('not_found_ar.json', 'w') as f:
                        f.write('') 
                if not os.path.exists('not_found_en.json'):
                   with open('not_found_en.json', 'w') as f:
                        f.write('')
    def correction_english_text(self, text):

        if not self.correct.get() in [self.M15 ,self.M128]:
           return text
        if not text or isinstance(text, float) or text==" " :
            return text
        self.open_5()
        corrected_words=self.spellingcheck(text) 
        return corrected_words
    def correction_Farsi_text(self, text):
        if not self.coorrect_aktive==True and not self.Farsi_text_edit.get() and not self.correct.get()in [self.M15 ,self.M128]:
            return text
        try:
            if self.S6==False:
               self.open_4() 
            correct_text = self.Farsi_correction_words(text)
            return correct_text
        except :
            return text
    def Farsi_correction_words(self, text):
        self.not_save_active=False
        try:
            if not self.correct.get()  in [self.M15 , self.M128]:
                return text   
            text=self.replace_characters(text,False)                
            words = self.Farsi_check_singel_char(text)  
            corrected_words = []
            for word in words:
                if self.coorrect_aktive==False:
                    corrected_words.append(word)
                    continue
                if len(word)<2 or any(re.search(pattern, word) for pattern in self.patterns):
                    corrected_words.append(word)
                    continue
                if  isinstance(word, int) or word == " " or not word or word in self.fa_words or word in self.fa_w_not_found :
                    corrected_words.append(word)
                    continue
                if word in self.replaced_words: 
                    rep_line_ok=False       
                    for line in self.replaced_lines:
                        if word in line.capitalize() :                          
                            rep_line = line.replace(word, '')
                         #   rep_line = line.replace(' ', '')
                            rep_line_ok=True
            
                            corrected_words.append(rep_line) 
                            break 
                    if rep_line_ok==True :
                      continue
                if  self.correct.get()==self.M128:
                    word= self.farsi_semi_auto_correction(word,text)
                    corrected_words.append(word)
                    continue
                else:
                    self.update_notification(f"  {self.M84} : {word}")
                    similar_word = self.farsi_auto_correction(word)
                    if similar_word is not None:
                        corrected_words.append(similar_word)
                        if similar_word not in self.replaced_words and  word!=similar_word: 
                           self.info2(f"{word} {self.M257} {similar_word}")
                           self.saveReplacedWords(word,similar_word)
                        else:
                           if not re.findall(r' ', similar_word)  and similar_word not in self.fa_w_not_found :
                              self.saveNewWord(similar_word)             
                        continue           
                    else:
                        corrected_words.append(word)
                        if word not in self.fa_w_not_found:
                            self.saveNewWord(word)
                        continue          
            return ' '.join(corrected_words)
        except:
            return text

    def farsi_auto_correction(self, word):
        
        if self.S6==False:
            self.open_4()
        if len(self.words2)<2:
           self.words2=farsi_tool.stop_words()
        def replace_and_check(i, char):
            new_word = word[:i] + char + word[i+1:]
            if  new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                return new_word
            return None
        if len(word) < 3 :
           reverse_word =word[::-1]
           if reverse_word in self.words2 or reverse_word in self.fa_words:
              return reverse_word
        try:
            similar_word = self.sim(word)
            if similar_word is not None:
                return similar_word
            similar_word = self.risheyabi(word)
            if similar_word is not None:
                return similar_word
            for i in range(len(word)):
                for char_group in self.char_groups:
                    if word[i] in char_group:
                        for char in char_group:
                            new_word = word[:i] + char + word[i+1:]
                            if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                               return new_word  

            for i in range(len(word)):
                for char in self.chars:
                    new_word = replace_and_check(i, char)
                    if new_word is not None:
                        return new_word
            return  None
        except :
            return None
    def risheyabi(self, word):
        try:
            self.save_word=False
            #self.words2=farsi_tool.stop_words()
            def replace_and_check(i, char):
                new_word = word[:i] + char + word[i+1:]
                if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    self.save_word=True
                    return new_word
                return None
            for i in range(len(word) - 1):
                for char_group in self.char_groups:
                    if word[i:i+1] in char_group:
                        for char in char_group:
                            new_word = replace_and_check(i, char)
                            if new_word is not None:
                                return new_word
            for i in range(len(word) + 1):
                for char in self.chars:
                    new_word = word[:i] + char + word[i:]
                    if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                        return new_word
        
            if  7>len(word) > 4:
                for suffix in self.start_with:            
                    if word.startswith(suffix):
                        new_word = word[:-len(suffix)]
                        if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                            self.save_word=True   
                            return word
                        else:
                            for suffix in self.suffixes:              
                                if new_word.endswith(suffix):
                                    new_word = new_word[:-len(suffix)]
                                    if new_word in self.fa_w_not_found or new_word in self.fa_words: 
                                        if  new_word in self.words2 and len(suffix)>2:
                                            return new_word+' '+suffix+' '
                                        self.save_word=True
                                        return word
                for suffix in self.suffixes:              
                    if word.endswith(suffix):
                        new_word = word[:-len(suffix)]
                        if new_word in self.fa_w_not_found or new_word in self.fa_words:   
                            if  new_word in self.words2 and len(suffix)>2:
                                return new_word+' '+suffix+' '
                            self.save_word=True
                            return word
                        else:
                            for suffix in self.start_with:            
                                if new_word.startswith(suffix):
                                    new_word = new_word[:-len(suffix)]
                                    if new_word in self.fa_w_not_found or new_word in self.fa_words:  
                                        self.save_word=True  
                                        return word                   
                for suffix in self.start_with:            
                    if word.startswith(suffix):
                        new_word = word[:-len(suffix)]
                        if new_word in self.fa_w_not_found or new_word in self.fa_words:  
                            self.save_word=True  
                            return word
            if  len(word) > 7 :  
                new_word = word[:-1]
                if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    return word
                new_word = word[:-2]
                if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if  new_word in self.words2 and len(word[:-2])>1:
                        return new_word +" "+ word[-2:]+' '
                    self.save_word=True
                    return word
                new_word = word[:-3]
                if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if len(word[:-3])>2:
                        return new_word +" "+ word[-3:]+' '
                    self.save_word=True
                    return word
                new_word = word[:-4]
                if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if len(word[:-4])>2:
                        return new_word +" "+ word[-4:]+' '
                    self.save_word=True
                    return word 
                new_word = word[:-5]
                if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if len(word[:-5])>2:
                        return new_word +" "+ word[-5:]+' '
                    self.save_word=True
                    return word            
                new_word = word[:-6]
                if new_word in self.words2 or new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if len(word[:-6])>2:
                        return new_word +" "+ word[-6:]+' '
                    self.save_word=True
                    return word

            return None
        except:
            return None
    def saveNewWord(self,newword):
            if self.S6==False:
                self.open_4()
            with open(self.not_found_file, 'a', encoding="utf-8") as f:
                    f.write(newword + '\n') 
            with open(self.not_found_file, 'r', encoding="utf-8") as f:
                    self.fa_w_not_found=f.read().splitlines() 
    def saveReplacedWords(self,word,replacedWord):
        if self.S6==False:
            self.open_4()
        self.replace='replace.json'
        #self.info2(f"{word+'   '+'جایگزین شد با '+'   '}{replacedWord}")
        with open(self.replace, 'a' , encoding="utf-8") as f:
            f.write(replacedWord+"  "+word + '\n')
        with open(self.replace, 'r', encoding='utf-8') as f:
            self.replaced_lines =f.read().splitlines()
        with open(self.replace, 'r', encoding='utf-8') as f:
            self.replaced_words=f.read().split()
    def sim(self, word):
        self.save_word=False
        try:
            for x in self.fa_w_not_found:
                ratio = Sequence.SequenceMatcher(None, word, x).ratio()
                if ratio > 0.90:
                   return x 
            similar_words = rapidfuz.process.extract(word, self.fa_words, limit=24)
            similar_words_same_length = [w[0] for w in similar_words if len(w[0]) == len(word)]
            if similar_words_same_length:
                return similar_words_same_length[0]          
            if  len(word)>4:
                similar_words_same_length = [w[0] for w in similar_words if  (len(word)-1>=len(w[0]) and len(word)+3<=len(w[0]))  or  (len(word)+1<=len(w[0]) and len(word)-2>=len(w[0])) ]
                if similar_words_same_length:
                    return similar_words_same_length[0] 
            if  len(word)>5:
                new_word = word[:-2]
                if new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if len(word[:-2])>1:
                       return new_word +" "+word[-2:]+" "
                    self.save_word=True                    
                    return new_word +word[-2:]
                new_word = word[:-3]
                if new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if len(word[:-3])>1:
                       return new_word +" "+ word[-3:]+" "
                    self.save_word=True
                    return new_word +word[-3:]
                new_word = word[:-4]
                if new_word in self.fa_w_not_found or new_word in self.fa_words: 
                    if len(word[:-4])>1:
                       return new_word +" "+ word[-4:]+" "
                    self.save_word=True
                    return new_word + word[-4:]
            return None
        except :
            return None
    def saveNewWord(self,newword):
        with open(self.not_found_file, 'a', encoding="utf-8") as f:
                f.write(newword + '\n') 
        with open(self.not_found_file, 'r', encoding="utf-8") as f:
                self.fa_w_not_found=f.read().splitlines() 
    def saveReplacedWords(self,word,replacedWord):
        self.replace='replace.json'
        self.info2(f"{word+'   '+'جایگزین شد با '+'   '}{replacedWord}")
        with open(self.replace, 'a' , encoding="utf-8") as f:
            f.write(replacedWord+" "+word + '\n')
        with open(self.replace, 'r', encoding='utf-8') as f:
            self.replaced_lines =f.read().splitlines()
        with open(self.replace, 'r', encoding='utf-8') as f:
            self.replaced_words=f.read().split()
        if not replacedWord in self.fa_w_not_found:
            self.saveNewWord(replacedWord)

    def check_single_char(self, words, i):
        try:
            if i > 0 and (words[i-1] + words[i]) in self.fa_words:
                return 'before'
            elif i < len(words) - 1 and (words[i] + words[i+1]) in self.fa_words:
                return 'after'
            else:
                return 'none'
        except:
             return 'none'
    def Farsi_check_singel_char(self, text):
        try:
            words=text.split()
            corrected_words = []   
            i = 0
            while i < len(words):
                word = words[i]
                corrected_word = word  # Initialize corrected_word with the original word
                if len(word) == 1 and word not in['و'] :
                    check_result = self.check_single_char(words, i)
                    if check_result == 'before':
                        corrected_word = corrected_words[-1] + word  # Join with the word before
                        corrected_words[-1] = corrected_word  # Update the last word in corrected_words
                        words[i]=''
                    elif check_result == 'after':
                        corrected_word = word + words[i+1]  # Join with the word after
                        corrected_words.append(corrected_word)  # Add the corrected word to corrected_words
                        words[i+1]=''
                    else:
                        corrected_words.append(word)  # Do not join
                    i += 1 
                else:
                    corrected_words.append(word)
                    i += 1
            return corrected_words
        except:
                return words
    def fix_space_word(self, text):       
        words = text.split()
        try:
            for i in range(len(words)): 
                for char in self.special_chars:
                    if char in words[i]:
                        words[i] = words[i].replace(char, ' ' + char + ' ' ) 
        except Exception :
            words[i] = ""
        return ' '.join(words) 
    def space_correction(self,text):
        try:
            text=self.fix_space_word(text)
            if not self.to_code in ['fa','ur','ar']:
                 text = ftfy.fix_text(text)           
            text = re.sub(r'([^\w\s])', r' \1' , text)
            text = re.sub(r'(\d+)', r' \1 ', text)
            text = re.sub(r'(\d+\))', r' \1 ', text)     
            text = re.sub(r'("  ")', r' \1 ', text)
            return text
        except:
            pass
    def format_text(self, text):
        try:
            if  self.format_box.cget('text') in [self.M182] :
                return self.reverse_text_ar(text)
            if  self.format_box.cget('text') in [self.M152] :
                return self.reverse_text_fa(text)
            if  not self.format_box.cget('text') in [self.M12 ,self.M153] :
                return text   
            formatted_imp_text = text
            if not self.skip_rtl==True and not self.console==True:
                if self.from_code in ["fa","ur"]:   
                    formatted_imp_text= self.reverse_text(text)
                else:
                    formatted_imp_text= self.reverse_text_ar(text)
            else:
                if not self.console==True:
                    formatted_imp_text=self.format_text_2(formatted_imp_text)                   
            self.rev=True       
            return formatted_imp_text
        except :
            return text

    def reverse_text(self, text):
        words = text.split()
        reverse_words = [word if not re.search(r'[\u0600-\u06FF]', word) else word[::-1] for word in words]
        return self.auto_reverse_mix_text(reverse_words)

    def auto_reverse_mix_text(self, words):
        grouped_words = []
        for i in range(len(words)):
            if not re.search(r'[\u0600-\u06FF]', words[i]) and (i+1 < len(words) and not re.search(r'[\u0600-\u06FF]', words[i+1])):
                grouped_word = words[i] + '~'
            else:                            
                grouped_word = words[i]+'_'
            grouped_words.append(grouped_word)
        new_text = "".join(grouped_words)
        new_words = new_text.split('_')
        new_words.reverse() 
        return str(" ".join(self.replace_reversed(new_words)))
    def replace_reversed(self, words):
        for i in range(len(words)): 
            words[i] = words[i].replace('~', ' ')  
            words[i] = words[i].replace('  ', ' ')                      
        return words

    def format_text_2(self, text):
        try:
            words=text.split()
            R_words=self.auto_reverse_mix_text(words)
          #  formatted_imp_text =self.replace_reversed(R_words)
            self.rev=True
            return R_words
        except :
            return words

    def reverse_text_ar(self, text):
        try:
            if self.from_code in ["fa","ur"]:
                text=self.reverse_words_and_sentence(text)
                return text
            text=self.reverse_words(text)
            return text
        except:
                pass
    def reverse_text_fa(self, text):
        text=get_display(text)
        return text 
    def reverse_words(self,text):
        try:
            reverse_words=[]
            words=text.split()
            for word in words:
                if  not re.search(r'[\u0600-\u06FF]', word):
                    reverse_word = word
                else:
                    reverse_word = word[::-1]
                reverse_words.append(reverse_word)
            return " ".join(reverse_words)
        except:
            return text
    def reverse_sentence(self,text):
        words=text.split()
        words.reverse()
        return " ".join(words)

    def reverse_words_and_sentence(self, text):
        try:
            words=text.split()
            reverse_words = []
            for word in words:
                if re.search(r'[A-Za-z0-9]+', word) :
                    reverse_word = word
                else:
                    reverse_word = word[::-1]
                reverse_words.append(reverse_word)
            reverse_words.reverse()
            return ' '.join(reverse_words)
        except:
            return text
    def farsi_semi_auto_correction(self, word,text):
        try:
            size = int(self.size_box.get())
            fo = self.font_box.cget("text")
            font = QtGui.QFont(fo)
            font.setPointSize(size)
            similar_words = rapidfuz.process.extract(word, self.fa_words, limit=14)
            synonyms_str = [word[0] for word in similar_words]  # Extract only the words
            app = QApplication(sys.argv)
            win = QMainWindow()
            win.setGeometry(66, 66, 769, 600)
            win.setWindowTitle(self.M133)
            label = QTextEdit(win)
            label.setFont(font)
            label.setText(text)
            self.highlight_words_curser(label, word)
            label.move(100, 20)
            label.resize(600,110)  # Change the width to 200
            self.Farhang = QTextEdit(win)
            self.Farhang.setFont(font)
            self.Farhang.setReadOnly(True)
           # self.Farhang.setText(text)
            self.Farhang.move(100, 250)
            self.Farhang.resize(600, 110)  # Change the width to 200
            self.button_Farhang = QPushButton(win)
            self.words_list = QPushButton(win)
            entry = QLineEdit(win)
            entry.setFont(font)
            entry.setText(word)
            label.setReadOnly(True)
            entry.move(100, 166)
            entry.resize(600, 35)
            label2 = QLabel(win)
            label2.setText(self.M131) 
            label2.move(10, 135)
            label2.resize(600, 35)  # Change the width to 200
            combo = QComboBox(win)
            combo.insertItem(0,self.M130)
            combo.addItems(synonyms_str)
            combo.setCurrentIndex(0)
            combo.move(100, 365)
            combo.resize(600, 35)
            hand_word = None
            def on_combobox_changed(text):
                nonlocal hand_word
                hand_word = text
                entry.setText(hand_word)
            combo.currentTextChanged.connect(on_combobox_changed)
            button = QPushButton(win)
            button.setText(self.M132)
            button.move(100, 500)
            button.resize(600, 55)
            button.clicked.connect(win.close)
            def on_checkbox_state_changed(state):
                self.correct.set(self.M84)
                self.coorrect_aktive=False
            checkbox = QCheckBox(self.M136, win)
            checkbox.stateChanged.connect(on_checkbox_state_changed)
            checkbox.move(100, 460)
            checkbox.resize(600, 35)
            checkbox_2 = QCheckBox(self.M142, win)
            checkbox_2.move(313, 460)
            checkbox_2.resize(600, 35)
            self.button_Farhang.move(100, 205)
            self.button_Farhang.resize(600, 35)

            self.words_list.resize(600, 35)
            self.words_list.move(100, 420)
            self.hand_word=word

    
            self.button_Farhang.clicked.connect(self.searching_farhang)
            self.button_Farhang.setText(self.M246)
            self.words_list.setText(self.M247)
            self.words_list.clicked.connect(self.words_list_F)

            win.show()
            app.exec_()
            if  entry.text():
                hand_word = entry.text()
            elif hand_word is None:
                hand_word = word
            self.hand_word=hand_word
            confirmation = QMessageBox.No
            if not checkbox_2.checkState() == Qt.Unchecked :
               self.not_save_active=True
            if  checkbox_2.checkState() == Qt.Unchecked and  self.not_save_active==False : 
                if hand_word!=word :
                    self.saveReplacedWords(word,hand_word)
                if  hand_word not in self.fa_w_not_found and hand_word not in self.fa_words  :
                    if  not re.findall(r' ', hand_word):
                        self.saveNewWord(hand_word)
            else:
                if hand_word!=word :
                    confirmation = QMessageBox.question(None,self.M245,f" {self.M245}   {word}  ->  {hand_word} " , QMessageBox.Yes |  QMessageBox.No)
                    if confirmation == QMessageBox.Yes:
                        self.saveReplacedWords(word,hand_word)
                        confirmation = QMessageBox.No
                if  hand_word not in self.fa_w_not_found and hand_word not in self.fa_words  :
                    if  not re.findall(r' ', hand_word):
                        confirmation = QMessageBox.question(None,self.M134,f"{self.M137} ' {hand_word} '" , QMessageBox.Yes |  QMessageBox.No)
                        if confirmation == QMessageBox.Yes:
                            self.saveNewWord(hand_word)

            return hand_word
        except Exception:
            return word
    def searching_farhang (self):
        try:
            self.clicked=False
            self.clicked_trueWords=False
            self.words_list.setText(self.M249)
            lines=[]
            self.skip_save=True
            similar_word = self.farsi_auto_correction(self.hand_word)
            if similar_word is not None:
                self.corrected_word= similar_word
                self.button_Farhang.setText(self.M250+" "+self.corrected_word)
            else:
                return self.button_Farhang.setText("کلمه صحیح یافت نشد")
            if len(self.farhang) <100:
                with open('fa_dic.json', 'r', encoding='utf-8') as f:
                    self.farhang =f.read().splitlines()
            for line in self.farhang:
                words = line.split()
                if self.corrected_word in words:
                    lines.append(line+'\n')
                else:continue
            if not len(lines)>0:
                FARHANG_RESULTS=self.corrected_word
            else:
                FARHANG_RESULTS=" ".join(lines) 
            self.Farhang.setText(FARHANG_RESULTS)
            self.highlight_words( self.Farhang,self.corrected_word)
            return lines
        except:
            pass
    def words_list_F(self):  
        try: 
            if  self.clicked_trueWords==False : 
                if not self.clicked==True  :
                    self.clicked=True
                    with open(self.replace, 'r', encoding='utf-8') as f:
                        self.replaced_=f.read()
                    self.Farhang.clear() 
                    self.listReplaced=len(self.replaced_)
                    self.Farhang.setText(self.replaced_)
                    self.Farhang.setReadOnly(False)
                    self.words_list.setText(f"{self.M248}  {self.M249} ")
                else:
                    if  len(self.Farhang.toPlainText()) >14:
                        with open(self.replace, 'w' , encoding="utf-8") as f:
                            f.write(self.Farhang.toPlainText())
                        with open(self.replace, 'r', encoding='utf-8') as f:
                            self.replaced_lines =f.read().splitlines()
                    self.Farhang.setReadOnly(True)
                    self.clicked=False
                    self.words_list.setText(self.M247)
                    self.clicked_trueWords=True
            else:
                if not self.clicked==True:   
                    self.clicked=True            
                    with open(self.not_found_file, 'r', encoding='utf-8') as f:
                        self.not_found_words=f.read()
                    self.list_newWords=len(self.replaced_)
                    self.Farhang.clear() 
                    self.Farhang.setText(self.not_found_words)
                    self.Farhang.setReadOnly(False)
                    self.words_list.setText(f"{self.M248}  {self.M247} ")
                else:
                    if  len(self.Farhang.toPlainText()) >5 :
                        with open(self.not_found_file, 'w' , encoding="utf-8") as f:
                            f.write(self.Farhang.toPlainText())
                        with open(self.not_found_file, 'r', encoding='utf-8') as f:
                            self.fa_w_not_found =f.read().splitlines()
                    self.Farhang.setReadOnly(True)
                    self.clicked=False
                    self.words_list.setText(self.M249)
                    self.clicked_trueWords=False
        except:
            self.Farhang.setReadOnly(True)
   # from PyQt5 import QtGui

    def highlight_words(self, text_edit, word):
        try:
            format = QtGui.QTextCharFormat()
            # Change the color of the text
            format.setForeground(QtGui.QBrush(QtGui.QColor("blue")))
            cursor = text_edit.textCursor()
            cursor.setPosition(0)
            while True:
                cursor = text_edit.document().find(word, cursor.position())
                if not cursor.isNull():
                    cursor.mergeCharFormat(format)
                else:
                    break
        except:
            pass
    def highlight_words_curser(self, text_edit, word):
        try:
            format = QtGui.QTextCharFormat()
            # Change the color of the text
            format.setForeground(QtGui.QBrush(QtGui.QColor("blue")))
            cursor = text_edit.textCursor()
            cursor.setPosition(0)
            while True:
                cursor = text_edit.document().find(word, cursor.position())
                if not cursor.isNull():
                    cursor.mergeCharFormat(format)
                   
                    text_edit.setTextCursor(cursor) 
                    text_edit.ensureCursorVisible() 
                else:
                    break
        except:
            pass

    def al110(self,state):
        self.internal=True
        if self.dgh==False :
            self.deghat.set(True)
            self.dgh=True
        else:
            self.deghat.set(False)  
            self.dgh=False 
    def on_button_clicked_2(self):
        try:
            QApplication.exit()
        except:
            pass
    def on_button_clicked(self):
        try:
            if self.entery_s:
               self.entery_s.clear() 
            if self.separate_search==True:
               #self.off_ketab()
               self.separate_search=False
            QApplication.exit()
        except:
            pass
    def save_file_in(self):
            self.export_docx_=True
            try:
                doc = dox.Document()
                size = int(self.size_box.get())
                new_font = self.font_box.cget("text")
                if self.color_code[0]:  # If a color is selected
                    rgb_color = tuple(int(self.color_code[0][i]) for i in range(3))  # Get the RGB color
                else:  # If no color is selected, use a default color
                    rgb_color = (0, 0, 0) 
               
                text=self.entry2.toPlainText().split('\n')                
                if  self.pdf_convert==True or len(text)>2 :
                    self.file_pat = filedialog.asksaveasfilename(
                        title="Export Translated Text",
                        filetypes=[("Word files", "*.docx")],
                        defaultextension=".docx" )             
                    if not self.file_pat:
                        return 
                    for para in text :                                
                        #para1=self.space_correction(para)
                        paragraph = doc.add_paragraph(para)
                        for run in paragraph.runs:
                            run.font.size = dox.shared.Pt(size)  # Set the font size
                            run.font.name = new_font  # Set the font style
                            run.font.color.rgb =dox.shared.RGBColor(*rgb_color)  # Set the font color                    
                    doc.save(self.file_pat)
                    os.startfile(self.file_pat)
                    self.process_finish()  
                    self.rev=False     
                else: self.update_notification(self.M197)                                                  
            except Exception as e:               
                self.process_finish()
                self.rev=False
                messagebox.showinfo(self.M44, f"{str(e)}")
    def research(self):
        self.total.setText(str(self.total_para))
        self.i=0
        self.entry2.clear() 
        self.stops = False
        try:
            lines=[]
            word= f"{self.entry.toPlainText()}"
            if len(word)<2:
                return None
            if self.dict in [self.M207] :
                results=self.abjad(word)
                self.entry2.clear()                      
                self.entry2.setText(results)
                self.separate_search=False
                return self.entry.toPlainText()
            elif self.search_process==self.M207:
                results=self.abjad(word)
                self.entry2.clear()                      
                self.entry2.setText(results)
                self.separate_search=False
                return self.entry.toPlainText()
            elif self.search_process in [self.M164,self.M163]:
                results =str(self.search_internet(word))
                self.entry2.clear()                      
                self.entry2.setText(results)
                self.highlight_words(self.entry2, word)
                self.separate_search=False
                return self.entry.toPlainText()
            else:               
                self.total_para=int(self.total.text())
                if not self.separate_search==True:
                    if  word in self.persian_conjunctions:
                        return self.entry.toPlainText()
                for line in self.ketab_:
                    if self.end==True:
                        break
                    if  self.deghat.get():
                        
                        if self.ktb==self.M194:
                            word=self.replace_characters(word,True)
                            line1=self.clean_text_arabic(line)
                            words = line1.split()
                            word=self.clean_text_arabic(word)
                        else:
                            words = line.split()
                            word=self.replace_characters(word,False)
                            #words = line.split()
                        if words:
                            if word in words:
                                lines.append(line+'\n')
                                self.i+=1
                                if len(lines)+1>self.total_para or self.stops==True: 
                                    break
                            else:
                                if word in words:
                                    lines.append(line+'\n')
                                    self.i+=1
                                    if len(lines)+1>self.total_para or self.stops==True: 
                                        break
                    else:
                        if self.ktb==self.M194:
                            word=self.replace_characters(word,True)
                            line1=self.clean_text_arabic(line)
                            word=self.clean_text_arabic(word)
                            if word in line1:
                                lines.append(line+'\n')
                                self.i+=1
                                if len(lines)+1>self.total_para or self.stops==True: 
                                    break
                        else:
                            word=self.replace_characters(word,False)
                            if word in line:
                                lines.append(line+'\n')
                                self.i+=1
                                if len(lines)+1>self.total_para or self.stops==True: 
                                    break
                self.entry2.clear() 
                self.total_cunt.setText(str(self.i))                               
                if not lines:
                    self.entry2.setText(f"Not found {word} ")
                    self.separate_search=False
                    return None
                else: 
                    results = '\n'.join(lines)
                    self.entry2.setText(results)
                    self.highlight_words(self.entry2, word)
                    self.separate_search=False
                    return self.entry.toPlainText()
        except:
            self.separate_search=False
            return None

    def paste_from_clipboard_qt5(self):
            try:
                self.entry.setText(self.root.clipboard_get())
            except :
                pass
    def copy_to_clipboard_qt5(self):
        try:
            selected_text = self.entry2.toPlainText()
            clipboard = QApplication.clipboard()
            if clipboard is not None:
                clipboard.clear()
            clipboard.setText(selected_text)
            self.update_notification(" متن کپی شد")
        except Exception:
            self.update_notification("کپی نشد")
    def update_total_para(self):
        self.total_para = int(self.total.text())   
 
    def book_search(self,text,word,results): 
        self.end=False
        self.clipboard_text=""
        try:
            self.dgh=False
            self.internal=False
            app = QApplication(sys.argv)
            win = QMainWindow()
            win.setGeometry(560,40, 786, 660)
            win.setWindowTitle(f"                                                                                                          {self.search_process}   "   )
            label = QTextEdit(win)
            label.setText(f"{text}")
            label.setReadOnly(True)
            label.move(14, 20)
            label.resize(765, 76)  
            self.entry = QTextEdit(win)
            size = 14
            font = QtGui.QFont()
            font.setPointSize(size)
            self.entry2 = QTextEdit(win)
            self.total = QLineEdit(win)
            self.total_cunt = QLineEdit(win)
            label2 = QLabel(win)
            self.totalable=QLabel(win)
            self.total_cuntlable=QLabel(win)
            button = QPushButton(win)
            button_search = QPushButton(win)
            button_copy = QPushButton(win)
            button_paste = QPushButton(win)
            button_save = QPushButton(win)
            checkbox = QCheckBox(self.M136, win)
            checkbox_2 = QCheckBox(self.M178, win)
            label2.setText(self.M158) 
            self.entry.setFont(font)
            self.entry.move(150, 140)
            self.entry.resize(520, 35)
            label2.move(300, 100)
            label2.resize(290, 35)  
            self.entry2.setFont(font)
            self.entry2.setAlignment(Qt.AlignRight)
            self.entry2.setReadOnly(True)
            self.entry2.resize(765, 360)
            self.entry2.move(14, 185)
            self.totalable.setFont(font)
            self.totalable.resize(270, 30)
            self.totalable.move(400, 550)
            self.totalable.setText(" حداکثر تعداد یافتن را وارد فرمایید")            
            self.total.setFont(font)
            self.total.resize(66, 30)
            self.total.move(320, 550)
            self.total.setText(str(self.total_para))
            self.total.textChanged.connect(self.update_total_para)
            self.total_cuntlable.setFont(font)
            self.total_cuntlable.resize(270, 30)
            self.total_cuntlable.move(400, 580)
            self.total_cuntlable.setText("تعداد یافته ها")

            self.total_cunt.setFont(font)
            self.total_cunt.resize(66, 30)
            self.total_cunt.move(320, 580)
          #  self.total_cunt.setText(str(self.i))
            self.total_cunt.setReadOnly(True)

            button.setText(self.M132)
            button.move(14, 610)
            button.resize(110, 35)
            self.entry.setText(word)
            self.entry2.setText(results)
            self.total_cunt.setText(str(self.i)) 
            if not self.books==self.M183 and not self.sina==self.M181:
                self.highlight_words(self.entry2, word)
            button_search.setText(self.M198)
            button_search.move(14, 140)
            button_search.resize(135, 35)
            button_paste.setText(self.M21)
            button_paste.move(669, 140)
            button_paste.resize(110, 35)
            button_copy.setText(self.M20)
            button_copy.move(135, 610)
            button_copy.resize(110, 35)
            button_save.setText(self.M4)
            button_save.move(280, 610)
            button_save.resize(210, 35)        
            if self.separate_search==True:
                self.research()
            button_search.clicked.connect(self.research)
            self.entery_s=self.entry2
            button.clicked.connect(self.on_button_clicked)
            button_paste.clicked.connect(self.paste_from_clipboard_qt5)
            button_copy.clicked.connect(self.copy_to_clipboard_qt5)
            button_save.clicked.connect(self.save_file_in)
            def on_checkbox_state_changed(state):               
                self.off_ketab()
            checkbox.stateChanged.connect(on_checkbox_state_changed)
            checkbox.move(14, 566)
            checkbox.resize(236, 35)
            checkbox_2.stateChanged.connect(self.al110)
            checkbox_2.move(14, 100)
            checkbox_2.resize(236, 35)
            win.show()
            app.exec_()
            if self.entry.toPlainText():
                hand_word = self.entry.toPlainText()
            elif  self.research():
                hand_word=self.research()
            if self.entry.toPlainText():
                hand_word=self.entry.toPlainText()
            else:
                hand_word=word 
            return hand_word                          
        except Exception as e :
            if self.separate_search==True:
               #self.off_ketab()
               self.separate_search=False
            return word
    def abjad(self,word):
        try:
            word = word.replace(' ', '')
            self.search_process=self.M207
            input_string=word
            abjad_dict = { 'ا': 1, 'ب': 2, 'پ': 2, 'ج': 3, 'چ': 3, 'د': 4, 'ه': 5, 'و': 6, 'ز': 7,
                        'ژ': 7, 'ح': 8, 'خ': 8, 'ط': 9, 'ی': 10, 'ک': 20, 'گ': 20, 'ل': 30, 'م': 40, 'ن': 50,
                            'س': 60, 'ع': 70, 'ف': 80, 'ص': 90, 'ق': 100, 'ر': 200, 'ش': 300, 'ت': 400, 'ث': 500,
                            'ذ': 700, 'ض': 800, 'ظ': 900, 'غ': 1000 }
            return str(sum(abjad_dict.get(char, 0) for char in input_string if char.isalpha()))
        except:
            return word

    def ketab(self, text):
        self.i=0
        self.end=False
        self.stops = False
        try:
            new_text = []
            lines=[]
            self.dgh=False
            self.internal=False
            if self.from_code in('fa','ar','ur') and self.ktb in [self.M194,self.M172]:
                text=self.replace_characters(text,True)
            else:
                if self.from_code in('fa','ar','ur'):
                   text=self.replace_characters(text,False)
            if self.sjmle.get():
                for_chunks =text.split()
                words = [' '.join(for_chunks[i:i+3]) for i in range(0, len(for_chunks), 3)]
            else:
                words=text.split()
            for word in words:
                self.i=0
                if  len(word)<3 or self.from_code in('fa','ar','ur') and not re.search(r'[\u0600-\u06FF]', word):
                    new_text.append(word)
                    continue                 
                if self.end==True:
                    new_text.append(word)
                    continue   
                if not self.sjmle.get():
                    if len(word)<3:
                        new_text.append(word)
                        continue   
                if self.search_Active==False:
                    new_text.append(word)
                    continue   
                if self.internet_aktive==True:
                    results=self.search_internet(word)                        
                    new_word=self.book_search(text,word,results)
                    if new_word:
                        new_text.append(new_word)
                    else:
                        new_text.append(word) 
                elif self.dict.get()==self.M207:
                    results=self.abjad(word)
                    new_word=self.book_search(text,word,results)
                    new_text.append(word) 
                else:
                    for line in self.ketab_:                                    
                        if  self.deghat.get():
                            if self.ktb ==self.M194:
                                line_celear=self.clean_text_arabic(line)                               
                                word=self.clean_text_arabic(word)
                                words = line_celear.split()
                            elif self.ktb ==self.M172:
                                line_celear=self.replace_characters(line,True)
                                word=self.replace_characters(word,True)
                                words = line_celear.split()

                            else:
                                words = line.split()
                            if words:
                                if  not self.book_aktive and (words[0] == word or words[-1] == word )or (len(words)>3 and (words[+1] == word or words[+2] == word)):
                                    lines.append(line+'\n')
                                    self.i+=1
                                    if len(lines)+1>self.total_para or self.stops==True:
                                        break
                                else:
                                    if word in words:
                                        lines.append(line+'\n')
                                        self.i+=1
                                        if len(lines)+1>self.total_para or self.stops==True:    
                                            break
                        else:
                            if self.ktb ==self.M194:
                                line_celear=self.clean_text_arabic(line)                               
                                word=self.clean_text_arabic(word) 
                                if word in line_celear:
                                    lines.append(line+'\n')
                                    self.i+=1
                                    if len(lines)+1>self.total_para or self.stops==True: 
                                        break
                            elif self.ktb ==self.M172:
                                line_celear=line
                                word=self.replace_characters(word,False)

                                if word in line:
                                    lines.append(line+'\n')
                                    self.i+=1
                                    if len(lines)+1>self.total_para or self.stops==True: 
                                        break
                            else:
                                word=self.replace_characters(word,False)
                                if word in line:
                                    lines.append(line+'\n')
                                    self.i+=1
                                    if len(lines)+1>self.total_para or self.stops==True: 
                                        break
                                                
                    if not lines:
                        new_text.append(word)                    
                    else:
                        results = '\n'.join(lines)
                        lines = []
                        new_word=self.book_search(text,word,results)
                        if word:
                            new_text.append(new_word)
                        else:
                            new_text.append(word)
            
            return ' '.join(new_text)
        except Exception:
            return text

    def jomlesaz(self,text):
        if self.empty_1==True and len(text) <1 :
           return text
        if len(text) <1:
            self.empty_1=True
        try:
            def on_button_clicked():
                nonlocal corrected_text
                corrected_text = new_word_entry.toPlainText()
                QApplication.exit()
            def on_checkbox_state_changed(state):
                self.jomlesazi.set(False)
            corrected_text = ""
            app = QApplication(sys.argv)
            win = QMainWindow()
            win.setGeometry(313, 313, 457, 396)
            win.setWindowTitle(self.M139)
            new_word_entry = QTextEdit(win)
            new_word_entry.move(10, 36)
            new_word_entry.resize(437, 269)
            new_word_entry.setText(text)
            button = QPushButton(win)
            button.setText(self.M138)
            button.move(157, 350)
            button.resize(157, 35)
            button.clicked.connect(on_button_clicked)
            checkbox = QCheckBox(self.M136, win)
            checkbox.stateChanged.connect(on_checkbox_state_changed)
            checkbox.move(20, 320)
            checkbox.resize(210, 25)
            win.show()
            app.exec_()
            corrected_text = new_word_entry.toPlainText()
            return corrected_text
        except Exception as e:
                self.update_notification(f"  {self.M44}  : {e}")
    def Tem_default(self):               
       self.myColor = '#0deab3'  

       self.regui()
    def Tem_red(self):               
        self.myColor = '#ff9999'  

        self.regui()
    def Tem_blue(self):               
        self.myColor = '#d5e2ff'  

        self.regui()
    def Tem_black(self):               
        self.myColor = '#ffffff'  

        self.regui()
    def Tem_green(self):               
        self.myColor = '#99cc99'  

        self.regui()
    def Tem_nocolor(self):               
        self.myColor = '#ffffff'  

        self.regui()
    def replace_ascii_digits_with_farsi(self,text):
        text = text.replace('0', '۰')
        text = text.replace('1', '۱')
        text = text.replace('2', '۲')
        text = text.replace('3', '۳')
        text = text.replace('4', '۴')
        text = text.replace('5', '۵')
        text = text.replace('6', '۶')
        text = text.replace('7', '۷')
        text = text.replace('8', '۸')
        text = text.replace('9', '۹')
        return text
    def clean_text_arabic(self,text):
            # Normalize the text
            normalized_text = unicodedata.normalize('NFD', text)
            # Remove diacritics
            cleaned_text = "".join(c for c in normalized_text if unicodedata.category(c) != 'Mn')
            return cleaned_text       
    def replace_characters(self, text,research):       
        words = text.split()
        try:
            for i in range(len(words)): 
                for char in self.special_chars:
                    if char in words[i]:
                        words[i] = words[i].replace(char, ' ' + char + ' ' ) 
                words[i] = words[i].replace('.', ' . ')
                words[i] = words[i].replace('َ', '')
                words[i] = words[i].replace('ُ', '')
                words[i] = words[i].replace('  ', ' ')

                if research==True:
                   words[i] = words[i].replace('ی', 'ي') 
                   words[i] = words[i].replace('ئ', 'ي')
                else:
                    words[i] = words[i].replace('ي', 'ی')
                    words[i] = words[i].replace('ئ', 'ی')
                words[i] = words[i].replace(u"\ufeb7", "ش")
                words[i] = words[i].replace(u"\ufeae", "ر")
                words[i] = words[i].replace(u"\ufe96", "ت")
                words[i] = words[i].replace(u"\ufedf", "ل")
                words[i] = words[i].replace(u"\ufeb0", "ز")
                words[i] = words[i].replace(u"\ufee8", "ن")
                words[i] = words[i].replace(u"\ufb93", "گ")
                words[i] = words[i].replace(u"\ufeb3", "س")
                words[i] = words[i].replace(u"\ufeec", "ه")
                words[i] = words[i].replace(u"\ufee3", "م")
                words[i] = words[i].replace(u"\ufecb", "ع")
                words[i] = words[i].replace('ؤ', 'و')
                words[i] = words[i].replace('ﻮ', 'و')
                
                words[i] = words[i].replace('ك', 'ک')
                words[i] = words[i].replace('ﻚ', 'ک')
                words[i] = words[i].replace('ﺑ', 'ب')
                words[i] = words[i].replace('ﺎ', 'ا')
                words[i] = words[i].replace('ك', 'ک')
                words[i] = words[i].replace('ﻚ', 'ک')
                words[i] = words[i].replace('ﺎ', 'ا')

                words[i] = words[i].replace('ﺪ', 'د')
                words[i] = words[i].replace('ﯿ', 'ی')
                words[i] = words[i].replace('ﻪ', 'ه')


                
                words[i] = words[i].replace('ﻬ', 'ه')
                words[i] = words[i].replace('ﻪ', 'ه')
                words[i] = words[i].replace('ﻼ', 'لا')
                words[i] = words[i].replace('ﯿ', 'ی')
                words[i] = words[i].replace('ﻨ', 'ن')

                
                words[i] = words[i].replace('ﺮ', 'ر')
                words[i] = words[i].replace('ﺷ', 'ش')
                words[i] = words[i].replace('ﺐ', 'ب')
   

                words[i] = words[i].replace('ﯽ', 'ی')
                words[i] = words[i].replace('ﻐ', 'غ')     


                words[i] = words[i].replace('ﻌ', 'ع') 

                words[i] = words[i].replace('ﺢ ', 'ح')
   
                words[i] = words[i].replace('إ', 'ا')
                words[i] = words[i].replace('أ', 'ا')
                words[i] = words[i].replace('إ', 'ا')
                
                words[i] = words[i].replace('ۀ', "ه")
                words[i] = words[i].replace('أ', 'ا')
                words[i] = words[i].replace('%', '٪')
                words[i] = words[i].replace('ْ', '')
                words[i] = words[i].replace('ٍ', '')
                words[i] = words[i].replace('ْ', '')
                words[i] = words[i].replace('ٍ', '')
                words[i] = words[i].replace('ً', '')
                words[i] = words[i].replace('ٌ', '')
                words[i] = words[i].replace('ٍ', '')
                words[i] = words[i].replace('َ', '')
                words[i] = words[i].replace('ُ', '')
                words[i] = words[i].replace('', '')
                words[i] = words[i].replace(" ّ", "")
                words[i] = words[i].replace(u"\u200F", "")
                words[i] = words[i].replace(u"\u200E", "")
                words[i] = words[i].replace(u"\u200D", "")
                words[i] = words[i].replace(u"\u200B", "")
                words[i] = words[i].replace(u"\u200A", "")
                words[i] = words[i].replace(u"\u00A0", "")
                words[i] = words[i].replace(u"\u0640", "")
                words[i] = words[i].replace(u"\u0640", "")
                words[i] = words[i].replace(" ْ", "")

                words[i] = words[i].replace(' ِ', '')
                words[i] = words[i].replace('  ', ' ')
                words[i] = words[i].replace('  ', ' ')
                if words[i] == "اهلل":
                    words[i]='الله'
                words[i] = words[i].replace("الل ه", "الله")
                words[i] = words[i].replace("اهلل", "الله")
        except Exception :
            words[i] = ""
        return ' '.join(words)   
    def search_internet(self,text):
        try:
            if self.internet.get()==self.M164:
                self.search_process=self.M164
                wikipedia.set_lang(self.from_code)
                ny=wikipedia.page(text)
                return ny.content
            else:
                self.search_process=self.M163
                r = ggl.ggl(text, lang=self.from_code, max_results=36)
                r=self.clean_text(str(r))
                return(r)
        except:
            return text
    def clean_text(self, text):
        pattern = r'http(s)?://\S+'
        text = re.sub(pattern, '\n', text)
        pattern = r"[\u200c'}S'body''title''search?num''200'{,href]"
        text = text.replace("\\", " ")
        text = text.replace("'body'", " ")
        text = text.replace("search?num", " ")
        text = text.replace("body", " ")
        text = text.replace("\u200c", " ")
        text = text.replace("'href'", " ")
        text = re.sub(pattern, ' ', text)
        pattern = r'(?<=\.|\?)\s'
        texts = re.split(pattern, text)
        text='\n'.join(texts)
        pattern = r"('href'| 'body'|'title'|{}|'u200c')"
        text = re.sub(pattern,'\n', text)
        pattern = r'(?<=[.!?]) +'
        texts = re.split(pattern, text)
        new_text = '\n'.join(texts)
        return new_text
    def motaradef(self, text):
        new_text = []
        self.end=False
        try:
            words=text.split()
            app = QApplication(sys.argv)
            win = QMainWindow()
            win.setGeometry(200, 200, 796, 500)
            win.setWindowTitle('مترادف ها')
            label = QTextEdit(win)
            label.setText(text)
            label.move(100, 20)
            #label.setWordWrap(True)
            label.resize(600, 96)  
            entry = QLineEdit(win)
            label2 = QLabel(win) 
            combo = QComboBox(win)
            button = QPushButton(win)
            checkbox_2 = QCheckBox(self.M178, win)
            checkbox = QCheckBox(self.M136, win)
            entry.move(100, 166)
            entry.resize(600, 35)
            label2.setText('لغت مترادف را بنویسید یا از لیست ریز انتخاب کنید') 
            label2.move(100, 110)
            label2.resize(600, 35) 
            combo.insertItem(0,'انتخاب مترادف')
            combo.setCurrentIndex(0)
            combo.move(100, 240)
            combo.resize(600, 35)
            button.move(100, 300)
            button.resize(600, 35)
            checkbox.move(100, 360)
            checkbox.resize(600, 35)
            checkbox_2.move(100, 400)
            checkbox_2.resize(236, 35)
            if  self.S5==False:
                self.open_3()
            lin_word=[]
            lines=""
            word1=''
            self.dgh=False 
            for word in words:
                if self.end==True:
                    new_text.append(word)
                    continue
                adi=False
                for line in self.synonyms_str:
                    l=line.split()
                    if  word in l :
                        lines +=line
                    elif self.deghat.get() and not word in line:
                        if len(word)>2:
                            for suffix3 in self.suffixes:              
                                if word.endswith(suffix3):
                                    new_word = word[:-len(suffix3)]
                                    if new_word in l:                         
                                        word1= new_word
                                        adii=suffix3
                                        adi=True
                                        lines +=line
                if not lines:
                    new_text.append(word)
                    continue
                else:
                    combo.clear()
                    lin_word=lines.split()
                    lines=""
                    combo.addItems(lin_word)
                    lin_word=""
                    hand_word = ""
                    def on_combobox_changed(text):
                        nonlocal hand_word
                        if not adi :
                            hand_word = text
                        else:hand_word=text+adii
                        entry.setText(hand_word)     
                    combo.currentTextChanged.connect(on_combobox_changed)
                    entry.setText(word)
                    button.setText(self.M132)
                    button.clicked.connect(win.close)
                    def on_checkbox_state_changed(state):
                        combo.clear()
                        self.off_ketab()
                    checkbox.stateChanged.connect(on_checkbox_state_changed)
                    checkbox_2.stateChanged.connect(self.al110)
                    win.show()
                    app.exec_()
                    if  entry.text():
                        hand_word = entry.text()
                    elif hand_word =="" and word:
                        if not adi :  
                            hand_word = word
                        else:
                            hand_word = word1+adii                      
                    new_text.append(hand_word)
                    combo.clear()
                    continue
        except Exception as e:
            self.update_notification(f"  {self.M44}  : {e}") 
        return " ".join(new_text) 
    def download_L(self):    
        self.url = self.download_entry.toPlainText()
        if  self.thread_active==False:
            self.thread_active = True
            self.stops=False 
            self.download_url=self.url
            threading.Thread(target=self.download_ul).start()
        else:
            self.update_notification(self.M24) 
        QApplication.exit()
    def combobox_changed(self,text): 
        self.selected_url = text
        self.download_entry.setText(self.selected_url)
    def paste_downlod(self):
        self.download_entry.insertPlainText(self.root.clipboard_get())
    def download_manager(self):
        try:
            with open(self.url_downloads, 'r', encoding="utf-8") as f:
                 self.url_down=f.read().splitlines() 
            app = QApplication(sys.argv)
            win = QMainWindow()
            win.setGeometry(313, 313, 457, 396)
            win.setWindowTitle(self.M73)
            label2 = QLabel(win)
            label2.setText(self.M240) 
            label2.move(15, 14)
            label2.resize(420, 35) 
            self.download_entry = QTextEdit(win)
            self.download_entry.move(15, 66)
            self.download_entry.resize(420, 150)
            button = QPushButton(win)
            button.setText(self.M73)
            button.move(15, 350)
            button.resize(420, 35)
            button.clicked.connect(self.download_L)
            button_search = QPushButton(win)
            combo = QComboBox(win)
            combo.insertItem(0,'History Downloads')
            combo.setCurrentIndex(0)
            combo.move(15, 240)
            combo.resize(420, 45)
            combo.currentTextChanged.connect(self.combobox_changed)
            #combo.clear()
            combo.addItems(self.url_down)
            button_search.setText(self.M21)
            button_search.move(15, 300)
            button_search.resize(420, 35)
            button_search.clicked.connect(self.paste_downlod)
            win.show()
            app.exec_()
        except Exception as e:
                self.update_notification(f"  {self.M44}  : {e}")
    def download_ul(self, ):
        try:
            download_url=self.download_url
            session = requests.Session()
            parsed_url = urllib.urlparse(download_url)
            if not all([parsed_url.scheme, parsed_url.netloc]):
                self.update_notification(f"  Download Error: Please check internet connection and retry")
            filename = os.path.basename(parsed_url.path)
            save_path = os.path.join(os.path.expanduser("~"), "downloads", filename)
            confirmation = messagebox.askyesno(f"{self.M86}", f"{self.M239} \n {download_url}  .")
            if not confirmation:
                self.update_notification(f"    {self.M78} ")
                self.process_finish()
            else:
                if not os.path.exists(self.url_downloads):
                    with open(self.url_downloads, 'w') as f:
                        f.write('') 
                    with open(self.url_downloads, 'r', encoding="utf-8") as f:
                        self.url_down=f.read().splitlines() 
                self.thread_active = True
                self.stops=False
                self.pack_downloaded = True
                headers = {}
                downloaded = 0 
                if os.path.exists(save_path):
                    downloaded = os.path.getsize(save_path)
                    headers['Range'] = f'bytes={downloaded}-'
                try:
                    response = session.get(download_url, headers=headers, stream=True)
                    content_disposition = response.headers.get('content-disposition')
                    if content_disposition:
                        filename = re.findall('filename=(.+)', content_disposition)[0]
                        filename = filename.replace('"', '')
                        save_path = os.path.join(os.path.expanduser("~"), "downloads", filename)
                    if download_url not in self.url_down:
                        with open(self.url_downloads, 'r', encoding="utf-8") as f:
                            lines = f.readlines()
                        lines.insert(0, download_url + '\n')
                        with open(self.url_downloads, 'w', encoding="utf-8") as f:
                            f.writelines(lines)
                        with open(self.url_downloads, 'r', encoding="utf-8") as f:
                            self.url_down=f.read().split("\n") 
                    total_size = int(response.headers.get('Content-Length', 0)) + downloaded
                    hash_object = hashlib.sha256()  # Change to desired hash function if needed
                    with open(save_path, 'ab') as file:
                        for data in response.iter_content(chunk_size=1048576):
                            if self.stops==True:
                                break
                            file.write(data)
                            downloaded += len(data)
                            total_size_MB = int(total_size / 1048576)
                            percent = (downloaded) / 1048576
                            self.update_notification(f"  {self.M75} {total_size_MB} {self.M73} {percent:.2f}")
                            hash_object.update(data)
                    if downloaded >= total_size:
                        confirmation = messagebox.askyesno(f"{self.M77}", f"{self.M85}  .")
                        if  confirmation:
                            self.update_notification(f"    {self.M77}   {self.M85}  ")
                            os.startfile(save_path)
                    self.process_finish()
                    self.update_notification(f"  {self.M77} ")
                except requests.exceptions.RequestException as e:
                    messagebox.showinfo("download Error", f"Internet connection {str(e)}")
                    self.pack_downloaded = True
                    self.process_finish()
                    self.update_notification(f"  Download Error: Please check internet connection or link address.")
                    #time.sleep(6) 
        except:
            pass   
    def stop(self):
        self.excel=False 
        self.stops=True   
        self.thread_active = False
        self.skip_rtl=False
        self.pdf=False
        self.docx=False
        self.console=False 
        self.pdf_convert=False
        self.update_notification(self.M42)  
        
    def process_finish(self):
        self.pack_downloaded = False
        self.excel=False 
        self.thread_active = False
        self.skip_rtl=False
        self.pdf=False 
        self.docx=False
        self.stops=False
        self.console=False 
        self.pdf_convert=False
        gc.collect()
        self.update_notification(self.M42)  
        
        #self.info2("")
if __name__== "__main__":
    root = tkinter.Tk()
    app = TranslationWindow(root)
    root.protocol("WM_DELETE_WINDOW", on_close)
    root.mainloop()
    "الحمد الله"
    "اللهم صل علي محمد و آل محمد و عجل فرجهم و اهلک والعن اعداهم"



